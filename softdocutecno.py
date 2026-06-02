import streamlit as st
from datetime import datetime, date, time, timedelta
import json
from pathlib import Path
import os
import tempfile

try:
    from docx import Document
    from docx.shared import Inches as DocxInches, Pt as DocxPt, Cm as DocxCm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    Document = None

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    from openpyxl.utils import get_column_letter
except ImportError:
    Workbook = None

try:
    from openai import OpenAI
except ImportError:
    OpenAI = None

try:
    from reportlab.lib.pagesizes import letter, landscape
    from reportlab.pdfgen import canvas
    from reportlab.lib import colors
    from reportlab.lib.utils import ImageReader
    from reportlab.pdfbase.pdfmetrics import stringWidth
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
        Table,
        TableStyle,
        Image,
        PageBreak,
    )
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
    from reportlab.lib.units import cm
except ImportError:
    canvas = None


# =====================================================
# GENERADOR DE DOCUMENTOS TECNOPARQUE / SENA
# =====================================================

st.set_page_config(
    page_title="Generador de Documentos Tecnoparque",
    page_icon="📄",
    layout="wide"
)

# =====================================================
# CONFIGURACIÓN GENERAL
# =====================================================
FORMATO_CODIGO = "GOR-F-084 V02"
FORMATO_CONFIDENCIALIDAD = "GIC-F-041 V03"

LUGAR_ENLACE_DEFAULT = "Tecnoparque Angostura - km 38 vía al sur de Neiva"
DIRECCION_REGIONAL_CENTRO_DEFAULT = "Dirección de formación profesional / HUILA / Centro De Formación Agroindustrial"
DEPENDENCIA_TALENTO_DEFAULT = "EMPRENDEDOR"
DEPENDENCIA_EXPERTO_DEFAULT = "SENA"
ANEXOS_DEFAULT = "NO APLICA"

RUTA_LOGO_SENA = "recursos/logo_sena.png"
CARPETA_FIRMAS = "recursos/firmas"
CARPETA_SALIDA = "documentos_generados"

# =====================================================
# ESTILOS STREAMLIT
# =====================================================
st.markdown(
    """
    <style>
    .main-title {
        font-size: 2rem;
        font-weight: 700;
        color: #2e7d32;
        margin-bottom: 0.2rem;
    }
    .subtitle {
        font-size: 1rem;
        color: #555;
        margin-bottom: 1.5rem;
    }
    .info-box {
        background-color: #f4f9f4;
        border-left: 5px solid #39a935;
        padding: 1rem;
        border-radius: 0.5rem;
        margin-bottom: 1rem;
    }
    div.stButton > button {
        width: 100%;
        min-height: 3.2rem;
        border-radius: 0.7rem;
        font-weight: 600;
        border: 1px solid #39a935;
    }
    </style>
    """,
    unsafe_allow_html=True
)

# =====================================================
# ESTADO DE NAVEGACIÓN
# =====================================================
if "fase_seleccionada" not in st.session_state:
    st.session_state.fase_seleccionada = None

if "documento_seleccionado" not in st.session_state:
    st.session_state.documento_seleccionado = None

if "datos_acta_generada" not in st.session_state:
    st.session_state.datos_acta_generada = None

if "ruta_pdf_generado" not in st.session_state:
    st.session_state.ruta_pdf_generado = None

if "datos_infraestructura_generada" not in st.session_state:
    st.session_state.datos_infraestructura_generada = None

if "ruta_pdf_infraestructura_generado" not in st.session_state:
    st.session_state.ruta_pdf_infraestructura_generado = None

if "datos_confidencialidad_generada" not in st.session_state:
    st.session_state.datos_confidencialidad_generada = None

if "ruta_pdf_confidencialidad_generado" not in st.session_state:
    st.session_state.ruta_pdf_confidencialidad_generado = None

if "ruta_firma_talento_tmp" not in st.session_state:
    st.session_state.ruta_firma_talento_tmp = None
if "datos_cronograma_generado" not in st.session_state:
    st.session_state.datos_cronograma_generado = None

if "ruta_pdf_cronograma_generado" not in st.session_state:
    st.session_state.ruta_pdf_cronograma_generado = None

if "datos_estado_arte_generado" not in st.session_state:
    st.session_state.datos_estado_arte_generado = None

if "ruta_pdf_estado_arte_generado" not in st.session_state:
    st.session_state.ruta_pdf_estado_arte_generado = None

if "datos_acta_ejecucion_generada" not in st.session_state:
    st.session_state.datos_acta_ejecucion_generada = None

if "ruta_pdf_acta_ejecucion_generado" not in st.session_state:
    st.session_state.ruta_pdf_acta_ejecucion_generado = None

if "datos_acta_cierre_generada" not in st.session_state:
    st.session_state.datos_acta_cierre_generada = None

if "ruta_pdf_acta_cierre_generado" not in st.session_state:
    st.session_state.ruta_pdf_acta_cierre_generado = None

if "datos_informe_tecnico_final_generado" not in st.session_state:
    st.session_state.datos_informe_tecnico_final_generado = None

if "ruta_docx_informe_tecnico_final_generado" not in st.session_state:
    st.session_state.ruta_docx_informe_tecnico_final_generado = None

if "datos_lean_canvas_generado" not in st.session_state:
    st.session_state.datos_lean_canvas_generado = None

if "ruta_pdf_lean_canvas_generado" not in st.session_state:
    st.session_state.ruta_pdf_lean_canvas_generado = None


# =====================================================
# FUNCIONES GENERALES
# =====================================================
def obtener_api_key() -> str | None:
    api_key = os.getenv("OPENAI_API_KEY")
    if api_key:
        return api_key
    try:
        return st.secrets.get("OPENAI_API_KEY")
    except Exception:
        return None


def calcular_hora_fin(fecha_acta: date, hora_inicio: time) -> time:
    inicio_dt = datetime.combine(fecha_acta, hora_inicio)
    fin_dt = inicio_dt + timedelta(minutes=30)
    return fin_dt.time()


def construir_objetivo_reunion(nombre_proyecto: str) -> str:
    return f"Fijar el alcance, objetivo general y objetivos específicos del proyecto: {nombre_proyecto}"


def guardar_datos_json(datos: dict, ruta: str = "datos_acta_inicio.json") -> None:
    Path(ruta).write_text(json.dumps(datos, ensure_ascii=False, indent=4), encoding="utf-8")


def limpiar_respuesta_json(texto: str) -> str:
    """
    Limpia respuestas de IA que pueden venir con markdown o texto adicional.
    Extrae el primer objeto JSON válido entre llaves cuando sea posible.
    """
    import re

    if not texto:
        return ""

    texto = str(texto).strip()
    texto = texto.replace("```json", "").replace("```JSON", "").replace("```", "").strip()

    if texto.startswith("{") and texto.endswith("}"):
        return texto

    inicio = texto.find("{")
    fin = texto.rfind("}")
    if inicio != -1 and fin != -1 and fin > inicio:
        return texto[inicio:fin + 1].strip()

    match = re.search(r"\{.*\}", texto, re.DOTALL)
    if match:
        return match.group(0).strip()

    return texto


def seleccionar_fase(nombre_fase: str) -> None:
    st.session_state.fase_seleccionada = nombre_fase
    st.session_state.documento_seleccionado = None


def seleccionar_documento(nombre_documento: str) -> None:
    st.session_state.documento_seleccionado = nombre_documento


def safe_filename(texto: str) -> str:
    limpio = "".join(ch if ch.isalnum() or ch in "-_ " else "" for ch in str(texto))
    limpio = "_".join(limpio.split())
    return limpio[:80] or "documento"


def fecha_larga_espanol(fecha: date) -> str:
    meses = {
        1: "enero",
        2: "febrero",
        3: "marzo",
        4: "abril",
        5: "mayo",
        6: "junio",
        7: "julio",
        8: "agosto",
        9: "septiembre",
        10: "octubre",
        11: "noviembre",
        12: "diciembre",
    }
    return f"{fecha.day} días del mes de {meses[fecha.month]} de {fecha.year}"


def fecha_larga_espanol_con_del(fecha: date) -> str:
    meses = {
        1: "Enero",
        2: "Febrero",
        3: "Marzo",
        4: "Abril",
        5: "Mayo",
        6: "Junio",
        7: "Julio",
        8: "Agosto",
        9: "Septiembre",
        10: "Octubre",
        11: "Noviembre",
        12: "Diciembre",
    }
    return f"{fecha.day} días del mes de {meses[fecha.month]} del {fecha.year}"


def ruta_firma(nombre_archivo: str) -> str:
    return str(Path(CARPETA_FIRMAS) / nombre_archivo)


def guardar_archivo_subido(uploaded_file, prefijo: str) -> str | None:
    if uploaded_file is None:
        return None

    sufijo = Path(uploaded_file.name).suffix.lower()
    if sufijo not in [".png", ".jpg", ".jpeg"]:
        raise ValueError("La firma debe estar en formato PNG, JPG o JPEG.")

    carpeta_tmp = Path(tempfile.gettempdir()) / "firmas_tecnoparque"
    carpeta_tmp.mkdir(parents=True, exist_ok=True)

    ruta = carpeta_tmp / f"{prefijo}_{datetime.now().strftime('%Y%m%d_%H%M%S')}{sufijo}"
    ruta.write_bytes(uploaded_file.getbuffer())
    return str(ruta)


# =====================================================
# FUNCIONES IA - ACTA DE INICIO
# =====================================================
def generar_textos_modo_prueba(nombre_proyecto: str, descripcion_proyecto: str) -> dict:
    return {
        "objetivo_general": (
            f"Desarrollar una propuesta técnica para el proyecto {nombre_proyecto}, orientada a la definición, "
            "estructuración y validación inicial de una solución de base tecnológica que responda a la necesidad identificada."
        ),
        "objetivos_especificos": [
            "Identificar los requerimientos técnicos, funcionales y operativos asociados a la necesidad planteada en la descripción del proyecto.",
            "Definir los componentes principales de la solución tecnológica, considerando los recursos disponibles y las condiciones de uso previstas.",
            "Estructurar una ruta inicial de desarrollo para orientar las actividades de diseño, prototipado, validación básica y documentación técnica.",
            "Validar preliminarmente el alcance del proyecto y los entregables esperados, de acuerdo con la información suministrada."
        ],
        "alcance": (
            f"El alcance del proyecto {nombre_proyecto} contempla la asesoría técnica inicial para la definición de la necesidad, "
            "la identificación de requerimientos, la estructuración de objetivos y la delimitación de actividades orientadas al desarrollo "
            "de una solución de base tecnológica. Incluye lineamientos para diseño, prototipado, validación básica y documentación."
        )
    }


def generar_textos_con_chatgpt(nombre_proyecto: str, descripcion_proyecto: str, modelo: str = "gpt-4.1-mini") -> dict:
    if OpenAI is None:
        raise ImportError("No está instalada la librería openai. Instálala con: pip install openai")

    api_key = obtener_api_key()
    if not api_key:
        raise ValueError("No se encontró OPENAI_API_KEY. Configúrala como variable de entorno o en .streamlit/secrets.toml.")

    client = OpenAI(api_key=api_key)

    instrucciones = """
Eres un experto en formulación de proyectos de base tecnológica, innovación,
desarrollo tecnológico y prototipado en el marco de Tecnoparque SENA.

Redacta en lenguaje formal, técnico e institucional.
No inventes fechas, códigos, nombres de personas, entidades o información no suministrada.
Los objetivos deben iniciar con verbos en infinitivo.
El objetivo general debe ser una sola oración clara.
Los objetivos específicos deben ser exactamente cuatro.
El alcance debe explicar qué incluye el proyecto, qué se desarrollará y cuáles son los entregables esperados.
No uses lenguaje comercial ni promesas no verificables.
Responde únicamente en JSON válido, sin texto antes ni después.
"""

    entrada = f"""
Genera los siguientes campos para un acta de inicio de proyecto:

1. objetivo_general
2. objetivos_especificos: exactamente 4 objetivos específicos
3. alcance

Nombre del proyecto:
{nombre_proyecto}

Descripción general del proyecto:
{descripcion_proyecto}

Formato obligatorio:
{{
  "objetivo_general": "...",
  "objetivos_especificos": ["...", "...", "...", "..."],
  "alcance": "..."
}}
"""

    respuesta = client.responses.create(
        model=modelo,
        instructions=instrucciones,
        input=entrada,
        temperature=0.3
    )

    texto = limpiar_respuesta_json(respuesta.output_text)
    try:
        datos = json.loads(texto)
    except json.JSONDecodeError:
        datos = {
            "objetivo_general": "No se pudo interpretar automáticamente la respuesta de la IA.",
            "objetivos_especificos": [],
            "alcance": texto
        }

    if not isinstance(datos.get("objetivos_especificos", []), list):
        datos["objetivos_especificos"] = []

    return datos


# =====================================================
# FUNCIONES PDF - ACTA DE INICIO
# =====================================================
def wrap_text(texto: str, font_name: str, font_size: float, max_width: float) -> list[str]:
    texto = str(texto or "").replace("\n", " ").strip()
    if not texto:
        return [""]

    palabras = texto.split()
    lineas = []
    linea = ""

    for palabra in palabras:
        intento = palabra if not linea else f"{linea} {palabra}"
        if stringWidth(intento, font_name, font_size) <= max_width:
            linea = intento
        else:
            if linea:
                lineas.append(linea)
            while stringWidth(palabra, font_name, font_size) > max_width and len(palabra) > 5:
                palabra = palabra[:-1]
            linea = palabra

    if linea:
        lineas.append(linea)

    return lineas


def calcular_font_para_celda(texto, w, h, font="Helvetica", size=7.5, min_size=4.8, leading_factor=1.18, label=None):
    padding_x = 5
    padding_y = 5
    available_w = max(w - padding_x * 2, 10)
    available_h = max(h - padding_y * 2, 8)

    current = size
    while current >= min_size:
        if label:
            label_width = stringWidth(label, "Helvetica-Bold", current)
            if label_width >= available_w * 0.85:
                label_lines = len(wrap_text(label, "Helvetica-Bold", current, available_w))
                text_width = available_w
            else:
                label_lines = 0
                text_width = max(available_w - label_width - 3, 10)
        else:
            label_lines = 0
            text_width = available_w

        lineas = wrap_text(texto, font, current, text_width)
        leading = current * leading_factor
        needed_h = max(1, len(lineas) + label_lines) * leading
        if needed_h <= available_h:
            return current, leading
        current -= 0.25

    return min_size, min_size * leading_factor


def draw_wrapped_text(c, texto, x, y, w, h, font="Helvetica", size=7.5, label=None, center=False):
    padding_x = 5
    padding_y = 5
    texto = str(texto or "").strip()
    label_text = str(label or "").strip()

    font_size, leading = calcular_font_para_celda(texto, w, h, font=font, size=size, label=label_text or None)
    cursor_y = y + h - padding_y - font_size
    max_w = w - padding_x * 2

    if center and not label_text:
        lineas = wrap_text(texto, font, font_size, max_w)
        total_h = len(lineas) * leading
        cursor_y = y + (h + total_h) / 2 - font_size
        c.setFont(font, font_size)
        for linea in lineas:
            c.drawCentredString(x + w / 2, cursor_y, linea)
            cursor_y -= leading
        return

    if label_text:
        label_width = stringWidth(label_text, "Helvetica-Bold", font_size)
        if label_width < max_w * 0.80:
            c.setFont("Helvetica-Bold", font_size)
            c.drawString(x + padding_x, cursor_y, label_text)
            c.setFont(font, font_size)
            text_x = x + padding_x + label_width + 3
            text_w = max_w - label_width - 3
            lineas = wrap_text(texto, font, font_size, text_w)
            if lineas:
                c.drawString(text_x, cursor_y, lineas[0])
                cursor_y -= leading
                lineas = lineas[1:]
            for linea in lineas:
                if cursor_y < y + 2:
                    break
                c.drawString(x + padding_x, cursor_y, linea)
                cursor_y -= leading
        else:
            c.setFont("Helvetica-Bold", font_size)
            for linea in wrap_text(label_text, "Helvetica-Bold", font_size, max_w):
                if cursor_y < y + 2:
                    break
                c.drawString(x + padding_x, cursor_y, linea)
                cursor_y -= leading
            c.setFont(font, font_size)
            for linea in wrap_text(texto, font, font_size, max_w):
                if cursor_y < y + 2:
                    break
                c.drawString(x + padding_x, cursor_y, linea)
                cursor_y -= leading
        return

    c.setFont(font, font_size)
    for linea in wrap_text(texto, font, font_size, max_w):
        if cursor_y < y + 2:
            break
        c.drawString(x + padding_x, cursor_y, linea)
        cursor_y -= leading


def draw_cell(c, x, y, w, h, texto="", label=None, font="Helvetica", size=7.5, center=False, fill=None):
    c.setStrokeColor(colors.black)
    c.setLineWidth(0.55)
    if fill:
        c.setFillColor(fill)
        c.rect(x, y, w, h, stroke=1, fill=1)
        c.setFillColor(colors.black)
    else:
        c.rect(x, y, w, h, stroke=1, fill=0)

    draw_wrapped_text(c, texto, x, y, w, h, font=font, size=size, label=label, center=center)


def draw_logo(c, page_width, top_y):
    logo_path = Path(RUTA_LOGO_SENA)

    if logo_path.exists():
        try:
            img = ImageReader(str(logo_path))
            logo_w = 52
            logo_h = 44
            logo_x = page_width / 2 - logo_w / 2
            logo_y = top_y - logo_h
            c.drawImage(img, logo_x, logo_y, width=logo_w, height=logo_h, mask="auto")
            return
        except Exception:
            pass

    c.setFillColor(colors.HexColor("#69B342"))
    c.setFont("Helvetica-Bold", 14)
    c.drawCentredString(page_width / 2, top_y - 18, "SENA")
    c.setFont("Helvetica-Bold", 14)
    c.drawCentredString(page_width / 2, top_y - 36, "▲")
    c.setFillColor(colors.black)


def footer_codigo(c, page_width):
    c.setFillColor(colors.grey)
    c.setFont("Helvetica", 11)
    c.drawCentredString(page_width / 2, 20, FORMATO_CODIGO)
    c.setFillColor(colors.black)


def generar_pdf_acta_inicio(datos: dict) -> str:
    if canvas is None:
        raise ImportError("No está instalada reportlab. Instálala con: pip install reportlab")

    Path(CARPETA_SALIDA).mkdir(exist_ok=True)
    nombre_archivo = f"Acta_Inicio_{safe_filename(datos.get('codigo_acta', 'sin_codigo'))}.pdf"
    ruta_pdf = str(Path(CARPETA_SALIDA) / nombre_archivo)

    c = canvas.Canvas(ruta_pdf, pagesize=letter)
    page_width, page_height = letter

    x0 = 24
    table_w = page_width - 48
    logo_top_y = 785
    y_top_content = 705
    y_safe_bottom = 62

    FONT_TITLE = 8.5
    FONT_SECTION = 8.0
    FONT_BODY = 7.2
    FONT_SMALL = 6.8
    FONT_TINY = 6.3

    def iniciar_pagina() -> float:
        draw_logo(c, page_width, logo_top_y)
        return y_top_content

    def cerrar_pagina() -> None:
        footer_codigo(c, page_width)
        c.showPage()

    def asegurar_espacio(y_actual: float, alto_requerido: float) -> float:
        if y_actual - alto_requerido < y_safe_bottom:
            cerrar_pagina()
            return iniciar_pagina()
        return y_actual

    def alto_texto(texto: str, ancho: float, base: float = 24, font_size: float = FONT_BODY, label: str | None = None, max_h: float = 70) -> float:
        padding = 10
        usable_w = max(ancho - padding, 20)
        texto_total = f"{label or ''} {texto or ''}".strip()
        lineas = wrap_text(texto_total, "Helvetica", font_size, usable_w)
        h = max(base, len(lineas) * (font_size + 2) + 12)
        return min(h, max_h)

    y = iniciar_pagina()

    h = 22
    draw_cell(c, x0, y - h, table_w, h, f"ACTA NO. 1 DEL PROYECTO: {datos.get('codigo_acta', '')}", font="Helvetica-Bold", size=FONT_TITLE, center=True)
    y -= h

    nombre_comite = f"Acta de inicio del proyecto {datos.get('codigo_acta', '')} - {datos.get('nombre_proyecto', '')}"
    h = alto_texto(nombre_comite, table_w, base=42, font_size=FONT_BODY, label="NOMBRE DEL COMITÉ O DE LA REUNIÓN:", max_h=58)
    draw_cell(c, x0, y - h, table_w, h, nombre_comite, label="NOMBRE DEL COMITÉ O DE LA REUNIÓN:", size=FONT_BODY)
    y -= h

    h = 38
    w1 = table_w * 0.58
    w2 = table_w * 0.21
    w3 = table_w * 0.21
    draw_cell(c, x0, y - h, w1, h, f"{datos.get('ciudad', '')} (Huila) - {datos.get('fecha_iso', datos.get('fecha_acta', ''))}", label="CIUDAD Y FECHA:", size=FONT_BODY)
    draw_cell(c, x0 + w1, y - h, w2, h, datos.get("hora_inicio", ""), label="HORA INICIO:", size=FONT_BODY)
    draw_cell(c, x0 + w1 + w2, y - h, w3, h, datos.get("hora_fin", ""), label="HORA FIN:", size=FONT_BODY)
    y -= h

    h = 46
    w_lugar = table_w * 0.58
    w_dir = table_w * 0.42
    draw_cell(c, x0, y - h, w_lugar, h, LUGAR_ENLACE_DEFAULT, label="LUGAR Y/O ENLACE:", size=FONT_BODY)
    draw_cell(c, x0 + w_lugar, y - h, w_dir, h, DIRECCION_REGIONAL_CENTRO_DEFAULT, label="DIRECCIÓN / REGIONAL / CENTRO:", size=FONT_BODY)
    y -= h

    agenda = (
        "1. Caracterización del proyecto de acuerdo con los objetivos y alcance propuestos.\n"
        "2. Documentación que soportan el inicio del proyecto."
    )
    h = 42
    draw_cell(c, x0, y - h, table_w, h, agenda, label="AGENDA O PUNTOS PARA DESARROLLAR:", size=FONT_BODY)
    y -= h

    h = alto_texto(datos.get("objetivo_reunion", ""), table_w, base=44, font_size=FONT_BODY, label="OBJETIVO(S) DE LA REUNIÓN:", max_h=60)
    draw_cell(c, x0, y - h, table_w, h, datos.get("objetivo_reunion", ""), label="OBJETIVO(S) DE LA REUNIÓN:", size=FONT_BODY)
    y -= h

    h = 22
    y = asegurar_espacio(y, h)
    draw_cell(c, x0, y - h, table_w, h, "DESARROLLO DE LA REUNIÓN", font="Helvetica-Bold", size=FONT_SECTION, center=True)
    y -= h

    codigo_nombre = f"{datos.get('codigo_acta', '')} - {datos.get('nombre_proyecto', '')}"
    h = alto_texto(codigo_nombre, table_w, base=26, font_size=FONT_BODY, label="Código y nombre del Proyecto:", max_h=48)
    y = asegurar_espacio(y, h)
    draw_cell(c, x0, y - h, table_w, h, codigo_nombre, label="Código y nombre del Proyecto:", size=FONT_BODY)
    y -= h

    h = alto_texto(datos.get("linea_sublinea", ""), table_w, base=24, font_size=FONT_BODY, label="Linea y sublinea:", max_h=38)
    y = asegurar_espacio(y, h)
    draw_cell(c, x0, y - h, table_w, h, datos.get("linea_sublinea", ""), label="Linea y sublinea:", size=FONT_BODY)
    y -= h

    bloque_talentos = 22 + 22 + 24
    y = asegurar_espacio(y, bloque_talentos)
    draw_cell(c, x0, y - 22, table_w, 22, "TALENTOS QUE PARTICIPAN EN EL PROYECTO", font="Helvetica-Bold", size=FONT_SECTION, center=True)
    y -= 22
    inter_w = table_w * 0.18
    draw_cell(c, x0, y - 22, inter_w, 22, "Interlocutor", font="Helvetica-Bold", size=FONT_SMALL)
    draw_cell(c, x0 + inter_w, y - 22, table_w - inter_w, 22, "Talento", font="Helvetica-Bold", size=FONT_SMALL)
    y -= 22
    draw_cell(c, x0, y - 24, inter_w, 24, "SI", size=FONT_SMALL)
    draw_cell(c, x0 + inter_w, y - 24, table_w - inter_w, 24, datos.get("nombre_talento", ""), size=FONT_SMALL)
    y -= 24

    h = 22
    y = asegurar_espacio(y, h)
    draw_cell(c, x0, y - h, table_w, h, "OBJETIVOS DEL PROYECTO Y ALCANCE", font="Helvetica-Bold", size=FONT_SECTION, center=True)
    y -= h

    h = alto_texto(datos.get("objetivo_general", ""), table_w, base=46, font_size=FONT_BODY, label="OBJETIVO GENERAL:", max_h=70)
    y = asegurar_espacio(y, h)
    draw_cell(c, x0, y - h, table_w, h, datos.get("objetivo_general", ""), label="OBJETIVO GENERAL:", size=FONT_BODY)
    y -= h

    h = 21
    y = asegurar_espacio(y, h)
    draw_cell(c, x0, y - h, table_w, h, "OBJETIVOS ESPECÍFICOS", font="Helvetica-Bold", size=FONT_SECTION, center=True)
    y -= h

    objetivos = datos.get("objetivos_especificos", [])[:4]
    while len(objetivos) < 4:
        objetivos.append("")

    num_w = 52
    obj_w = table_w - num_w
    for idx, obj in enumerate(objetivos, start=1):
        h = alto_texto(obj, obj_w, base=24, font_size=FONT_SMALL, max_h=42)
        y = asegurar_espacio(y, h)
        draw_cell(c, x0, y - h, num_w, h, str(idx), font="Helvetica-Bold", size=FONT_SMALL, center=True)
        draw_cell(c, x0 + num_w, y - h, obj_w, h, obj, size=FONT_SMALL)
        y -= h

    h = alto_texto(datos.get("alcance", ""), table_w, base=36, font_size=FONT_TINY, label="ALCANCE DEL PROYECTO:", max_h=82)
    y = asegurar_espacio(y, h)
    draw_cell(c, x0, y - h, table_w, h, datos.get("alcance", ""), label="ALCANCE DEL PROYECTO:", size=FONT_TINY)
    y -= h

    bloque_conclusiones = 22 + 28
    y = asegurar_espacio(y, bloque_conclusiones)
    draw_cell(c, x0, y - 22, table_w, 22, "CONCLUSIONES", font="Helvetica-Bold", size=FONT_SECTION, center=True)
    y -= 22
    conclusion = "Se fijan alcance y objetivos entre el talento y experto participantes de la reunión."
    draw_cell(c, x0, y - 28, table_w, 28, conclusion, size=FONT_BODY)
    y -= 28

    asistentes_alto = 22 + 34 + 34 + 34
    y = asegurar_espacio(y, asistentes_alto)
    draw_cell(c, x0, y - 22, table_w, 22, "DE: ASISTENTES Y APROBACIÓN DE DECISIONES:", font="Helvetica-Bold", size=FONT_SECTION, center=True)
    y -= 22

    col_w = [table_w * 0.18, table_w * 0.17, table_w * 0.20, table_w * 0.22, table_w * 0.23]
    headers = ["NOMBRE", "DEPENDENCIA / EMPRESA", "APRUEBA (SI/NO)", "OBSERVACIÓN", "FIRMA O PARTICIPACIÓN VIRTUAL"]
    x = x0
    for w, h_text in zip(col_w, headers):
        draw_cell(c, x, y - 34, w, 34, h_text, font="Helvetica-Bold", size=FONT_SMALL, center=True)
        x += w
    y -= 34

    filas = [
        [datos.get("nombre_talento", ""), DEPENDENCIA_TALENTO_DEFAULT, "SI", "", ""],
        [datos.get("nombre_experto", ""), DEPENDENCIA_EXPERTO_DEFAULT, "SI", "", ""],
    ]

    for fila in filas:
        x = x0
        for w, value in zip(col_w, fila):
            center = value in [DEPENDENCIA_TALENTO_DEFAULT, DEPENDENCIA_EXPERTO_DEFAULT, "SI", ""]
            draw_cell(c, x, y - 34, w, 34, value, font="Helvetica-Bold" if x == x0 else "Helvetica", size=FONT_SMALL, center=center)
            x += w
        y -= 34

    bloque_final_alto = 55 + 32
    y = asegurar_espacio(y, bloque_final_alto)

    proteccion = (
        "De acuerdo con La Ley 1581 de 2012, Protección de Datos Personales, el Servicio Nacional de Aprendizaje SENA, "
        "se compromete a garantizar la seguridad y protección de los datos personales que se encuentran almacenados en este documento, "
        "y les dará el tratamiento correspondiente en cumplimiento de lo establecido legalmente."
    )
    draw_cell(c, x0, y - 55, table_w, 55, proteccion, size=FONT_BODY)
    y -= 55

    texto_anexos = "ANEXOS\n" + ANEXOS_DEFAULT
    draw_cell(c, x0, y - 32, table_w, 32, texto_anexos, font="Helvetica-Bold", size=FONT_SECTION, center=True)

    footer_codigo(c, page_width)
    c.save()
    return ruta_pdf


# =====================================================
# PDF - USO DE INFRAESTRUCTURA
# =====================================================
def generar_pdf_uso_infraestructura(datos: dict) -> str:
    if canvas is None:
        raise ImportError("No está instalada reportlab. Instálala con: pip install reportlab")

    from html import escape

    Path(CARPETA_SALIDA).mkdir(parents=True, exist_ok=True)

    nombre_archivo = (
        f"Uso_Infraestructura_"
        f"{safe_filename(datos.get('codigo_proyecto', 'proyecto'))}_"
        f"{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    )
    ruta_pdf = str(Path(CARPETA_SALIDA) / nombre_archivo)

    page_width, page_height = letter

    def encabezado_logo(c, doc):
        c.saveState()
        logo_path = Path(RUTA_LOGO_SENA)
        if logo_path.exists():
            try:
                logo = ImageReader(str(logo_path))
                logo_w = 58
                logo_h = 58
                c.drawImage(
                    logo,
                    (page_width - logo_w) / 2,
                    page_height - 92,
                    width=logo_w,
                    height=logo_h,
                    preserveAspectRatio=True,
                    mask="auto",
                )
            except Exception:
                c.setFillColor(colors.HexColor("#39a935"))
                c.setFont("Helvetica-Bold", 18)
                c.drawCentredString(page_width / 2, page_height - 60, "SENA")
                c.setFillColor(colors.black)
        else:
            c.setFillColor(colors.HexColor("#39a935"))
            c.setFont("Helvetica-Bold", 18)
            c.drawCentredString(page_width / 2, page_height - 60, "SENA")
            c.setFillColor(colors.black)

        c.restoreState()

    doc = SimpleDocTemplate(
        ruta_pdf,
        pagesize=letter,
        rightMargin=3.0 * cm,
        leftMargin=3.0 * cm,
        topMargin=3.4 * cm,
        bottomMargin=2.2 * cm,
    )

    estilo_titulo = ParagraphStyle(
        name="Titulo",
        fontName="Helvetica-Bold",
        fontSize=11,
        leading=14,
        alignment=TA_CENTER,
        spaceAfter=14,
    )

    estilo_normal = ParagraphStyle(
        name="NormalJustificado",
        fontName="Helvetica",
        fontSize=10.5,
        leading=14.5,
        alignment=TA_JUSTIFY,
        spaceAfter=8,
    )

    estilo_normal_sin_espacio = ParagraphStyle(
        name="NormalSinEspacio",
        fontName="Helvetica",
        fontSize=10.5,
        leading=14.5,
        alignment=TA_JUSTIFY,
        spaceAfter=3,
    )

    estilo_negrita = ParagraphStyle(
        name="Negrita",
        fontName="Helvetica-Bold",
        fontSize=10.5,
        leading=14.5,
        alignment=TA_LEFT,
        spaceBefore=8,
        spaceAfter=2,
    )

    estilo_tabla = ParagraphStyle(
        name="Tabla",
        fontName="Helvetica",
        fontSize=9.5,
        leading=11.5,
        alignment=TA_LEFT,
    )

    estilo_tabla_centro = ParagraphStyle(
        name="TablaCentro",
        fontName="Helvetica",
        fontSize=9.5,
        leading=11.5,
        alignment=TA_CENTER,
    )

    historia = []

    ciudad = escape(str(datos.get("ciudad", "")).strip())
    fecha_documento = datos.get("fecha_documento")

    if isinstance(fecha_documento, date):
        fecha_texto = fecha_larga_espanol(fecha_documento)
        fecha_corta = fecha_documento.strftime("%d/%m/%Y")
        fecha_iso = fecha_documento.strftime("%Y-%m-%d")
    else:
        fecha_texto = escape(str(datos.get("fecha_texto", "")).strip())
        fecha_corta = escape(str(datos.get("fecha_corta", "")).strip())
        fecha_iso = escape(str(datos.get("fecha_iso", "")).strip())

    codigo_proyecto = escape(str(datos.get("codigo_proyecto", "")).strip())
    nombre_proyecto = escape(str(datos.get("nombre_proyecto", "")).strip())
    nombre_talento = escape(str(datos.get("nombre_talento", "")).strip())
    telefono_talento = escape(str(datos.get("telefono_talento", "")).strip())
    nombre_experto = escape(str(datos.get("nombre_experto", "")).strip())
    linea_experto = escape(str(datos.get("linea_experto", "")).strip())

    historia.append(
        Paragraph(
            "MANUAL DE PRESTAMO Y USO DE INFRAESTRUCTURA RED TECNOPARQUE<br/>COLOMBIA",
            estilo_titulo,
        )
    )

    historia.append(Paragraph(f"En la Ciudad de {ciudad} a los {fecha_texto}", estilo_normal))

    historia.append(
        Paragraph(
            f'Luego de aceptado el Proyecto <b>“{codigo_proyecto} {nombre_proyecto}”.</b>',
            estilo_normal,
        )
    )

    historia.append(
        Paragraph(
            "Con el fin de brindar un mejor servicio y asegurar un uso apropiado de la "
            "infraestructura de la Red Tecnoparque del SENA, usted como Talento y Gestor "
            "Tecnoparque, deberá tener en cuenta los siguientes puntos y cumplirlos "
            "respectivamente: Se leen, socializan y comprenden.",
            estilo_normal,
        )
    )

    puntos_infraestructura = [
        "Todo laboratorio de la Red Tecnoparque debe tener un manual de normas y comportamientos básicos sin importar el tipo de laboratorio que sea, como, por ejemplo: no ingresar con alimentos y bebidas al interior de los laboratorios, de ser necesario utilizar elementos de sonido mantener el volumen adecuado, obrar con honestidad, respeto, responsabilidad, mantener un tono de voz adecuado, etc., y las normas que cada gestor considere necesarias para una armonía de trabajo.",
        "Antes de comenzar a utilizar los laboratorios, los Gestores deben realizar una capacitación o transferencia básica para el cuidado y manejo de los laboratorios. En el interior de cada laboratorio debe haber un manual o instructivo de las condiciones de operación de los equipos con que el talento necesite trabajar, y que al mismo tiempo informe sobre la indumentaria adecuada para la operación de los equipos de ser necesaria (Utilizar los elementos de seguridad dispuestos para la operación de los equipos).",
        "Cada gestor debe indicar el estado de los equipos y herramientas antes de que sean utilizados por los Talentos y verificar el estado de los mismos una vez el talento termine o entregue los equipos y herramientas.",
        "Cuando se solicite alguna herramienta de corte o se ingrese a un laboratorio para manipular equipos que requieren un manejo especial, se debe usar todos los equipos de protección personal según el caso, y se debe presentar sin excepción el carné de la EPS actualizado a la fecha de uso, el cual debe ser entregado al gestor encargado durante su trabajo con los equipos, este será devuelto al regresar las herramientas.",
        "Siempre que una herramienta o equipo le sea prestado, este deberá ser registrado por uno de los asesores en el formato común de préstamos, en ningún momento el Talento debe recibir o entregar un equipo o herramienta sin que sea registrado el préstamo o devolución de la misma, ya que se asumirá que aún no lo ha devuelto.",
        "Las herramientas se deben regresar al finalizar el día, esto quiere decir que toda herramienta y equipo debe ser entregado antes de las 5:30 p.m, si no es entregada en este horario debe ser reportada por el Talento al día siguiente antes de las 9:00 am, y dejar anotación en la minuta de la empresa de seguridad del Nodo. Si por razones extraordinarias, los equipos o herramientas salen del Nodo o del Centro para un acompañamiento a los Proyectos, cada Gestor encargado debe gestionar los Seguros correspondientes con el Almacén del Centro padrino SENA e informar al Cuentadante del equipo.",
        "El Talento se responsabilizará de lo que ocurra con los equipos y herramientas, al incumplir con la hora de entrega acordada de los equipos y herramientas, esta responsabilidad va desde cubrir gastos de reparación, hasta costos de reposición por pérdidas. Toda herramienta y equipo se entrega en perfecto estado de funcionamiento, si es entregada por el gestor, en mal estado debe ser reportado por el Talento de inmediato, de lo contrario tendrá que asumir los daños no reportados.",
        "Si un equipo o herramienta sufre daños por mal uso durante el tiempo de préstamo, este daño debe ser asumido por el Talento y serán suspendidos los servicios de Tecnoparque del SENA al proyecto mientras el daño no sea cubierto por el Talento.",
        "Al terminar de usar las herramientas, equipos, laboratorios o infraestructura en general, debe quedar limpio, ordenado y en buen estado, en las condiciones como se recibieron.",
        "Las herramientas no se deben retirar del piso, del ambiente o laboratorio al cual corresponden, de ser retirados deben tener la previa autorización del Cuentadante del bien y Gestor encargado.",
        "Si existe algún convenio, alianza, carta de intención o de cooperación con otro Nodo, Centro de formación u otra institución para uso de infraestructura compartida, se debe soportar por escrito el uso de los equipos y herramientas para beneficio de los usuarios o Talentos. Los tiempos y espacios de trabajo deben quedar pactados y oficialmente definidos con horarios y tiempos de uso. Así como los responsables por daños o deterioro de la infraestructura.",
    ]

    for i, punto in enumerate(puntos_infraestructura, start=1):
        tabla_punto = Table(
            [[Paragraph(f"{i}.", estilo_normal_sin_espacio), Paragraph(escape(punto), estilo_normal_sin_espacio)]],
            colWidths=[1.0 * cm, 14.0 * cm],
        )
        tabla_punto.setStyle(
            TableStyle(
                [
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 0),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 3),
                    ("TOPPADDING", (0, 0), (-1, -1), 0),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
                ]
            )
        )
        historia.append(tabla_punto)

    historia.append(Spacer(1, 4))
    historia.append(Paragraph("En cuanto al software:", estilo_normal))

    puntos_software = [
        "Todo equipo que registre un software ilegal será asociado con el Talento que tenía bajo préstamo ese equipo según la fecha y la hora de préstamo, lo que hace responsable al Talento de asumir todos los perjuicios y efectos legales que esta acción conlleve.",
        "Aquellos equipos que contengan archivos peligrosos y material escandaloso serán asociados con el Talento según la hora y fecha de creación, lo que dará la suspensión de los servicios de Tecnoparque del SENA para el proyecto.",
        "La Red Tecnoparque no se responsabiliza de los archivos y documentos guardados en los equipos de cómputo ni el uso que a ellos den otros Talentos.",
        "El cambiar contraseñas de equipos o crear nuevas sesiones dará suspensión de los servicios de Tecnoparque para el proyecto.",
        "Se prohíbe el ingreso a páginas cuyo fin sea la pornografía, la estafa o el ocio.",
    ]

    for i, punto in enumerate(puntos_software, start=1):
        tabla_punto = Table(
            [[Paragraph(f"{i}.", estilo_normal_sin_espacio), Paragraph(escape(punto), estilo_normal_sin_espacio)]],
            colWidths=[1.0 * cm, 14.0 * cm],
        )
        tabla_punto.setStyle(
            TableStyle(
                [
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 0),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 3),
                    ("TOPPADDING", (0, 0), (-1, -1), 0),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
                ]
            )
        )
        historia.append(tabla_punto)

    historia.append(Paragraph("EXCEPCIÓN:", estilo_negrita))

    excepcion = (
        "Aquellos Talentos que requieran trabajar en el desarrollo de sus proyectos entre semana después "
        "de las 5:30 pm y los fines de semana, deberá contar con la respectiva autorización por escrito "
        "del subdirector del Centro, del Dinamizador del Nodo y Experto a cargo con el visto bueno del "
        "coordinador administrativo y previo aviso con 4 días de anterioridad."
    )

    tabla_excepcion = Table(
        [[Paragraph("1.", estilo_normal_sin_espacio), Paragraph(escape(excepcion), estilo_normal_sin_espacio)]],
        colWidths=[1.0 * cm, 14.0 * cm],
    )
    tabla_excepcion.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 3),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    historia.append(tabla_excepcion)

    historia.append(Paragraph("Para constancia, se firman y comprometen al cumplimiento.", estilo_normal))

    tabla_firmas_data = [
        [
            Paragraph("<b>N</b>", estilo_tabla_centro),
            Paragraph("<b>Nombre del Talento</b>", estilo_tabla_centro),
            Paragraph("<b>Teléfono.</b>", estilo_tabla_centro),
            Paragraph("<b>Firma</b>", estilo_tabla_centro),
        ],
        [
            Paragraph("1", estilo_tabla_centro),
            Paragraph(nombre_talento, estilo_tabla),
            Paragraph(telefono_talento, estilo_tabla_centro),
            Paragraph("", estilo_tabla_centro),
        ],
        [
            Paragraph("<b>N</b>", estilo_tabla_centro),
            Paragraph("<b>Nombre del Experto encargado</b>", estilo_tabla_centro),
            Paragraph("<b>Línea</b>", estilo_tabla_centro),
            Paragraph("<b>Firma</b>", estilo_tabla_centro),
        ],
        [
            Paragraph("1", estilo_tabla_centro),
            Paragraph(nombre_experto, estilo_tabla),
            Paragraph(linea_experto, estilo_tabla_centro),
            Paragraph("", estilo_tabla_centro),
        ],
    ]

    tabla_firmas = Table(
        tabla_firmas_data,
        colWidths=[1.0 * cm, 6.0 * cm, 3.8 * cm, 4.5 * cm],
        rowHeights=[0.65 * cm, 1.15 * cm, 0.65 * cm, 1.25 * cm],
    )

    tabla_firmas.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.6, colors.black),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("ALIGN", (0, 0), (0, -1), "CENTER"),
                ("ALIGN", (2, 0), (2, -1), "CENTER"),
                ("ALIGN", (3, 0), (3, -1), "CENTER"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]
        )
    )

    historia.append(Spacer(1, 8))
    historia.append(tabla_firmas)

    doc.build(
        historia,
        onFirstPage=encabezado_logo,
        onLaterPages=encabezado_logo,
    )

    datos_json = {
        "tipo_documento": "Uso de infraestructura",
        "codigo_proyecto": datos.get("codigo_proyecto", ""),
        "nombre_proyecto": datos.get("nombre_proyecto", ""),
        "nombre_talento": datos.get("nombre_talento", ""),
        "telefono_talento": datos.get("telefono_talento", ""),
        "nombre_experto": datos.get("nombre_experto", ""),
        "linea_experto": datos.get("linea_experto", ""),
        "ciudad": datos.get("ciudad", ""),
        "fecha_corta": fecha_corta,
        "fecha_iso": fecha_iso,
        "ruta_pdf": ruta_pdf,
    }

    guardar_datos_json(datos_json, ruta="datos_uso_infraestructura.json")

    return ruta_pdf


# =====================================================
# PDF - CONFIDENCIALIDAD Y COMPROMISO
# =====================================================

def obtener_ruta_logo_sena() -> str | None:
    posibles_rutas = [
        Path(RUTA_LOGO_SENA),
        Path(CARPETA_FIRMAS) / "logo_sena.png",
        Path(CARPETA_FIRMAS) / "logo_sena.jpg",
        Path(CARPETA_FIRMAS) / "logo_sena.jpeg",
    ]

    for ruta in posibles_rutas:
        if ruta.exists():
            return str(ruta)

    return None


def crear_parrafo(texto: str, estilo: ParagraphStyle):
    return Paragraph(str(texto).replace("\n", "<br/>"), estilo)


def firma_img(path: str | None, width: float = 4.0 * cm, height: float = 1.05 * cm):
    if path and Path(path).exists():
        try:
            return Image(path, width=width, height=height)
        except Exception:
            return Paragraph("", ParagraphStyle(name="vacio", fontSize=8))
    return Paragraph("", ParagraphStyle(name="vacio", fontSize=8))


def generar_pdf_confidencialidad(datos: dict) -> str:
    if canvas is None:
        raise ImportError("No está instalada reportlab. Instálala con: pip install reportlab")

    Path(CARPETA_SALIDA).mkdir(parents=True, exist_ok=True)

    nombre_archivo = (
        f"Confidencialidad_Compromiso_"
        f"{safe_filename(datos.get('codigo_proyecto', 'proyecto'))}_"
        f"{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    )
    ruta_pdf = str(Path(CARPETA_SALIDA) / nombre_archivo)

    fecha_documento = datos["fecha_documento"]
    ciudad = datos["ciudad"]
    nombre_proyecto = datos["nombre_proyecto"]
    codigo_proyecto = datos["codigo_proyecto"]
    nombre_talento = datos["nombre_talento"]
    cedula_talento = datos["cedula_talento"]
    ciudad_expedicion = datos["ciudad_expedicion"]
    ruta_firma_talento = datos.get("ruta_firma_talento")

    page_width, page_height = letter

    def encabezado_pie(c, doc):
        c.saveState()

        # Logo SENA centrado en encabezado
        ruta_logo = obtener_ruta_logo_sena()
        if ruta_logo:
            try:
                logo = ImageReader(ruta_logo)
                logo_w = 2.2 * cm
                logo_h = 2.2 * cm
                c.drawImage(
                    logo,
                    (page_width - logo_w) / 2,
                    page_height - 2.6 * cm,
                    width=logo_w,
                    height=logo_h,
                    preserveAspectRatio=True,
                    mask="auto",
                )
            except Exception:
                c.setFillColor(colors.HexColor("#39a935"))
                c.setFont("Helvetica-Bold", 16)
                c.drawCentredString(page_width / 2, page_height - 1.7 * cm, "SENA")
                c.setFillColor(colors.black)
        else:
            c.setFillColor(colors.HexColor("#39a935"))
            c.setFont("Helvetica-Bold", 16)
            c.drawCentredString(page_width / 2, page_height - 1.7 * cm, "SENA")
            c.setFillColor(colors.black)

        # Código únicamente en pie de página centrado
        c.setFont("Helvetica", 9)
        c.setFillColor(colors.grey)
        c.drawCentredString(page_width / 2, 1.1 * cm, FORMATO_CONFIDENCIALIDAD)
        c.setFillColor(colors.black)

        c.restoreState()

    doc = SimpleDocTemplate(
        ruta_pdf,
        pagesize=letter,
        rightMargin=2.35 * cm,
        leftMargin=2.35 * cm,
        topMargin=3.4 * cm,
        bottomMargin=2.2 * cm,
    )

    estilo_normal = ParagraphStyle(
        name="NormalConfidencialidad",
        fontName="Helvetica",
        fontSize=10,
        leading=14,
        alignment=TA_JUSTIFY,
        spaceAfter=8,
    )

    estilo_titulo = ParagraphStyle(
        name="TituloConfidencialidad",
        fontName="Helvetica-Bold",
        fontSize=11,
        leading=15,
        alignment=TA_CENTER,
        spaceAfter=12,
    )

    estilo_subtitulo = ParagraphStyle(
        name="SubtituloConfidencialidad",
        fontName="Helvetica-Bold",
        fontSize=10.5,
        leading=14,
        alignment=TA_CENTER,
        spaceBefore=8,
        spaceAfter=8,
    )

    estilo_tabla = ParagraphStyle(
        name="TablaConfidencialidad",
        fontName="Helvetica",
        fontSize=8.8,
        leading=10.5,
        alignment=TA_LEFT,
    )

    estilo_tabla_negrita = ParagraphStyle(
        name="TablaConfidencialidadNegrita",
        fontName="Helvetica-Bold",
        fontSize=8.8,
        leading=10.5,
        alignment=TA_CENTER,
    )

    historia = []

    historia.append(
        crear_parrafo(
            "PROCESO GESTIÓN DE INNOVACIÓN Y COMPETITIVIDAD<br/>"
            "FORMATO DE CONFIDENCIALIDAD Y COMPROMISOS RED TECNOPARQUE",
            estilo_titulo,
        )
    )

    texto_intro = (
        f"En la ciudad de {ciudad} a los {fecha_larga_espanol_con_del(fecha_documento)}, "
        "se celebra la presente Acta de Confidencialidad y Compromisos entre la Red Tecnoparque "
        "SENA Nodo Angostura representado por los firmantes abajo en este documento, y por otra "
        f"parte el Titular del Proyecto {nombre_talento}, identificado con cedula N° {cedula_talento} "
        f"de {ciudad_expedicion}, quien en adelante se denominará Titular y talento interlocutor del "
        f"proyecto {codigo_proyecto} - {nombre_proyecto} previas las siguientes consideraciones:"
    )
    historia.append(crear_parrafo(texto_intro, estilo_normal))

    historia.append(crear_parrafo("CONSIDERACIONES", estilo_subtitulo))
    historia.append(
        crear_parrafo(
            "Para dar inicio a la etapa de desarrollo del proyecto de base tecnológica enunciado "
            "anteriormente, es necesario establecer acuerdos entre las partes en las siguientes "
            "materias: i) el manejo de la información confidencial, ii) los derechos de propiedad "
            "intelectual, iii) los compromisos entre la Red TecnoParque SENA y la(s)/el(los) Titular(es) "
            "del proyecto, acorde con lo dispuesto por los Acuerdos 09 de 2010 y 03 de 2012 del Consejo "
            "Directivo Nacional y la Guía Metodológica Red TecnoParque Colombia, o las normas y "
            "documentos que los sustituyan.",
            estilo_normal,
        )
    )

    historia.append(crear_parrafo("En mérito de lo expuesto se", estilo_normal))
    historia.append(crear_parrafo("ACUERDA:", estilo_subtitulo))

    historia.append(crear_parrafo("CAPITULO I<br/>DE LA CONFIDENCIALIDAD", estilo_subtitulo))

    historia.append(
        crear_parrafo(
            "<b>PRIMERO. INFORMACIÓN CONFIDENCIAL.</b> De conformidad con lo dispuesto por el "
            "Acuerdo 03 de 2012, constituirá “Información Confidencial” las metodologías, tecnologías, "
            "planos, prototipos, programas de computador y propiedad intelectual e ideas de la(s)/el(los) "
            "Titular(es)) del proyecto. Esto es, para mayor detalle sin limitarse a lo enunciado: las obras "
            "protegidas por el derecho de autor, nuevas creaciones o signos distintivos objeto de propiedad "
            "industrial, técnicas, modelos, invenciones, know-how, procesos, algoritmos, programas, ejecutables, "
            "investigaciones, detalles de diseño, información financiera, lista de clientes, bases de datos, "
            "inversionistas, empleados, relaciones de negocios y contractuales, pronósticos de negocios, planes "
            "de mercadeo. Se considera igualmente información confidencial, a) cualquier información revelada, "
            "sobre terceras personas y que no sea de dominio público u obvio, antes de la firma de la presente "
            "acta, b) la que no sea de fácil acceso, c) aquella información que no esté sujeta a medidas de "
            "protección razonables, de acuerdo con las circunstancias del caso, a fin de mantener su carácter "
            "confidencial. Los resultados del proyecto pueden considerarse confidenciales, si la(s)/el(los) "
            "Titular(es) así lo define(n).",
            estilo_normal,
        )
    )

    historia.append(
        crear_parrafo(
            "<b>SEGUNDO. MANEJO DE LA INFORMACIÓN CONFIDENCIAL.</b> Las partes acuerdan que "
            "cualquier información confidencial intercambiada, facilitada o creada entre ellas durante el "
            "desarrollo del proyecto, será mantenida en estricta reserva. El/la experto/a, Dinamizador/a, "
            "Infocenter y en general la Red Tecnoparque sólo podrá revelar información confidencial a quienes "
            "la necesiten y estén autorizados previamente por la(s)/el(los) Titular(es) que firman este documento. "
            "Así mismo, la(s)/el(los) Titular(es) que se incorpora(n) a la Red TecnoParque, deberán mantener en "
            "total reserva la información confidencial obtenida de otros Titulares, Experto, etc.",
            estilo_normal,
        )
    )

    historia.append(
        crear_parrafo(
            "En el caso de que la Red TecnoParque SENA requiera usar información catalogada como "
            "confidencial para el desarrollo de otros proyectos, deberá ser autorizado por escrito por "
            "la(s)/el(los) Titular(es) propietaria(os) de esta información.",
            estilo_normal,
        )
    )

    historia.append(
        crear_parrafo(
            "El SENA podrá hacer uso de los resultados de los proyectos, cuando ello sea necesario o "
            "requerido por los procesos de formación profesional, respetando siempre los derechos de "
            "propiedad intelectual sobre los resultados y la reserva de la información confidencial.",
            estilo_normal,
        )
    )

    historia.append(
        crear_parrafo(
            "<b>TERCERO. EXCEPCIONES.</b> Lo datos intercambiados entre las partes no constituyen "
            "información confidencial en los siguientes casos: a) Cuando la parte receptora tenga "
            "evidencia de que conoce previamente la información recibida; b) Cuando la información "
            "recibida sea de dominio público; c) Cuando la información es revelada por el propietario y "
            "este acepta que puede ser utilizada como información de dominio público.",
            estilo_normal,
        )
    )

    historia.append(
        crear_parrafo(
            "<b>CUARTO. DURACIÓN.</b> Las condiciones para el manejo de información confidencial que "
            "asume la(s)/el(los) Titular(es) a través de este acuerdo regirán durante el tiempo que dure "
            "el desarrollo del proyecto y cinco (5) años más, en ausencia de un acuerdo diferente entre "
            "la(s)/el(los) Titular(es) del proyecto y el Nodo de la Red Tecnoparque.",
            estilo_normal,
        )
    )

    historia.append(crear_parrafo("CAPITULO II<br/>DE LOS DERECHOS DE PROPIEDAD INTELECTUAL", estilo_subtitulo))

    textos_capitulo_ii = [
        "<b>QUINTO. TITULAR DEL PROYECTO.</b> Para los efectos del presente acuerdo y la prestación de servicios de Red Tecnoparque Colombia, se entiende por Titular del proyecto la(s) persona(s) u entidad(es) que lo idea(n), formula(n) y trabaja(n) en su implementación y son propietario(s) de los derechos patrimoniales derivados.",
        "<b>SEXTO. INTERLOCUTOR DEL PROYECTO.</b> Para los efectos del presente acuerdo y la prestación de servicios de Red Tecnoparque Colombia, se entiende por Interlocutor del proyecto la persona actúa en representación, manejo de las comunicaciones y toma de decisiones con la Red Tecnoparque para el (los) Titular(es) del proyecto.",
        "<b>SÉPTIMO. EJECUTOR DEL PROYECTO.</b> Para los efectos del presente acuerdo y la prestación de servicios de Red Tecnoparque Colombia, se entiende por Ejecutor del proyecto la(s) persona(s) que formula(n), trabaja(n) y apoya(n) en la implementación del proyecto con o sin relación contractual con el tercero, persona natural o jurídica, Titular del proyecto.",
        "<b>OCTAVO. TITULARIDAD DE LOS DERECHOS DE PROPIEDAD INTELECTUAL.</b> De conformidad con el Acuerdo 09 de 2010, por el cual se establecen las políticas para el programa de TecnoAcademias y TecnoParque, Capítulo II de los TecnoParque, los derechos de propiedad intelectual del proyecto desarrollado en la Red TecnoParque serán de sus autores/inventores, es decir, de la(s)/el(los) Titular(es) del proyecto, de conformidad con las normas vigentes que regulan la materia.",
        "Es responsabilidad de la(s)/el(los) Titular(es) iniciar los procesos de protección de la propiedad industrial de los productos, procesos o diseños que resulte de su proyecto, si considera que son susceptibles de algún mecanismo de protección nacional o internacional.",
        "El manejo de derechos de autor respecto del Interlocutor/a y ejecutor/a, se realizará en los siguientes términos:",
        "En los casos en que resulten del proyecto obras susceptibles de protección de derechos de autor, acorde con las disposiciones que regulan la materia, el/la Interlocutor/a del proyecto y/o ejecutor/a, será(n) titular(es) de los derechos morales de autor, siempre que su contribución a la obra no haya sido puramente física o mecánica, es decir, en los casos en que pueda predicarse su calidad de coautor(es)/a(s). La titularidad de los derechos patrimoniales de autor, por su parte, obedecerá a lo pactado en la relación contractual entre el/la Interlocutor/a y la persona natural o jurídica contratante que representa; en ausencia de pacto, la titularidad de los derechos patrimoniales se definirá en los términos que se acuerde entre el/la Interlocutor/a y/o el/la ejecutor/a y la persona natural o jurídica Titular del proyecto, o por lo establecido en el artículo 28 de la Ley 1450 de 2011 que modifica el artículo 20 de la Ley 23 de 1982.",
    ]

    for texto in textos_capitulo_ii:
        historia.append(crear_parrafo(texto, estilo_normal))

    historia.append(crear_parrafo("CAPITULO III<br/>DE LOS COMPROMISOS Y EL DESARROLLO DE PROYECTOS Y PROTOTIPOS", estilo_subtitulo))

    historia.append(
        crear_parrafo(
            "<b>NOVENO. COMPROMISOS.</b> El desarrollo de proyectos de base tecnológica en la Red "
            "TecnoParque SENA conlleva asumir los compromisos que se enuncian en el presente "
            "documento, encaminados a optimizar el tiempo, los recursos invertidos en el desarrollo del "
            "proyecto, así como sus resultados y beneficios.",
            estilo_normal,
        )
    )

    historia.append(
        crear_parrafo(
            "<b>DÉCIMO. COMPROMISOS DE LA RED TECNOPARQUE SENA.</b> Es responsable de ofrecer "
            "sin ningún costo, asesoría técnica especializada y personalizada, herramientas e infraestructura "
            "necesaria para el desarrollo de iniciativas novedosas de productos y servicios de base tecnológica, "
            "susceptible de ser materializada en prototipos funcionales, ofreciendo adicionalmente:",
            estilo_normal,
        )
    )

    compromisos_red = [
        "Orientación sobre entidades de fortalecimiento empresarial y financiación.",
        "Acceso y uso de la infraestructura tecnológica en los horarios de servicio establecidos por cada Nodo.",
        "Oportunidades para participar en diferentes eventos como ferias, transferencia de tecnología, talleres y seminarios técnicos, encuentros tecnológicos, muestras empresariales, ruedas de negocios, entre otros, teniendo en cuenta los parámetros de selección que defina la Red Tecnoparque y el SENA.",
        "Cumplimiento del cronograma de trabajo definido entre los Titulares y los Experto de la Red TecnoParque SENA, en donde los Experto de la Red, cumplen con el servicio de asesoría técnica especializada y personalizada a los proyectos.",
        "Ofrecer el servicio de acceso a laboratorios en óptimas condiciones, garantizando el buen uso de la infraestructura.",
        "Contar con profesionales idóneos para ofrecer un servicio de calidad en el acompañamiento a la ejecución y asesoría a las iniciativas innovadoras de base tecnológica que se desarrollan al interior de la Red TecnoParque.",
    ]

    for i, item in enumerate(compromisos_red, start=1):
        historia.append(crear_parrafo(f"{i}. {item}", estilo_normal))

    historia.append(
        crear_parrafo(
            "<b>PARÁGRAFO.</b> La Red TecnoParque SENA, NO financia ninguna clase de materiales, "
            "insumos, equipos, membresías, pagos, viajes, papelería, para el desarrollo de proyectos, "
            "construcción o comercialización de prototipos.",
            estilo_normal,
        )
    )

    historia.append(
        crear_parrafo(
            "<b>DECIMOPRIMERO. COMPROMISOS DE LOS TITULARES Y/O SUS INTERLOCUTORES.</b> "
            "Por medio de la presente acta se comprometen a:",
            estilo_normal,
        )
    )

    compromisos_titulares = [
        "Elaborar los documentos de Inicio, planeación, ejecución y cierre avalados por la Red TecnoParque SENA, luego de la firma de la presente Acta.",
        "Entregar a tiempo todos los documentos y las evidencias solicitadas por los expertos del Nodo, utilizando las herramientas de gestión dispuestas para tal fin.",
        "Cumplir con un horario de asistencia mínimo de horas semanales de trabajo autónomo presencial en el Nodo, el cual es establecido en común acuerdo con el experto asignado. Adicionalmente deberá cumplir con dos (2) horas semanales de acompañamiento técnico por el/la experto/a para el desarrollo del proyecto.",
        "Asistir al comité de seguimiento del proyecto y presentar al experto asignado los avances en un informe, en donde se dará cumplimiento a los objetivos planteados al inicio del proceso y con las respectivas evidencias (fotos, videos, simulaciones, diseños, entre otras) que lo respalden, lo anterior como mecanismo de autoevaluación y seguimiento, por lo tanto, es de carácter obligatorio.",
        "Asistir a las reuniones programadas por el Nodo.",
        "Del comportamiento: a) Mantener en todos los momentos (eventos, talleres, seminarios, trabajo en laboratorios, etc.) y espacios institucionales del SENA, un trato de respeto y buena convivencia. b) Utilizar la indumentaria y los elementos de protección personal dispuestos y/o solicitados por el Experto a cargo del laboratorio. c) Conservar y mantener en buen estado, orden y aseo, las instalaciones físicas, equipos y herramientas de la entidad o que estén a cargo de ésta, respondiendo por los daños ocasionados a éstos intencionalmente o por descuido, debidamente comprobados.",
        "No realizar actividades diferentes a las requeridas por el proyecto dentro de instalaciones del Nodo o no avaladas por la Red, en caso de presentarse la necesidad deberá contar con la autorización del Experto asignado al proyecto.",
        "Una vez finalizado el proyecto, se firmará un Acta de Cierre, en donde la(s)/el(los) Titular(es) entregará(n) las evidencias de finalización como fotos, videos, simulaciones, diseños e informes correspondientes.",
        "La(s)/el(los) Titular(es) en contraprestación a los servicios recibidos por la Red, realizará promoción y difusión de la Imagen Red TecnoParque SENA, esto durante el desarrollo del proyecto y una vez finalizado. Para ello utilizará el Logo SENA/Tecnoparque, el cual estará acompañado de la siguiente frase: “Apoyado por la Red Tecnoparque”, impreso y pegado sobre el prototipo del producto/servicio. Nunca en productos comerciales ni en prototipos en proceso de patente.",
        "Una vez finalizado el proyecto, asistir a la rueda de iniciativas empresariales, evento programado por el Nodo para la muestra, proyección y difusión de las iniciativas gestadas con el apoyo de la institución, para ello se deberán tener en cuenta las pautas para la selección de las iniciativas empresariales a presentar en el evento, estas pautas son diseñadas acorde a las particularidades de la región y el Nodo en el que se desarrollaron los proyectos.",
        "Conocer, aceptar y dar cumplimiento a los términos para uso de infraestructura adecuado de los diferentes laboratorios y equipos de la Red Tecnoparque SENA, incluyendo las medidas de Bioseguridad pertinentes en cada nodo y laboratorio.",
        "Programar, coordinar y asegurar la asistencia y las actividades de trabajo del equipo de Talentos Ejecutores del Proyecto.",
        "Cumplir con todos los protocolos de bioseguridad de los diferentes espacios a utilizar en cada Nodo.",
    ]

    for i, item in enumerate(compromisos_titulares, start=1):
        historia.append(crear_parrafo(f"{i}. {item}", estilo_normal))

    historia.append(
        crear_parrafo(
            "<b>DECIMOSEGUNDO. COMPROMISOS DE LOS TALENTOS EJECUTORES.</b> Por medio de la presente acta se comprometen a:",
            estilo_normal,
        )
    )

    compromisos_ejecutores = [
        "Entregar a tiempo todos los documentos y las evidencias solicitadas por los expertos del Nodo, utilizando las herramientas de gestión dispuestas para tal fin.",
        "Cumplir con el horario de asistencia acordado con el Talento interlocutor.",
        "Asistir al comité de seguimiento del proyecto y presentar al experto asignado los avances en un informe, en donde se dará cumplimiento a los objetivos planteados al inicio del proceso y con las respectivas evidencias (fotos, videos, simulaciones, diseños, entre otras) que lo respalden, lo anterior como mecanismo de autoevaluación y seguimiento, por lo tanto, es de carácter obligatorio.",
        "Asistir a las reuniones programadas por el Nodo.",
        "Del comportamiento: a) Mantener en todos los momentos (eventos, talleres, seminarios, trabajo en laboratorios, etc.) y espacios institucionales del SENA, un trato de respeto y buena convivencia. b) Utilizar la indumentaria y los elementos de protección personal dispuestos y/o solicitados por el Experto a cargo del laboratorio. c) Conservar y mantener en buen estado, orden y aseo, las instalaciones físicas, equipos y herramientas de la entidad o que estén a cargo de ésta, respondiendo por los daños ocasionados a éstos intencionalmente o por descuido, debidamente comprobados.",
        "No realizar actividades diferentes a las requeridas por el proyecto dentro de instalaciones del Nodo o no avaladas por la Red, en caso de presentarse la necesidad deberá contar con la autorización del Experto asignado al proyecto.",
        "Conocer, aceptar y dar cumplimiento a los términos para uso de infraestructura adecuado de los diferentes laboratorios y equipos de la Red Tecnoparque SENA, incluyendo las medidas de Bioseguridad pertinentes en cada nodo y laboratorio.",
        "Cumplir con todos los protocolos de bioseguridad de los diferentes espacios a utilizar en cada Nodo.",
    ]

    for i, item in enumerate(compromisos_ejecutores, start=1):
        historia.append(crear_parrafo(f"{i}. {item}", estilo_normal))

    historia.append(
        crear_parrafo(
            "<b>DECIMOTERCERO. TRANSFERENCIA DE CONOCIMIENTO.</b> En contrapartida por haber "
            "recibido el servicio de Asesoría técnica especializada y usos de infraestructura, en el desarrollo "
            "de proyectos de Base Tecnológica, la(s)/el(los) Titular(es) cumplirá(n) con alguna(s) de las siguientes "
            "actividades de Transferencia de Conocimiento. Éstas se ejecutan dentro del tiempo en el que "
            "la(s)/el(los) Titular(es) está(n) recibiendo el servicio mencionado y se definen y cumplen de acuerdo "
            "con los cronogramas de trabajo que se construyen entre Experto y Titulares en la etapa de planeación "
            "del proyecto:",
            estilo_normal,
        )
    )

    transferencias = [
        "Desarrollar Charlas Informativas de casos de éxito.",
        "Participar como ponente en un evento de divulgación tecnológica hacia empresa o academia.",
        "Participar en eventos en representación del SENA.",
        "Apoyar procesos de actualización a Experto Tecnoparque a través de transferencias de conocimiento.",
    ]

    for i, item in enumerate(transferencias, start=1):
        historia.append(crear_parrafo(f"{i}. {item}", estilo_normal))

    historia.append(
        crear_parrafo(
            "La(s)/el(los) Titular(es) además debe(n) diligenciar en su totalidad los documentos entregables "
            "y entregar uno de los productos, de acuerdo con la fase en la que se encuentre el proyecto y con "
            "los formatos establecidos por la Red, para ello deberá entregar una cuenta de correo personal o "
            "empresarial y compartir los documentos solicitados por el Experto a cargo de las asesorías del proyecto. "
            "Los documentos en los que debe participar son:",
            estilo_normal,
        )
    )

    documentos_soporte = [
        "a. Acta de inicio",
        "b. Actas de ejecución",
        "c. Encuesta de satisfacción",
        "d. Acta de cierre",
        "e. Documentos soporte",
    ]

    for item in documentos_soporte:
        historia.append(crear_parrafo(item, estilo_normal))

    historia.append(
        crear_parrafo(
            "<b>DECIMOCUARTO. INCUMPLIMIENTO DE COMPROMISOS DEL TITULAR.</b> El incumplimiento "
            "de los compromisos adquiridos por la(s)/el(los) Titular(es) dará lugar a la aplicación de las medidas "
            "restrictivas que se contemplan a continuación, dependiendo de la naturaleza del incumplimiento.",
            estilo_normal,
        )
    )

    historia.append(
        crear_parrafo(
            "<b>DECIMOQUINTO MEDIDAS RESTRICTIVAS.</b> Dependiendo de la naturaleza del incumplimiento "
            "de los compromisos adquiridos por la(s) persona(s) Titular(es), se aplicarán las siguientes medidas restrictivas:",
            estilo_normal,
        )
    )

    medidas = [
        "<b>Restricción de acceso:</b> restricción de acceso a herramientas, laboratorios, equipos especializados y asesorías y pérdida de privilegios de horarios, durante un (1) mes, cuando la(s) persona(s) Titular(es), incumplan reiteradamente las citas y horarios programados con los expertos, que estén incumpliendo con el plan de trabajo injustificadamente y/o cuando no presenten los informes con evidencias de avance.",
        "<b>Restricción temporal de eventos:</b> restricción de acceso a cierto tipo de eventos durante un periodo de tres (3) meses. Aplica para personas que se hayan inscrito en talleres, charlas y actividades complementarias y no hayan asistido quitándole el cupo o la oportunidad a otras personas de participar.",
        "<b>Suspensión temporal:</b> suspensión de todos los servicios ofrecidos por la Red TecnoParque SENA, durante un periodo igual 30 días hábiles, cuando la(s) persona(s) Titular(es) no asista(n) al comité de seguimiento o se ausente(n) por más de cuatro (4) semanas al Nodo sin previa notificación o justificación.",
        "<b>Cancelación del proyecto:</b> se presenta cuando la(s) persona(s) Titular(es) se ausente(n) de las actividades de la Red Tecnoparque por un periodo superior a un (1) mes sin previa notificación o justificación. Durante los seis (6) meses siguientes a la cancelación del proyecto, no se podrá prestar proyectos al Comité de Ideas.",
    ]

    for i, item in enumerate(medidas, start=1):
        historia.append(crear_parrafo(f"{i}. {item}", estilo_normal))

    historia.append(
        crear_parrafo(
            "<b>DECIMOSEXTO. MODIFICACIÓN O TERMINACIÓN.</b> Este acuerdo sólo podrá ser modificado "
            "o darse por terminado con el consentimiento expreso por escrito de ambas partes antes o en el Acta de Cierre del proyecto.",
            estilo_normal,
        )
    )

    historia.append(
        crear_parrafo(
            "<b>DECIMOSÉPTIMO. FIRMA DEL DOCUMENTO.</b> Para la firma de este documento en los casos "
            "de los menores de edad, este deberá ser avalado y firmado por un acudiente mayor de edad quien "
            "también firmará el presente acuerdo, aceptando todas las políticas y manuales vigentes de la Red "
            "TecnoParque SENA.",
            estilo_normal,
        )
    )

    historia.append(
        crear_parrafo(
            f"Para constancia, se firma el documento en la ciudad de {ciudad} a los "
            f"{fecha_larga_espanol_con_del(fecha_documento)}, por las partes:",
            estilo_normal,
        )
    )

    historia.append(Spacer(1, 10))
    historia.append(crear_parrafo("Firmas", estilo_subtitulo))

    firmas = [
        {
            "firma": ruta_firma_talento,
            "nombre": nombre_talento,
            "cargo": "Nombre Talento Interlocutor",
            "cedula": f"C.C. {cedula_talento}",
        },
        {
            "firma": ruta_firma("fsergio.png"),
            "nombre": "Sergio Andrés Cabrera",
            "cargo": "Nombre del Experto Tecnoparque",
            "cedula": "C.C. 1.110.454.504",
        },
        {
            "firma": ruta_firma("fcaro.png"),
            "nombre": "Carolina Garcia Monje",
            "cargo": "Nombre del Experto Tecnoparque",
            "cedula": "C.C. 36.301.495",
        },
        {
            "firma": ruta_firma("fdiego.png"),
            "nombre": "Diego Alfonso Polania",
            "cargo": "Nombre del Experto Tecnoparque",
            "cedula": "C.C. 7.684.683",
        },
        {
            "firma": ruta_firma("fcesar.png"),
            "nombre": "Cesar Augusto Pérez Tafur",
            "cargo": "Nombre del Experto Tecnoparque",
            "cedula": "C.C. 7.728.013",
        },
        {
            "firma": ruta_firma("fmaria.png"),
            "nombre": "Maria Andrea Qimbaya",
            "cargo": "Nombre del Apoyo Administrativo",
            "cedula": "C.C. 1003.812.026",
        },
        {
            "firma": ruta_firma("ffelix.png"),
            "nombre": "Felix Augusto Reyes Gutierrez",
            "cargo": "Profesional Grado 10.",
            "cedula": "C.C. 93407279",
        },
    ]

    tabla_firmas_data = [
        [
            crear_parrafo("<b>Nombre, cargo y cédula</b>", estilo_tabla_negrita),
            crear_parrafo("<b>Firma</b>", estilo_tabla_negrita),
        ]
    ]

    for firmante in firmas:
        tabla_firmas_data.append(
            [
                crear_parrafo(
                    f"<b>{firmante['nombre']}</b><br/>{firmante['cargo']}<br/>{firmante['cedula']}",
                    estilo_tabla,
                ),
                firma_img(firmante["firma"], width=4.0 * cm, height=1.05 * cm),
            ]
        )

    tabla_firmas = Table(
        tabla_firmas_data,
        colWidths=[10.2 * cm, 5.0 * cm],
        rowHeights=[0.75 * cm] + [1.65 * cm for _ in firmas],
    )

    tabla_firmas.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
                ("BACKGROUND", (0, 0), (-1, 0), colors.whitesmoke),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("ALIGN", (0, 0), (0, -1), "LEFT"),
                ("ALIGN", (1, 0), (1, -1), "CENTER"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )

    historia.append(tabla_firmas)

    doc.build(
        historia,
        onFirstPage=encabezado_pie,
        onLaterPages=encabezado_pie,
    )

    datos_json = dict(datos)
    datos_json["fecha_documento"] = fecha_documento.strftime("%d/%m/%Y")
    datos_json["ruta_pdf"] = ruta_pdf
    guardar_datos_json(datos_json, ruta="datos_confidencialidad_compromiso.json")

    return ruta_pdf

# =====================================================
# CRONOGRAMA DE ACTIVIDADES - FASE DE PLANEACIÓN
# =====================================================

RUTA_LOGO_TECNOPARQUE = "recursos/logo_tecnoparque.png"


def obtener_ruta_logo_tecnoparque() -> str | None:
    posibles_rutas = [
        Path(RUTA_LOGO_TECNOPARQUE),
        Path("recursos/logo_tecnoparque.jpg"),
        Path("recursos/logo_tecnoparque.jpeg"),
        Path(RUTA_LOGO_SENA),
    ]

    for ruta in posibles_rutas:
        if ruta.exists():
            return str(ruta)

    return None


def normalizar_dia_semana(nombre_dia: str) -> int:
    mapa = {
        "lunes": 0,
        "martes": 1,
        "miércoles": 2,
        "miercoles": 2,
        "jueves": 3,
        "viernes": 4,
        "sábado": 5,
        "sabado": 5,
        "domingo": 6,
    }
    return mapa[nombre_dia.lower().strip()]


def obtener_fechas_programadas(fecha_inicio: date, fecha_fin: date, dias_semana: list[str]) -> list[date]:
    dias_num = [normalizar_dia_semana(dia) for dia in dias_semana]
    fechas = []
    actual = fecha_inicio

    while actual <= fecha_fin:
        if actual.weekday() in dias_num:
            fechas.append(actual)
        actual += timedelta(days=1)

    return fechas


def dividir_fechas_por_actividad(fechas: list[date], cantidad_actividades: int) -> list[list[date]]:
    if cantidad_actividades <= 0:
        return []

    if not fechas:
        return [[] for _ in range(cantidad_actividades)]

    bloques = []
    total_fechas = len(fechas)

    for i in range(cantidad_actividades):
        inicio = round(i * total_fechas / cantidad_actividades)
        fin = round((i + 1) * total_fechas / cantidad_actividades)
        bloque = fechas[inicio:fin]

        if not bloque and fechas:
            bloque = [fechas[min(i, total_fechas - 1)]]

        bloques.append(bloque)

    return bloques


def generar_actividades_cronograma_modo_prueba(descripcion_proyecto: str, cantidad_actividades: int) -> list[str]:
    actividades_base = [
        "Revisión técnica y conceptual del proyecto",
        "Identificación de requerimientos técnicos, funcionales y operativos",
        "Definición de alternativas de diseño y criterios de selección",
        "Diseño preliminar de la solución propuesta",
        "Modelado, simulación o representación técnica de la solución",
        "Validación técnica de componentes, materiales o procesos",
        "Ajustes del diseño de acuerdo con la validación realizada",
        "Construcción o integración del prototipo funcional",
        "Pruebas de funcionamiento y verificación técnica",
        "Documentación de resultados y consolidación de entregables",
    ]

    if cantidad_actividades <= len(actividades_base):
        return actividades_base[:cantidad_actividades]

    actividades = actividades_base[:]
    while len(actividades) < cantidad_actividades:
        actividades.append(f"Actividad técnica complementaria {len(actividades) + 1}")

    return actividades


def generar_actividades_cronograma_con_chatgpt(
    descripcion_proyecto: str,
    cantidad_actividades: int,
    modelo: str = "gpt-4.1-mini"
) -> list[str]:
    if OpenAI is None:
        raise ImportError("No está instalada la librería openai. Instálala con: pip install openai")

    api_key = obtener_api_key()
    if not api_key:
        raise ValueError("No se encontró OPENAI_API_KEY. Configúrala como variable de entorno o en .streamlit/secrets.toml.")

    client = OpenAI(api_key=api_key)

    instrucciones = """
Eres un experto en formulación, planeación y seguimiento de proyectos de base tecnológica
en el marco de la Red Tecnoparque SENA.

Debes generar actividades técnicas para un cronograma de planeación tipo diagrama de Gantt.
Las actividades deben ser claras, verificables, técnicas y orientadas al desarrollo del proyecto.
No inventes nombres de personas, fechas, códigos ni entidades.
No incluyas numeración.
No incluyas fechas.
No incluyas explicaciones adicionales.
Responde únicamente en JSON válido.
"""

    entrada = f"""
Genera exactamente {cantidad_actividades} actividades técnicas para el cronograma del siguiente proyecto.

Descripción del proyecto:
{descripcion_proyecto}

Formato obligatorio:
{{
  "actividades": [
    "Actividad 1",
    "Actividad 2"
  ]
}}
"""

    respuesta = client.responses.create(
        model=modelo,
        instructions=instrucciones,
        input=entrada,
        temperature=0.35
    )

    texto = limpiar_respuesta_json(respuesta.output_text)

    try:
        datos = json.loads(texto)
        actividades = datos.get("actividades", [])
    except json.JSONDecodeError:
        actividades = []

    actividades_limpias = []
    for actividad in actividades:
        actividad = str(actividad).strip()
        if actividad:
            actividades_limpias.append(actividad)

    if len(actividades_limpias) < cantidad_actividades:
        actividades_extra = generar_actividades_cronograma_modo_prueba(
            descripcion_proyecto,
            cantidad_actividades - len(actividades_limpias)
        )
        actividades_limpias.extend(actividades_extra)

    return actividades_limpias[:cantidad_actividades]


def nombre_mes_es(fecha: date) -> str:
    meses = {
        1: "Enero",
        2: "Febrero",
        3: "Marzo",
        4: "Abril",
        5: "Mayo",
        6: "Junio",
        7: "Julio",
        8: "Agosto",
        9: "Septiembre",
        10: "Octubre",
        11: "Noviembre",
        12: "Diciembre",
    }
    return meses[fecha.month]


def agrupar_fechas_por_bloques_de_meses(fechas: list[date], max_meses_por_hoja: int = 3) -> list[list[date]]:
    """
    Agrupa las fechas programadas en bloques de máximo 3 meses calendario.
    Cada bloque se renderiza como una hoja diferente del cronograma.
    """
    if not fechas:
        return []

    fechas_ordenadas = sorted(fechas)
    bloques = []
    bloque_actual = []
    meses_actuales = []

    for fecha in fechas_ordenadas:
        clave_mes = (fecha.year, fecha.month)

        if clave_mes not in meses_actuales:
            if len(meses_actuales) >= max_meses_por_hoja:
                bloques.append(bloque_actual)
                bloque_actual = []
                meses_actuales = []
            meses_actuales.append(clave_mes)

        bloque_actual.append(fecha)

    if bloque_actual:
        bloques.append(bloque_actual)

    return bloques


def generar_pdf_cronograma(datos: dict) -> str:
    if canvas is None:
        raise ImportError("No está instalada reportlab. Instálala con: pip install reportlab")

    Path(CARPETA_SALIDA).mkdir(parents=True, exist_ok=True)

    nombre_archivo = (
        f"Cronograma_Actividades_"
        f"{safe_filename(datos.get('codigo_proyecto', 'proyecto'))}_"
        f"{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    )
    ruta_pdf = str(Path(CARPETA_SALIDA) / nombre_archivo)

    page_size = landscape(letter)
    page_width, page_height = page_size

    # Márgenes laterales ampliados para que el cronograma no quede pegado al borde.
    margen_lateral = 1.35 * cm
    ancho_disponible = page_width - (margen_lateral * 2)

    doc = SimpleDocTemplate(
        ruta_pdf,
        pagesize=page_size,
        rightMargin=margen_lateral,
        leftMargin=margen_lateral,
        topMargin=2.6 * cm,
        bottomMargin=1.25 * cm,
    )

    def encabezado_pie(c, doc):
        c.saveState()

        ruta_logo = obtener_ruta_logo_tecnoparque()
        if ruta_logo:
            try:
                logo = ImageReader(ruta_logo)
                c.drawImage(
                    logo,
                    margen_lateral,
                    page_height - 2.2 * cm,
                    width=6.2 * cm,
                    height=1.5 * cm,
                    preserveAspectRatio=True,
                    mask="auto",
                )
            except Exception:
                c.setFont("Helvetica-Bold", 14)
                c.setFillColor(colors.HexColor("#39a935"))
                c.drawString(margen_lateral, page_height - 1.4 * cm, "SENA Tecnoparque")
                c.setFillColor(colors.black)
        else:
            c.setFont("Helvetica-Bold", 14)
            c.setFillColor(colors.HexColor("#39a935"))
            c.drawString(margen_lateral, page_height - 1.4 * cm, "SENA Tecnoparque")
            c.setFillColor(colors.black)

        c.setFont("Helvetica-Bold", 13)
        c.setFillColor(colors.black)
        c.drawCentredString(page_width / 2, page_height - 1.3 * cm, "CRONOGRAMA DE ACTIVIDADES")

        c.setFont("Helvetica", 8)
        c.setFillColor(colors.grey)
        c.drawCentredString(page_width / 2, 0.55 * cm, "TP-PEPBT V.1")

        c.restoreState()

    estilo_normal = ParagraphStyle(
        name="NormalCronograma",
        fontName="Helvetica",
        fontSize=7.0,
        leading=8.2,
        alignment=TA_LEFT,
    )

    estilo_negrita = ParagraphStyle(
        name="NegritaCronograma",
        fontName="Helvetica-Bold",
        fontSize=7.0,
        leading=8.2,
        alignment=TA_LEFT,
    )

    estilo_centro = ParagraphStyle(
        name="CentroCronograma",
        fontName="Helvetica-Bold",
        fontSize=5.0,
        leading=5.8,
        alignment=TA_CENTER,
    )

    estilo_mes = ParagraphStyle(
        name="MesCronograma",
        fontName="Helvetica-Bold",
        fontSize=5.0,
        leading=5.6,
        alignment=TA_CENTER,
    )

    estilo_dia = ParagraphStyle(
        name="DiaCronograma",
        fontName="Helvetica",
        fontSize=4.8,
        leading=5.2,
        alignment=TA_CENTER,
    )

    historia = []

    fechas_todas = datos["fechas_programadas"]
    actividades = datos["actividades"]
    bloques_fechas_actividad = datos["bloques_fechas"]

    bloques_hojas = agrupar_fechas_por_bloques_de_meses(fechas_todas, max_meses_por_hoja=3)
    if not bloques_hojas:
        bloques_hojas = [[]]

    for numero_hoja, fechas_hoja in enumerate(bloques_hojas, start=1):
        if numero_hoja > 1:
            historia.append(PageBreak())

        if fechas_hoja:
            periodo_hoja = f"{fechas_hoja[0].strftime('%d/%m/%Y')} al {fechas_hoja[-1].strftime('%d/%m/%Y')}"
        else:
            periodo_hoja = f"{datos['fecha_inicio'].strftime('%d/%m/%Y')} al {datos['fecha_fin'].strftime('%d/%m/%Y')}"

        # =====================================================
        # DATOS GENERALES DEL CRONOGRAMA
        # =====================================================
        datos_generales = [
            [
                Paragraph("<b>NOMBRE DEL PROYECTO</b>", estilo_negrita),
                Paragraph(datos["nombre_proyecto"], estilo_normal),
                Paragraph("<b>NOMBRE DEL TALENTO</b>", estilo_negrita),
                Paragraph(datos["nombre_talento"], estilo_normal),
            ],
            [
                Paragraph("<b>CÓDIGO DEL PROYECTO</b>", estilo_negrita),
                Paragraph(datos["codigo_proyecto"], estilo_normal),
                Paragraph("<b>EXPERTO</b>", estilo_negrita),
                Paragraph(datos["nombre_experto"], estilo_normal),
            ],
            [
                Paragraph("<b>LÍNEA</b>", estilo_negrita),
                Paragraph(datos["linea"], estilo_normal),
                Paragraph("<b>TIEMPO DE EJECUCIÓN</b>", estilo_negrita),
                Paragraph(
                    f"{datos['fecha_inicio'].strftime('%d/%m/%Y')} al {datos['fecha_fin'].strftime('%d/%m/%Y')}",
                    estilo_normal,
                ),
            ],
            [
                Paragraph("<b>DÍAS PROGRAMADOS</b>", estilo_negrita),
                Paragraph(", ".join(datos["dias_semana"]), estilo_normal),
                Paragraph("<b>PERIODO MOSTRADO</b>", estilo_negrita),
                Paragraph(periodo_hoja, estilo_normal),
            ],
        ]

        ancho_col_1 = 3.8 * cm
        ancho_col_3 = 3.8 * cm
        ancho_col_2 = ancho_disponible * 0.38
        ancho_col_4 = ancho_disponible - ancho_col_1 - ancho_col_2 - ancho_col_3

        tabla_datos = Table(
            datos_generales,
            colWidths=[ancho_col_1, ancho_col_2, ancho_col_3, ancho_col_4],
            rowHeights=[0.65 * cm, 0.55 * cm, 0.55 * cm, 0.55 * cm],
        )

        tabla_datos.setStyle(
            TableStyle(
                [
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.black),
                    ("BACKGROUND", (0, 0), (0, -1), colors.whitesmoke),
                    ("BACKGROUND", (2, 0), (2, -1), colors.whitesmoke),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 4),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                    ("TOPPADDING", (0, 0), (-1, -1), 2),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
                ]
            )
        )

        historia.append(tabla_datos)
        historia.append(Spacer(1, 0.25 * cm))

        # =====================================================
        # TABLA GANTT POR HOJA - MÁXIMO 3 MESES
        # =====================================================
        ancho_item = 1.0 * cm
        ancho_desc = 8.0 * cm
        ancho_fechas_total = ancho_disponible - ancho_item - ancho_desc

        if fechas_hoja:
            ancho_fecha = max(0.42 * cm, ancho_fechas_total / len(fechas_hoja))
        else:
            ancho_fecha = 0.6 * cm

        col_widths = [ancho_item, ancho_desc] + [ancho_fecha for _ in fechas_hoja]

        fila_meses = [
            Paragraph("<b>ITEM</b>", estilo_centro),
            Paragraph("<b>DESCRIPCIÓN ACTIVIDAD</b>", estilo_centro),
        ]

        for fecha in fechas_hoja:
            fila_meses.append(Paragraph(nombre_mes_es(fecha), estilo_mes))

        fila_dias = [
            Paragraph("", estilo_centro),
            Paragraph("", estilo_centro),
        ]

        for fecha in fechas_hoja:
            fila_dias.append(Paragraph(str(fecha.day), estilo_dia))

        tabla_data = [fila_meses, fila_dias]

        for idx, actividad in enumerate(actividades, start=1):
            fila = [
                Paragraph(str(idx), estilo_centro),
                Paragraph(actividad, estilo_normal),
            ]

            for _fecha in fechas_hoja:
                fila.append(Paragraph("", estilo_centro))

            tabla_data.append(fila)

        row_heights = [0.42 * cm, 0.38 * cm] + [0.75 * cm for _ in actividades]

        tabla_gantt = Table(
            tabla_data,
            colWidths=col_widths,
            rowHeights=row_heights,
            repeatRows=2,
        )

        estilos_tabla = [
            ("GRID", (0, 0), (-1, -1), 0.35, colors.black),
            ("BACKGROUND", (0, 0), (-1, 1), colors.whitesmoke),
            ("BACKGROUND", (0, 0), (1, -1), colors.whitesmoke),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("ALIGN", (0, 0), (0, -1), "CENTER"),
            ("ALIGN", (2, 0), (-1, -1), "CENTER"),
            ("LEFTPADDING", (0, 0), (-1, -1), 1),
            ("RIGHTPADDING", (0, 0), (-1, -1), 1),
            ("TOPPADDING", (0, 0), (-1, -1), 1),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
        ]

        # Unificar celdas de meses consecutivos dentro de la hoja actual.
        if fechas_hoja:
            inicio_mes = 2
            mes_actual = fechas_hoja[0].month
            anio_actual = fechas_hoja[0].year

            for idx_fecha, fecha in enumerate(fechas_hoja, start=2):
                es_ultima = idx_fecha == len(fechas_hoja) + 1
                cambia_mes = fecha.month != mes_actual or fecha.year != anio_actual

                if cambia_mes:
                    fin_mes = idx_fecha - 1
                    if fin_mes > inicio_mes:
                        estilos_tabla.append(("SPAN", (inicio_mes, 0), (fin_mes, 0)))

                    inicio_mes = idx_fecha
                    mes_actual = fecha.month
                    anio_actual = fecha.year

                if es_ultima:
                    fin_mes = idx_fecha
                    if fin_mes > inicio_mes:
                        estilos_tabla.append(("SPAN", (inicio_mes, 0), (fin_mes, 0)))

        # Marcar fechas de cada actividad en verde, únicamente cuando estén dentro de la hoja actual.
        verde = colors.HexColor("#39a935")

        for idx_actividad, fechas_actividad in enumerate(bloques_fechas_actividad, start=2):
            fila_tabla = idx_actividad
            for idx_fecha, fecha in enumerate(fechas_hoja, start=2):
                if fecha in fechas_actividad:
                    estilos_tabla.append(
                        ("BACKGROUND", (idx_fecha, fila_tabla), (idx_fecha, fila_tabla), verde)
                    )

        tabla_gantt.setStyle(TableStyle(estilos_tabla))
        historia.append(tabla_gantt)

    # No se agrega tabla inferior de entregables.
    # El PDF termina con el diagrama de Gantt, paginado por bloques de máximo 3 meses.

    doc.build(
        historia,
        onFirstPage=encabezado_pie,
        onLaterPages=encabezado_pie,
    )

    datos_json = dict(datos)
    datos_json["fecha_inicio"] = datos["fecha_inicio"].strftime("%d/%m/%Y")
    datos_json["fecha_fin"] = datos["fecha_fin"].strftime("%d/%m/%Y")
    datos_json["fechas_programadas"] = [
        f.strftime("%d/%m/%Y") for f in datos["fechas_programadas"]
    ]
    datos_json["bloques_fechas"] = [
        [f.strftime("%d/%m/%Y") for f in bloque]
        for bloque in datos["bloques_fechas"]
    ]
    datos_json["ruta_pdf"] = ruta_pdf

    guardar_datos_json(datos_json, ruta="datos_cronograma_actividades.json")

    return ruta_pdf
# =====================================================
# ESTADO DEL ARTE - FASE DE PLANEACIÓN
# =====================================================

def limpiar_lista_tecnologias(texto: str) -> list[str]:
    if not texto:
        return []

    separadores = ["\n", ";", ","]
    tecnologias = [texto]

    for sep in separadores:
        nuevas = []
        for item in tecnologias:
            nuevas.extend(item.split(sep))
        tecnologias = nuevas

    return [t.strip() for t in tecnologias if t.strip()]


def generar_estado_arte_modo_prueba(
    nombre_proyecto: str,
    codigo_proyecto: str,
    descripcion_proyecto: str,
    tecnologias_previstas: list[str],
) -> dict:
    tecnologias_base = tecnologias_previstas or ["Inteligencia artificial", "Aplicaciones web", "Bases de datos", "Visualización de información"]

    return {
        "introduccion": (
            f"El presente Estado del Arte tiene como propósito establecer una base conceptual, técnica e investigativa "
            f"para el proyecto {nombre_proyecto}, identificado con el código {codigo_proyecto}. El documento busca "
            "revisar antecedentes, referentes nacionales e internacionales, tecnologías relevantes y tendencias "
            "emergentes que permitan orientar la toma de decisiones durante la fase de planeación."
        ),
        "objetivos": [
            "Identificar referentes técnicos, académicos y aplicados relacionados con el proyecto.",
            "Analizar el contexto nacional e internacional del sector en el cual se desarrolla la iniciativa.",
            "Reconocer tecnologías relevantes y emergentes que puedan fortalecer la solución propuesta.",
            "Construir una base documental que oriente las decisiones técnicas de la fase de planeación."
        ],
        "antecedentes_contexto": (
            f"{descripcion_proyecto}\n\n"
            "A partir de la información suministrada, el proyecto se configura como una iniciativa de base tecnológica "
            "orientada a resolver una necesidad específica mediante el uso de herramientas digitales, procesos de innovación "
            "y desarrollo técnico. Su origen se relaciona con la identificación de una oportunidad de mejora en un contexto "
            "productivo, educativo, ambiental, cultural o social, donde la incorporación de tecnología puede generar valor "
            "agregado. El talento vinculado al proyecto cumple un papel central, dado que aporta conocimiento del problema, "
            "experiencia del entorno y motivación para transformar la necesidad identificada en una solución viable. "
            "Desde la perspectiva de innovación, el proyecto resulta pertinente porque combina componentes técnicos, "
            "metodológicos y funcionales que pueden diferenciarlo de soluciones convencionales. La revisión de antecedentes "
            "permite reconocer que iniciativas similares suelen requerir una adecuada articulación entre análisis del problema, "
            "selección tecnológica, diseño de prototipos, validación funcional y documentación de resultados. En este sentido, "
            "el Estado del Arte permite delimitar los enfoques existentes, identificar oportunidades de mejora y proponer "
            "criterios técnicos que orienten el desarrollo del proyecto en sus siguientes fases."
        ),
        "contexto_nacional_internacional": (
            "En el contexto nacional, el desarrollo de proyectos de base tecnológica se ha fortalecido gracias a la adopción "
            "de herramientas digitales, procesos de automatización, analítica de datos, prototipado rápido e integración de "
            "tecnologías emergentes en sectores productivos y sociales. En Colombia, las iniciativas orientadas a innovación "
            "aplicada suelen articular actores académicos, institucionales y empresariales, promoviendo soluciones ajustadas "
            "a las necesidades territoriales. En el ámbito internacional, se evidencia una tendencia creciente hacia el uso "
            "de inteligencia artificial, sistemas inteligentes, plataformas web, aplicaciones móviles, Internet de las Cosas, "
            "modelado digital y tecnologías de visualización para mejorar procesos, optimizar recursos y generar experiencias "
            "más eficientes. Estas tendencias permiten ubicar el proyecto dentro de una dinámica global de transformación "
            "digital e innovación aplicada."
        ),
        "proyectos_similares": [
            {
                "numero": 1,
                "nombre": "Proyecto de innovación tecnológica aplicado al sector relacionado",
                "enlace": "https://www.sena.edu.co",
                "referencia_apa": "Servicio Nacional de Aprendizaje. (2025). Proyectos de innovación tecnológica. SENA."
            },
            {
                "numero": 2,
                "nombre": "Experiencia internacional de transformación digital aplicada",
                "enlace": "https://www.oecd.org",
                "referencia_apa": "OECD. (2024). Digital transformation and innovation practices. OECD Publishing."
            },
            {
                "numero": 3,
                "nombre": "Solución tecnológica basada en sistemas inteligentes",
                "enlace": "https://ieeexplore.ieee.org",
                "referencia_apa": "IEEE. (2024). Intelligent systems and applied technology cases. IEEE Xplore."
            },
            {
                "numero": 4,
                "nombre": "Aplicación de tecnologías emergentes en proyectos productivos",
                "enlace": "https://www.sciencedirect.com",
                "referencia_apa": "ScienceDirect. (2024). Emerging technologies in applied innovation. Elsevier."
            },
            {
                "numero": 5,
                "nombre": "Caso de uso de prototipado tecnológico para innovación",
                "enlace": "https://scholar.google.com",
                "referencia_apa": "Google Scholar. (2025). Applied prototyping and innovation projects."
            },
        ],
        "tecnologias_relevantes": [
            {
                "tecnologia": tecnologia,
                "analisis": (
                    f"{tecnologia} se considera una tecnología relevante para el proyecto porque puede aportar capacidades "
                    "técnicas orientadas a mejorar el diseño, desarrollo, validación o implementación de la solución. Su uso "
                    "debe evaluarse de acuerdo con los recursos disponibles, el nivel de complejidad técnica y los resultados esperados."
                ),
                "cita_apa": "Autor, A. (2024). Aplicaciones tecnológicas emergentes en proyectos de innovación. Revista de Innovación Aplicada, 12(2), 45-60."
            }
            for tecnologia in tecnologias_base
        ],
        "tecnologias_emergentes": [
            {
                "tecnologia": "Inteligencia artificial generativa",
                "analisis": "La inteligencia artificial generativa permite apoyar procesos de creación, análisis, síntesis documental y generación de contenidos técnicos, facilitando la producción de información estructurada para la toma de decisiones.",
                "cita_apa": "OpenAI. (2025). Generative artificial intelligence applications in professional workflows. OpenAI."
            },
            {
                "tecnologia": "Internet de las Cosas",
                "analisis": "El Internet de las Cosas facilita la conexión de sensores, dispositivos y plataformas para capturar datos en tiempo real, lo cual puede fortalecer procesos de monitoreo, automatización y validación técnica.",
                "cita_apa": "Atzori, L., Iera, A., & Morabito, G. (2024). Internet of Things: Concepts and applications. Computer Networks."
            },
            {
                "tecnologia": "Analítica de datos",
                "analisis": "La analítica de datos permite convertir información recolectada en indicadores, patrones y criterios de decisión, aportando evidencia para validar el comportamiento del proyecto o prototipo.",
                "cita_apa": "Chen, M., Mao, S., & Liu, Y. (2024). Data analytics for intelligent systems. Information Sciences."
            },
            {
                "tecnologia": "Modelado y simulación digital",
                "analisis": "El modelado y la simulación permiten representar componentes, procesos o escenarios antes de su implementación física, reduciendo riesgos técnicos y costos de experimentación.",
                "cita_apa": "Banks, J. (2023). Simulation and modeling for engineering applications. Engineering Systems Journal."
            },
            {
                "tecnologia": "Computación en la nube",
                "analisis": "La computación en la nube aporta escalabilidad, almacenamiento y disponibilidad para plataformas digitales, permitiendo desplegar soluciones con mayor flexibilidad operativa.",
                "cita_apa": "Mell, P., & Grance, T. (2023). Cloud computing models and applications. National Institute of Standards and Technology."
            },
        ],
        "conclusiones": (
            "El Estado del Arte evidencia que el proyecto se articula con tendencias actuales de innovación tecnológica, "
            "transformación digital y desarrollo de soluciones aplicadas. La revisión de referentes y tecnologías permite "
            "identificar oportunidades para fortalecer el diseño, la validación y la implementación de la iniciativa. Asimismo, "
            "la comparación entre tecnologías previstas y emergentes ofrece criterios para tomar decisiones técnicas más sólidas "
            "durante la fase de planeación."
        ),
        "bibliografia": [
            "Atzori, L., Iera, A., & Morabito, G. (2024). Internet of Things: Concepts and applications. Computer Networks.",
            "Banks, J. (2023). Simulation and modeling for engineering applications. Engineering Systems Journal.",
            "Chen, M., Mao, S., & Liu, Y. (2024). Data analytics for intelligent systems. Information Sciences.",
            "Mell, P., & Grance, T. (2023). Cloud computing models and applications. National Institute of Standards and Technology.",
            "OpenAI. (2025). Generative artificial intelligence applications in professional workflows. OpenAI.",
        ],
    }


def generar_estado_arte_con_chatgpt(
    nombre_proyecto: str,
    codigo_proyecto: str,
    descripcion_proyecto: str,
    tecnologias_previstas: list[str],
    modelo: str = "gpt-4.1-mini",
) -> dict:
    """
    Genera el Estado del Arte en dos pasos:
    1. Intenta obtener notas con búsqueda web.
    2. Convierte esas notas en JSON usando salida estructurada.

    Si la cuenta/modelo no soporta web_search_preview o text.format, cae a métodos alternos
    y nunca deja la app bloqueada por JSON inválido.
    """
    if OpenAI is None:
        raise ImportError("No está instalada la librería openai. Instálala con: pip install openai")

    api_key = obtener_api_key()
    if not api_key:
        raise ValueError("No se encontró OPENAI_API_KEY. Configúrala como variable de entorno o en .streamlit/secrets.toml.")

    client = OpenAI(api_key=api_key)

    tecnologias_texto = ", ".join(tecnologias_previstas) if tecnologias_previstas else "No especificadas"

    # -----------------------------------------------------
    # PASO 1: Investigación con búsqueda web, en texto libre.
    # -----------------------------------------------------
    prompt_busqueda = f"""
Investiga fuentes públicas y académicas para construir un Estado del Arte en español.

Proyecto: {nombre_proyecto}
Código: {codigo_proyecto}
Descripción: {descripcion_proyecto}
Tecnologías previstas: {tecnologias_texto}

Necesito notas verificables sobre:
1. Contexto nacional e internacional del sector.
2. Mínimo 5 proyectos o iniciativas similares con nombre, entidad y enlace.
3. Tecnologías relevantes y emergentes.
4. Mínimo 5 artículos, fuentes técnicas o documentos académicos de validación.
5. Referencias en APA 7.

Entrega notas ordenadas, con enlaces visibles cuando existan.
"""

    notas_investigacion = ""
    try:
        respuesta_busqueda = client.responses.create(
            model=modelo,
            tools=[{"type": "web_search_preview"}],
            input=prompt_busqueda,
            temperature=0.2,
        )
        notas_investigacion = getattr(respuesta_busqueda, "output_text", "") or ""
    except Exception:
        notas_investigacion = (
            "No fue posible ejecutar búsqueda web desde la API. "
            "Generar el documento con base en conocimiento general y marcar fuentes para verificación manual."
        )

    # -----------------------------------------------------
    # Esquema JSON estricto para Responses API.
    # -----------------------------------------------------
    estado_arte_schema = {
        "type": "object",
        "additionalProperties": False,
        "properties": {
            "introduccion": {"type": "string"},
            "objetivos": {"type": "array", "items": {"type": "string"}},
            "antecedentes_contexto": {"type": "string"},
            "contexto_nacional_internacional": {"type": "string"},
            "proyectos_similares": {
                "type": "array",
                "items": {
                    "type": "object",
                    "additionalProperties": False,
                    "properties": {
                        "numero": {"type": "integer"},
                        "nombre": {"type": "string"},
                        "enlace": {"type": "string"},
                        "descripcion_breve": {"type": "string"},
                        "referencia_apa": {"type": "string"},
                    },
                    "required": ["numero", "nombre", "enlace", "descripcion_breve", "referencia_apa"],
                },
            },
            "tecnologias_relevantes": {
                "type": "array",
                "items": {
                    "type": "object",
                    "additionalProperties": False,
                    "properties": {
                        "tecnologia": {"type": "string"},
                        "analisis": {"type": "string"},
                        "cita_apa": {"type": "string"},
                    },
                    "required": ["tecnologia", "analisis", "cita_apa"],
                },
            },
            "tecnologias_emergentes": {
                "type": "array",
                "items": {
                    "type": "object",
                    "additionalProperties": False,
                    "properties": {
                        "tecnologia": {"type": "string"},
                        "analisis": {"type": "string"},
                        "cita_apa": {"type": "string"},
                    },
                    "required": ["tecnologia", "analisis", "cita_apa"],
                },
            },
            "articulos_validacion": {
                "type": "array",
                "items": {
                    "type": "object",
                    "additionalProperties": False,
                    "properties": {
                        "numero": {"type": "integer"},
                        "tecnologia": {"type": "string"},
                        "articulo": {"type": "string"},
                        "enlace": {"type": "string"},
                        "referencia_apa": {"type": "string"},
                    },
                    "required": ["numero", "tecnologia", "articulo", "enlace", "referencia_apa"],
                },
            },
            "conclusiones": {"type": "string"},
            "bibliografia": {"type": "array", "items": {"type": "string"}},
        },
        "required": [
            "introduccion",
            "objetivos",
            "antecedentes_contexto",
            "contexto_nacional_internacional",
            "proyectos_similares",
            "tecnologias_relevantes",
            "tecnologias_emergentes",
            "articulos_validacion",
            "conclusiones",
            "bibliografia",
        ],
    }

    instrucciones_json = """
Actúa como investigador académico senior y formulador de proyectos de base tecnológica para la Red Tecnoparque SENA.
Genera un Estado del Arte profesional, académico e investigativo en español.
Respeta la estructura solicitada y usa exclusivamente JSON válido según el esquema.
No uses markdown.
No uses bloques de código.
El campo antecedentes_contexto debe tener mínimo 500 palabras.
Incluye 4 objetivos.
Incluye mínimo 5 proyectos similares.
Incluye las tecnologías previstas por el usuario y 5 o 6 tecnologías emergentes.
Incluye mínimo 5 artículos o fuentes técnicas de validación.
Usa APA 7 en referencias.
"""

    prompt_json = f"""
Datos del proyecto:
- Nombre del proyecto: {nombre_proyecto}
- Código del proyecto: {codigo_proyecto}
- Descripción detallada: {descripcion_proyecto}
- Tecnologías previstas: {tecnologias_texto}

Notas de investigación disponibles:
{notas_investigacion}
"""

    # -----------------------------------------------------
    # PASO 2: Intento principal con Structured Outputs.
    # -----------------------------------------------------
    datos = None
    try:
        respuesta_json = client.responses.create(
            model=modelo,
            instructions=instrucciones_json,
            input=prompt_json,
            text={
                "format": {
                    "type": "json_schema",
                    "name": "estado_arte_schema",
                    "schema": estado_arte_schema,
                    "strict": True,
                }
            },
            temperature=0.15,
        )
        texto_limpio = limpiar_respuesta_json(respuesta_json.output_text)
        datos = json.loads(texto_limpio)
    except Exception:
        datos = None

    # -----------------------------------------------------
    # PASO 3: Fallback con Chat Completions en modo JSON.
    # -----------------------------------------------------
    if datos is None:
        prompt_chat = instrucciones_json + "\nDevuelve únicamente JSON válido con las claves requeridas.\n\n" + prompt_json
        try:
            respuesta_chat = client.chat.completions.create(
                model=modelo,
                messages=[
                    {"role": "system", "content": "Devuelve únicamente JSON válido. Sin markdown."},
                    {"role": "user", "content": prompt_chat},
                ],
                response_format={"type": "json_object"},
                temperature=0.15,
            )
            texto_chat = respuesta_chat.choices[0].message.content or ""
            datos = json.loads(limpiar_respuesta_json(texto_chat))
        except Exception:
            datos = None

    # -----------------------------------------------------
    # PASO 4: Último fallback: modo prueba automático para no bloquear la app.
    # -----------------------------------------------------
    if datos is None:
        datos = generar_estado_arte_modo_prueba(
            nombre_proyecto=nombre_proyecto,
            codigo_proyecto=codigo_proyecto,
            descripcion_proyecto=descripcion_proyecto,
            tecnologias_previstas=tecnologias_previstas,
        )
        datos["introduccion"] = (
            "Nota técnica: la generación con búsqueda web/JSON estructurado falló en la API, "
            "por lo cual se generó una versión base editable para no bloquear el flujo. "
            + datos.get("introduccion", "")
        )

    # -----------------------------------------------------
    # Normalización final para que el PDF nunca falle por claves faltantes.
    # -----------------------------------------------------
    campos_lista = [
        "objetivos",
        "proyectos_similares",
        "tecnologias_relevantes",
        "tecnologias_emergentes",
        "articulos_validacion",
        "bibliografia",
    ]
    campos_texto = [
        "introduccion",
        "antecedentes_contexto",
        "contexto_nacional_internacional",
        "conclusiones",
    ]

    for campo in campos_lista:
        if campo not in datos or not isinstance(datos[campo], list):
            datos[campo] = []

    for campo in campos_texto:
        if campo not in datos or not isinstance(datos[campo], str):
            datos[campo] = ""

    if len(datos["objetivos"]) < 4:
        datos["objetivos"] += [
            "Identificar referentes técnicos y académicos relacionados con el proyecto.",
            "Analizar el contexto nacional e internacional del sector de aplicación.",
            "Reconocer tecnologías relevantes y emergentes aplicables a la solución propuesta.",
            "Orientar la toma de decisiones técnicas durante la fase de planeación.",
        ][len(datos["objetivos"]):]

    if not datos["tecnologias_relevantes"]:
        datos["tecnologias_relevantes"] = [
            {
                "tecnologia": tecnologia,
                "analisis": (
                    f"{tecnologia} se considera relevante para el proyecto porque puede aportar capacidades "
                    "técnicas para el diseño, implementación, validación o escalamiento de la solución propuesta."
                ),
                "cita_apa": "Referencia técnica pendiente de verificación.",
            }
            for tecnologia in tecnologias_previstas
        ]

    return datos


def generar_pdf_estado_arte(datos: dict) -> str:
    if canvas is None:
        raise ImportError("No está instalada reportlab. Instálala con: pip install reportlab")

    Path(CARPETA_SALIDA).mkdir(parents=True, exist_ok=True)

    nombre_archivo = (
        f"Estado_del_Arte_"
        f"{safe_filename(datos.get('codigo_proyecto', 'proyecto'))}_"
        f"{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    )
    ruta_pdf = str(Path(CARPETA_SALIDA) / nombre_archivo)

    page_width, page_height = letter

    doc = SimpleDocTemplate(
        ruta_pdf,
        pagesize=letter,
        rightMargin=2.0 * cm,
        leftMargin=2.0 * cm,
        topMargin=2.7 * cm,
        bottomMargin=1.7 * cm,
    )

    def encabezado_pie(c, doc):
        c.saveState()

        ruta_logo = obtener_ruta_logo_tecnoparque()
        if ruta_logo:
            try:
                logo = ImageReader(ruta_logo)
                c.drawImage(
                    logo,
                    1.5 * cm,
                    page_height - 2.1 * cm,
                    width=5.8 * cm,
                    height=1.4 * cm,
                    preserveAspectRatio=True,
                    mask="auto",
                )
            except Exception:
                c.setFont("Helvetica-Bold", 12)
                c.setFillColor(colors.HexColor("#39a935"))
                c.drawString(1.5 * cm, page_height - 1.4 * cm, "SENA Tecnoparque")
                c.setFillColor(colors.black)
        else:
            c.setFont("Helvetica-Bold", 12)
            c.setFillColor(colors.HexColor("#39a935"))
            c.drawString(1.5 * cm, page_height - 1.4 * cm, "SENA Tecnoparque")
            c.setFillColor(colors.black)

        c.setFont("Helvetica", 8)
        c.setFillColor(colors.grey)
        c.drawCentredString(page_width / 2, 0.8 * cm, "FORMATO ESTADO DEL ARTE - RED TECNOPARQUE COLOMBIA")
        c.setFillColor(colors.black)

        c.restoreState()

    estilo_titulo = ParagraphStyle(
        name="TituloEstadoArte",
        fontName="Helvetica-Bold",
        fontSize=13,
        leading=16,
        alignment=TA_CENTER,
        spaceAfter=10,
    )

    estilo_subtitulo = ParagraphStyle(
        name="SubtituloEstadoArte",
        fontName="Helvetica-Bold",
        fontSize=11,
        leading=14,
        alignment=TA_LEFT,
        spaceBefore=8,
        spaceAfter=6,
    )

    estilo_normal = ParagraphStyle(
        name="NormalEstadoArte",
        fontName="Helvetica",
        fontSize=9.5,
        leading=13,
        alignment=TA_JUSTIFY,
        spaceAfter=6,
    )

    estilo_tabla = ParagraphStyle(
        name="TablaEstadoArte",
        fontName="Helvetica",
        fontSize=7.2,
        leading=8.6,
        alignment=TA_LEFT,
    )

    estilo_tabla_centro = ParagraphStyle(
        name="TablaCentroEstadoArte",
        fontName="Helvetica",
        fontSize=7.2,
        leading=8.6,
        alignment=TA_CENTER,
    )

    estilo_tabla_header = ParagraphStyle(
        name="TablaHeaderEstadoArte",
        fontName="Helvetica-Bold",
        fontSize=7.4,
        leading=8.8,
        alignment=TA_CENTER,
    )

    historia = []

    contenido = datos["contenido_estado_arte"]

    historia.append(Paragraph("FORMATO ESTADO DEL ARTE<br/>RED TECNOPARQUE COLOMBIA", estilo_titulo))

    tabla_info = Table(
        [
            [
                Paragraph("<b>CÓDIGO DEL PROYECTO:</b>", estilo_tabla_header),
                Paragraph(datos["codigo_proyecto"], estilo_tabla),
            ],
            [
                Paragraph("<b>NOMBRE DEL PROYECTO:</b>", estilo_tabla_header),
                Paragraph(datos["nombre_proyecto"], estilo_tabla),
            ],
            [
                Paragraph("<b>NODO:</b>", estilo_tabla_header),
                Paragraph("ANGOSTURA", estilo_tabla),
            ],
            [
                Paragraph("<b>FECHA:</b>", estilo_tabla_header),
                Paragraph(datos["fecha_documento"].strftime("%B %Y").upper(), estilo_tabla),
            ],
        ],
        colWidths=[5.0 * cm, 11.8 * cm],
    )

    tabla_info.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.4, colors.black),
                ("BACKGROUND", (0, 0), (0, -1), colors.whitesmoke),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )

    historia.append(tabla_info)
    historia.append(Spacer(1, 0.3 * cm))

    historia.append(Paragraph("1. Introducción", estilo_subtitulo))
    historia.append(Paragraph(contenido.get("introduccion", ""), estilo_normal))

    historia.append(Paragraph("2. Objetivos del Estado del Arte", estilo_subtitulo))
    for objetivo in contenido.get("objetivos", []):
        historia.append(Paragraph(f"• {objetivo}", estilo_normal))

    historia.append(Paragraph("3. Antecedentes y Contexto", estilo_subtitulo))
    historia.append(Paragraph(contenido.get("antecedentes_contexto", ""), estilo_normal))

    historia.append(Paragraph("4. Contexto Nacional e Internacional", estilo_subtitulo))
    historia.append(Paragraph(contenido.get("contexto_nacional_internacional", ""), estilo_normal))

    historia.append(Paragraph("5. Proyectos Similares", estilo_subtitulo))

    proyectos_data = [
        [
            Paragraph("<b>No.</b>", estilo_tabla_header),
            Paragraph("<b>Nombre del Proyecto</b>", estilo_tabla_header),
            Paragraph("<b>Enlace de consulta</b>", estilo_tabla_header),
        ]
    ]

    for item in contenido.get("proyectos_similares", []):
        proyectos_data.append(
            [
                Paragraph(str(item.get("numero", "")), estilo_tabla_centro),
                Paragraph(str(item.get("nombre", "")), estilo_tabla),
                Paragraph(str(item.get("enlace", "")), estilo_tabla),
            ]
        )

    tabla_proyectos = Table(
        proyectos_data,
        colWidths=[1.2 * cm, 8.0 * cm, 7.6 * cm],
        repeatRows=1,
    )

    tabla_proyectos.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.4, colors.black),
                ("BACKGROUND", (0, 0), (-1, 0), colors.whitesmoke),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )

    historia.append(tabla_proyectos)

    historia.append(Paragraph("6. Estado del Arte de Tecnologías Relevantes", estilo_subtitulo))

    historia.append(Paragraph("6.1 Tecnologías previstas para el proyecto", estilo_subtitulo))
    for item in contenido.get("tecnologias_relevantes", []):
        historia.append(Paragraph(f"<b>{item.get('tecnologia', '')}.</b> {item.get('analisis', '')} ({item.get('cita_apa', '')})", estilo_normal))

    historia.append(Paragraph("6.2 Tecnologías emergentes contrastadas", estilo_subtitulo))
    for item in contenido.get("tecnologias_emergentes", []):
        historia.append(Paragraph(f"<b>{item.get('tecnologia', '')}.</b> {item.get('analisis', '')} ({item.get('cita_apa', '')})", estilo_normal))

    historia.append(Paragraph("6.3 Artículos científicos y técnicos de validación", estilo_subtitulo))

    articulos_data = [
        [
            Paragraph("<b>No.</b>", estilo_tabla_header),
            Paragraph("<b>Tecnología</b>", estilo_tabla_header),
            Paragraph("<b>Artículo / Fuente</b>", estilo_tabla_header),
            Paragraph("<b>Enlace</b>", estilo_tabla_header),
        ]
    ]

    for item in contenido.get("articulos_validacion", []):
        articulos_data.append(
            [
                Paragraph(str(item.get("numero", "")), estilo_tabla_centro),
                Paragraph(str(item.get("tecnologia", "")), estilo_tabla),
                Paragraph(str(item.get("articulo", "")), estilo_tabla),
                Paragraph(str(item.get("enlace", "")), estilo_tabla),
            ]
        )

    tabla_articulos = Table(
        articulos_data,
        colWidths=[1.1 * cm, 3.8 * cm, 6.0 * cm, 5.9 * cm],
        repeatRows=1,
    )

    tabla_articulos.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.4, colors.black),
                ("BACKGROUND", (0, 0), (-1, 0), colors.whitesmoke),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )

    historia.append(tabla_articulos)

    historia.append(Paragraph("7. Conclusiones", estilo_subtitulo))
    historia.append(Paragraph(contenido.get("conclusiones", ""), estilo_normal))

    historia.append(Paragraph("8. Referencias Bibliográficas", estilo_subtitulo))
    for referencia in contenido.get("bibliografia", []):
        historia.append(Paragraph(f"• {referencia}", estilo_normal))

    doc.build(
        historia,
        onFirstPage=encabezado_pie,
        onLaterPages=encabezado_pie,
    )

    datos_json = dict(datos)
    datos_json["fecha_documento"] = datos["fecha_documento"].strftime("%d/%m/%Y")
    datos_json["ruta_pdf"] = ruta_pdf

    guardar_datos_json(datos_json, ruta="datos_estado_arte.json")

    return ruta_pdf

# =====================================================
# ACTA DE EJECUCIÓN - ASESORÍAS Y USOS
# =====================================================

VALOR_HORA_EXPERTO = 25266


def formato_moneda_colombiana(valor: float | int) -> str:
    try:
        valor_int = int(round(float(valor)))
    except Exception:
        valor_int = 0

    return "$" + f"{valor_int:,.0f}".replace(",", ".")


def distribuir_fechas_para_asesorias(
    fecha_inicio: date,
    fecha_fin: date,
    dias_semana: list[str],
    cantidad_asesorias: int
) -> list[date]:
    fechas_disponibles = obtener_fechas_programadas(fecha_inicio, fecha_fin, dias_semana)

    if not fechas_disponibles:
        return []

    if cantidad_asesorias <= 0:
        return []

    fechas_resultado = []

    for i in range(cantidad_asesorias):
        indice = round(i * (len(fechas_disponibles) - 1) / max(cantidad_asesorias - 1, 1))
        fechas_resultado.append(fechas_disponibles[indice])

    return fechas_resultado


def generar_asesorias_ejecucion_modo_prueba(
    descripcion_proyecto: str,
    cantidad_asesorias: int,
    fechas_asesorias: list[date],
    horas_por_asesoria: float
) -> list[dict]:
    actividades_base = [
        "Revisión del avance técnico del proyecto y validación de requerimientos definidos en la fase de planeación.",
        "Asesoría para la estructuración de componentes técnicos y definición de criterios de diseño de la solución.",
        "Acompañamiento en la selección de tecnologías, materiales, herramientas o recursos requeridos para el desarrollo.",
        "Revisión del diseño preliminar, ajustes funcionales y recomendaciones para la construcción o implementación.",
        "Asesoría en pruebas iniciales, verificación de resultados y análisis de funcionamiento de la solución propuesta.",
        "Acompañamiento en ajustes técnicos derivados de la validación del prototipo o componente desarrollado.",
        "Revisión de evidencias, documentación técnica y consolidación de avances del proyecto.",
        "Asesoría para cierre parcial de actividades, identificación de mejoras y definición de siguientes pasos.",
    ]

    asesorias = []

    for i in range(cantidad_asesorias):
        descripcion = actividades_base[i % len(actividades_base)]
        fecha = fechas_asesorias[i] if i < len(fechas_asesorias) else date.today()

        asesorias.append(
            {
                "fecha": fecha.strftime("%d/%m/%Y"),
                "horas": float(horas_por_asesoria),
                "descripcion": descripcion,
            }
        )

    return asesorias


def generar_asesorias_ejecucion_con_chatgpt(
    descripcion_proyecto: str,
    cantidad_asesorias: int,
    fechas_asesorias: list[date],
    horas_por_asesoria: float,
    modelo: str = "gpt-4.1-mini"
) -> list[dict]:
    if OpenAI is None:
        raise ImportError("No está instalada la librería openai. Instálala con: pip install openai")

    api_key = obtener_api_key()
    if not api_key:
        raise ValueError("No se encontró OPENAI_API_KEY. Configúrala como variable de entorno o en .streamlit/secrets.toml.")

    fechas_texto = [f.strftime("%d/%m/%Y") for f in fechas_asesorias]

    if not fechas_texto:
        return generar_asesorias_ejecucion_modo_prueba(
            descripcion_proyecto,
            cantidad_asesorias,
            fechas_asesorias,
            horas_por_asesoria
        )

    client = OpenAI(api_key=api_key)

    instrucciones = """
Eres un experto en seguimiento técnico de proyectos de base tecnológica de la Red Tecnoparque SENA.

Debes generar actividades de asesoría y uso de infraestructura para un acta de ejecución.
Las actividades deben ser técnicas, claras, verificables y coherentes con el desarrollo del proyecto.
No inventes nombres de personas, códigos, valores económicos ni entidades.
Usa únicamente las fechas suministradas.
Responde únicamente con JSON válido, sin markdown y sin explicaciones adicionales.
"""

    entrada = f"""
Descripción general del proyecto:
{descripcion_proyecto}

Cantidad de asesorías requeridas:
{cantidad_asesorias}

Horas por asesoría:
{horas_por_asesoria}

Fechas disponibles para distribuir las asesorías:
{", ".join(fechas_texto)}

Genera exactamente {cantidad_asesorias} registros.

Formato JSON obligatorio:
{{
  "asesorias": [
    {{
      "fecha": "dd/mm/aaaa",
      "horas": {horas_por_asesoria},
      "descripcion": "Descripción técnica de la asesoría realizada"
    }}
  ]
}}
"""

    try:
        respuesta = client.responses.create(
            model=modelo,
            instructions=instrucciones,
            input=entrada,
            temperature=0.35,
        )

        texto = limpiar_respuesta_json(respuesta.output_text)
        datos = json.loads(texto)

        asesorias = datos.get("asesorias", [])

        asesorias_limpias = []
        fechas_validas = set(fechas_texto)

        for i, item in enumerate(asesorias):
            fecha_item = str(item.get("fecha", "")).strip()
            if fecha_item not in fechas_validas:
                fecha_item = fechas_texto[min(i, len(fechas_texto) - 1)]

            asesorias_limpias.append(
                {
                    "fecha": fecha_item,
                    "horas": float(horas_por_asesoria),
                    "descripcion": str(item.get("descripcion", "")).strip() or "Asesoría técnica para seguimiento del proyecto.",
                }
            )

        if len(asesorias_limpias) < cantidad_asesorias:
            faltantes = cantidad_asesorias - len(asesorias_limpias)
            adicionales = generar_asesorias_ejecucion_modo_prueba(
                descripcion_proyecto,
                faltantes,
                fechas_asesorias[-faltantes:] if faltantes <= len(fechas_asesorias) else fechas_asesorias,
                horas_por_asesoria
            )
            asesorias_limpias.extend(adicionales)

        return asesorias_limpias[:cantidad_asesorias]

    except Exception:
        return generar_asesorias_ejecucion_modo_prueba(
            descripcion_proyecto,
            cantidad_asesorias,
            fechas_asesorias,
            horas_por_asesoria
        )


def generar_pdf_acta_ejecucion(datos: dict) -> str:
    if canvas is None:
        raise ImportError("No está instalada reportlab. Instálala con: pip install reportlab")

    Path(CARPETA_SALIDA).mkdir(parents=True, exist_ok=True)

    nombre_archivo = (
        f"Acta_Ejecucion_"
        f"{safe_filename(datos.get('codigo_proyecto', 'proyecto'))}_"
        f"{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    )
    ruta_pdf = str(Path(CARPETA_SALIDA) / nombre_archivo)

    page_size = landscape(letter)
    page_width, page_height = page_size

    doc = SimpleDocTemplate(
        ruta_pdf,
        pagesize=page_size,
        rightMargin=1.2 * cm,
        leftMargin=1.2 * cm,
        topMargin=1.5 * cm,
        bottomMargin=1.0 * cm,
    )

    estilo_titulo = ParagraphStyle(
        name="TituloActaEjecucion",
        fontName="Helvetica-Bold",
        fontSize=10.5,
        leading=13,
        alignment=TA_CENTER,
    )

    estilo_header = ParagraphStyle(
        name="HeaderActaEjecucion",
        fontName="Helvetica-Bold",
        fontSize=7.4,
        leading=8.5,
        alignment=TA_CENTER,
    )

    estilo_celda = ParagraphStyle(
        name="CeldaActaEjecucion",
        fontName="Helvetica",
        fontSize=7.2,
        leading=8.5,
        alignment=TA_CENTER,
    )

    estilo_celda_left = ParagraphStyle(
        name="CeldaLeftActaEjecucion",
        fontName="Helvetica",
        fontSize=7.2,
        leading=8.5,
        alignment=TA_LEFT,
    )

    estilo_negrita_left = ParagraphStyle(
        name="NegritaLeftActaEjecucion",
        fontName="Helvetica-Bold",
        fontSize=7.4,
        leading=8.5,
        alignment=TA_LEFT,
    )

    historia = []

    # Encabezado principal
    logo_path = obtener_ruta_logo_sena()
    if logo_path:
        try:
            logo = Image(logo_path, width=2.2 * cm, height=1.8 * cm)
        except Exception:
            logo = Paragraph("SENA", estilo_titulo)
    else:
        logo = Paragraph("SENA", estilo_titulo)

    encabezado_data = [
        [
            logo,
            Paragraph("Seguimiento de Asesorías y Usos de Infraestructura", estilo_titulo),
        ],
        [
            "",
            Paragraph(f"ACTA No. 02 del proyecto No {datos['codigo_proyecto']}", estilo_titulo),
        ],
    ]

    tabla_encabezado = Table(
        encabezado_data,
        colWidths=[7.0 * cm, 19.0 * cm],
        rowHeights=[1.0 * cm, 0.8 * cm],
    )

    tabla_encabezado.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.45, colors.black),
                ("SPAN", (0, 0), (0, 1)),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("ALIGN", (0, 0), (0, 1), "CENTER"),
            ]
        )
    )

    historia.append(tabla_encabezado)

    # Información general
    gris = colors.HexColor("#BFBFBF")

    seccion_info = Table(
        [[Paragraph("Información general", estilo_header)]],
        colWidths=[26.0 * cm],
        rowHeights=[0.55 * cm],
    )
    seccion_info.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.45, colors.black),
                ("BACKGROUND", (0, 0), (-1, -1), gris),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]
        )
    )
    historia.append(seccion_info)

    info_data = [
        [
            Paragraph("Código del proyecto", estilo_header),
            Paragraph("Nombre del proyecto", estilo_header),
            Paragraph("Experto a cargo del proyecto", estilo_header),
            Paragraph("Sublínea tecnológica", estilo_header),
        ],
        [
            Paragraph(datos["codigo_proyecto"], estilo_celda),
            Paragraph(datos["nombre_proyecto"], estilo_celda),
            Paragraph(datos["nombre_experto"], estilo_celda),
            Paragraph(datos["sublinea_tecnologica"], estilo_celda),
        ],
    ]

    tabla_info = Table(
        info_data,
        colWidths=[5.8 * cm, 10.0 * cm, 5.2 * cm, 5.0 * cm],
        rowHeights=[0.55 * cm, 2.05 * cm],
    )

    tabla_info.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.45, colors.black),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]
        )
    )

    historia.append(tabla_info)

    # Talentos
    seccion_talentos = Table(
        [[Paragraph("Talentos del Proyecto", estilo_header)]],
        colWidths=[26.0 * cm],
        rowHeights=[0.55 * cm],
    )
    seccion_talentos.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.45, colors.black),
                ("BACKGROUND", (0, 0), (-1, -1), gris),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]
        )
    )
    historia.append(seccion_talentos)

    talentos_data = [
        [
            Paragraph("Número de documento", estilo_header),
            Paragraph("Nombres y apellidos", estilo_header),
            Paragraph("Número de contacto", estilo_header),
        ],
        [
            Paragraph(datos["documento_talento"], estilo_celda),
            Paragraph(datos["nombre_talento"], estilo_celda),
            Paragraph(datos["telefono_talento"], estilo_celda),
        ],
    ]

    tabla_talentos = Table(
        talentos_data,
        colWidths=[8.5 * cm, 11.5 * cm, 6.0 * cm],
        rowHeights=[0.55 * cm, 0.75 * cm],
    )

    tabla_talentos.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.45, colors.black),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]
        )
    )

    historia.append(tabla_talentos)

    # Asesorías y usos
    seccion_asesorias = Table(
        [[Paragraph("Asesorías y usos", estilo_header)]],
        colWidths=[26.0 * cm],
        rowHeights=[0.55 * cm],
    )
    seccion_asesorias.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.45, colors.black),
                ("BACKGROUND", (0, 0), (-1, -1), gris),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]
        )
    )
    historia.append(seccion_asesorias)

    asesorias_data = [
        [
            Paragraph("Fecha de la Asesoría y uso de infraestructura", estilo_header),
            Paragraph("Horas de Asesoría", estilo_header),
            Paragraph("Descripción", estilo_header),
        ]
    ]

    for item in datos["asesorias"]:
        asesorias_data.append(
            [
                Paragraph(str(item.get("fecha", "")), estilo_celda),
                Paragraph(str(item.get("horas", "")), estilo_celda),
                Paragraph(str(item.get("descripcion", "")), estilo_celda_left),
            ]
        )

    asesorias_data.append(
        [
            "",
            "",
            Paragraph(
                f"<b>Valor total de la asesoría (Valor hora: {formato_moneda_colombiana(VALOR_HORA_EXPERTO)}) "
                f"{formato_moneda_colombiana(datos['total_honorarios'])}</b>",
                estilo_negrita_left,
            ),
        ]
    )

    row_heights_asesorias = [0.55 * cm] + [0.75 * cm for _ in datos["asesorias"]] + [0.55 * cm]

    tabla_asesorias = Table(
        asesorias_data,
        colWidths=[7.0 * cm, 4.0 * cm, 15.0 * cm],
        rowHeights=row_heights_asesorias,
        repeatRows=1,
    )

    tabla_asesorias.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.45, colors.black),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("SPAN", (0, len(asesorias_data) - 1), (1, len(asesorias_data) - 1)),
            ]
        )
    )

    historia.append(tabla_asesorias)

    # Materiales y equipos
    seccion_materiales = Table(
        [[Paragraph("Materiales, equipos e insumos utilizados en el proyecto", estilo_header)]],
        colWidths=[26.0 * cm],
        rowHeights=[0.55 * cm],
    )
    seccion_materiales.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.45, colors.black),
                ("BACKGROUND", (0, 0), (-1, -1), gris),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]
        )
    )
    historia.append(seccion_materiales)

    materiales_data = [
        [
            Paragraph("Nombre del equipo/Material Usado", estilo_header),
            Paragraph("horas de uso/cantidad", estilo_header),
            Paragraph("valor total", estilo_header),
        ]
    ]

    for item in datos["equipos_materiales"]:
        materiales_data.append(
            [
                Paragraph(str(item.get("nombre", "")), estilo_celda_left),
                Paragraph(str(item.get("cantidad_horas", "")), estilo_celda),
                Paragraph(formato_moneda_colombiana(item.get("valor_total", 0)), estilo_celda),
            ]
        )

    materiales_data.append(
        [
            "",
            Paragraph("<b>Costo total uso de equipos e infraestructura</b>", estilo_negrita_left),
            Paragraph(f"<b>{formato_moneda_colombiana(datos['total_equipos_materiales'])}</b>", estilo_celda),
        ]
    )

    row_heights_materiales = [0.55 * cm] + [0.65 * cm for _ in datos["equipos_materiales"]] + [0.55 * cm]

    tabla_materiales = Table(
        materiales_data,
        colWidths=[8.0 * cm, 12.0 * cm, 6.0 * cm],
        rowHeights=row_heights_materiales,
        repeatRows=1,
    )

    tabla_materiales.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.45, colors.black),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("SPAN", (0, len(materiales_data) - 1), (0, len(materiales_data) - 1)),
            ]
        )
    )

    historia.append(tabla_materiales)

    # Total general
    total_general = Table(
        [
            [
                Paragraph("Costo total honorarios experto más valor de uso de equipos y materiales", estilo_negrita_left),
                Paragraph(f"<b>{formato_moneda_colombiana(datos['total_general'])}</b>", estilo_celda),
            ]
        ],
        colWidths=[20.0 * cm, 6.0 * cm],
        rowHeights=[0.65 * cm],
    )

    total_general.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.45, colors.black),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]
        )
    )

    historia.append(total_general)

    # Firmas
    seccion_firmas = Table(
        [[Paragraph("Firma Expertos y Talentos", estilo_header)]],
        colWidths=[26.0 * cm],
        rowHeights=[0.55 * cm],
    )
    seccion_firmas.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.45, colors.black),
                ("BACKGROUND", (0, 0), (-1, -1), gris),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]
        )
    )
    historia.append(seccion_firmas)

    firmas_data = [
        [
            Paragraph(f"{datos['nombre_experto']} - Experto", estilo_celda),
            Paragraph(f"{datos['nombre_talento']} - Talento Interlocutor", estilo_celda),
        ],
        [
            Paragraph("", estilo_celda),
            Paragraph("", estilo_celda),
        ],
    ]

    tabla_firmas = Table(
        firmas_data,
        colWidths=[13.0 * cm, 13.0 * cm],
        rowHeights=[0.75 * cm, 1.2 * cm],
    )

    tabla_firmas.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.45, colors.black),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]
        )
    )

    historia.append(tabla_firmas)

    doc.build(historia)

    datos_json = dict(datos)
    datos_json["ruta_pdf"] = ruta_pdf
    guardar_datos_json(datos_json, ruta="datos_acta_ejecucion.json")

    return ruta_pdf

# =====================================================
# ACTA DE CIERRE - FASE DE CIERRE
# =====================================================

FORMATO_ACTA_CIERRE = "GOR-F-084 V02"


def generar_objetivo_cierre(codigo_proyecto: str, nombre_proyecto: str) -> str:
    return (
        f"Dar por finalizada la ejecución del proyecto {codigo_proyecto} - {nombre_proyecto}, "
        "revisando los objetivos alcanzados, entregables, desempeño, y obtener la aprobación formal "
        "del cierre del proyecto."
    )

def generar_evidencias_cierre_modo_prueba(nombre_proyecto: str, evidencias_producto: str) -> dict:
    texto_lower = evidencias_producto.lower()

    if "susceptible de inscribir un nuevo proyecto" in texto_lower or "inscribir un nuevo proyecto" in texto_lower:
        conclusion_adicional = (
            "Adicionalmente, se identifica que el proyecto es susceptible de inscribir un nuevo proyecto "
            "o una nueva idea de base tecnológica, con el fin de continuar su fortalecimiento, validación, "
            "escalamiento o desarrollo de nuevas funcionalidades."
        )
    else:
        conclusion_adicional = (
            "Se evaluará la posibilidad de inscribir un nuevo prototipo, idea o proyecto de base tecnológica, "
            "de acuerdo con los resultados obtenidos, las oportunidades de mejora identificadas y el potencial "
            "de continuidad técnica del desarrollo alcanzado."
        )

    entregables = [
        item.strip(" -•0123456789.)")
        for item in evidencias_producto.replace(";", "\n").split("\n")
        if item.strip()
    ]

    if entregables:
        entregables_texto = " ".join(
            [
                f"Se evidencia el desarrollo y entrega de {entregable}, como componente técnico asociado al prototipo, producto o resultado obtenido durante la ejecución del proyecto."
                for entregable in entregables
            ]
        )
    else:
        entregables_texto = (
            "Se registran evidencias asociadas al desarrollo del prototipo, incluyendo productos, componentes, "
            "documentos técnicos, validaciones, diseños, pruebas o implementaciones generadas durante la ejecución del proyecto."
        )

    return {
        "evidencias_normatividad": (
            f"De acuerdo con la naturaleza técnica del proyecto {nombre_proyecto}, se identifican referentes normativos "
            "aplicables para orientar la validación, documentación y cierre técnico del prototipo desarrollado. "
            "Según el tipo de solución, pueden considerarse Normas Técnicas Colombianas NTC relacionadas con gestión "
            "de calidad, documentación técnica, seguridad de producto, requisitos de operación, trazabilidad, "
            "medición, validación funcional y buenas prácticas de desarrollo tecnológico. De manera complementaria, "
            "pueden tomarse como referencia normas internacionales ISO aplicables a sistemas de gestión de calidad, "
            "diseño de productos, procesos de ensayo, documentación de resultados, seguridad de operación, "
            "interoperabilidad tecnológica, pruebas funcionales y validación de componentes. Estas referencias permiten "
            "establecer criterios mínimos para verificar que el prototipo, producto o solución desarrollada cuente con "
            "evidencias técnicas suficientes, trazabilidad documental y condiciones adecuadas para futuras fases de "
            "fortalecimiento, transferencia, mejora o escalamiento."
        ),
        "evidencias_modelo_negocio": (
            "Se adjunta el Modelo Canvas aplicado al Proyecto de Base Tecnológica, como herramienta de análisis "
            "para la identificación de la propuesta de valor, segmentos de cliente, canales, recursos clave, "
            "actividades clave, aliados estratégicos, estructura de costos y fuentes de ingreso."
        ),
        "evidencias_pruebas_documentadas": (
            "Se adjunta el Informe Técnico Final, en el cual se documenta la metodología desarrollada, los procesos "
            "de validación, las pruebas realizadas, los resultados obtenidos y la implementación técnica del proyecto."
        ),
        "evidencias_prototipo": entregables_texto,
        "conclusion_adicional": conclusion_adicional,
    }




def generar_evidencias_cierre_con_chatgpt(
    nombre_proyecto: str,
    evidencias_producto: str,
    modelo: str = "gpt-4.1-mini"
) -> dict:
    if OpenAI is None:
        raise ImportError("No está instalada la librería openai. Instálala con: pip install openai")

    api_key = obtener_api_key()
    if not api_key:
        raise ValueError("No se encontró OPENAI_API_KEY. Configúrala como variable de entorno o en .streamlit/secrets.toml.")

    client = OpenAI(api_key=api_key)

    instrucciones = """
Eres un experto en cierre técnico de proyectos de base tecnológica de la Red Tecnoparque SENA.

Debes generar evidencias para un Acta de Cierre institucional.
Responde únicamente en JSON válido.
No uses markdown.
No inventes nombres de personas, fechas, códigos ni entidades.
Las normas deben ser genéricas y aplicables según el tipo de prototipo, citando NTC o normas internacionales cuando corresponda.
"""

    entrada = f"""
Proyecto:
{nombre_proyecto}

Texto genérico ingresado por el usuario en Evidencias del Producto:
{evidencias_producto}

Genera un JSON con esta estructura exacta:

{{
  "evidencias_normatividad": "Redacción técnica sobre NTC o normas internacionales aplicables al tipo de prototipo.",
  "evidencias_modelo_negocio": "Frase estándar indicando que se adjunta Modelo Canvas aplicado al PBT.",
  "evidencias_pruebas_documentadas": "Frase indicando que se adjunta Informe Técnico Final con metodología, validación e implementación.",
  "evidencias_prototipo": "Lista redactada de entregables específicos extraídos del texto de evidencias del producto. Debe entenderse como evidencias de prototipo y entregables desarrollados.",
  "conclusion_adicional": "Conclusión adicional solo si el texto menciona que el proyecto es susceptible de inscribir un nuevo proyecto. Si no aplica, dejar vacío."
}}
"""

    try:
        respuesta = client.responses.create(
            model=modelo,
            instructions=instrucciones,
            input=entrada,
            temperature=0.25,
        )

        texto = limpiar_respuesta_json(respuesta.output_text)
        datos = json.loads(texto)

        campos = [
            "evidencias_normatividad",
            "evidencias_modelo_negocio",
            "evidencias_pruebas_documentadas",
            "evidencias_prototipo",
            "conclusion_adicional",
        ]

        for campo in campos:
            if campo not in datos or not isinstance(datos[campo], str):
                datos[campo] = ""

        if not datos.get("evidencias_prototipo", "").strip():
            datos_base = generar_evidencias_cierre_modo_prueba(nombre_proyecto, evidencias_producto)
            datos["evidencias_prototipo"] = datos_base["evidencias_prototipo"]

        if not datos.get("evidencias_normatividad", "").strip() or len(datos["evidencias_normatividad"]) < 250:
            datos_base = generar_evidencias_cierre_modo_prueba(nombre_proyecto, evidencias_producto)
            datos["evidencias_normatividad"] = datos_base["evidencias_normatividad"]
                

        texto_lower = evidencias_producto.lower()

        if (
            "susceptible de inscribir un nuevo proyecto" in texto_lower
            or "inscribir un nuevo proyecto" in texto_lower
        ):
            if not datos["conclusion_adicional"]:
                datos["conclusion_adicional"] = (
                    "Adicionalmente, se identifica que el proyecto es susceptible de inscribir un nuevo proyecto "
                    "o una nueva idea de base tecnológica, con el fin de continuar su fortalecimiento, validación, "
                    "escalamiento o desarrollo de nuevas funcionalidades."
                )
        else:
            datos["conclusion_adicional"] = (
                "Se evaluará la posibilidad de inscribir un nuevo prototipo, idea o proyecto de base tecnológica, "
                "de acuerdo con los resultados obtenidos, las oportunidades de mejora identificadas y el potencial "
                "de continuidad técnica del desarrollo alcanzado."
            )
            

        return datos

    except Exception:
        return generar_evidencias_cierre_modo_prueba(nombre_proyecto, evidencias_producto)


def generar_pdf_acta_cierre(datos: dict) -> str:
    if canvas is None:
        raise ImportError("No está instalada reportlab. Instálala con: pip install reportlab")

    Path(CARPETA_SALIDA).mkdir(parents=True, exist_ok=True)

    nombre_archivo = (
        f"Acta_Cierre_"
        f"{safe_filename(datos.get('codigo_proyecto', 'proyecto'))}_"
        f"{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    )
    ruta_pdf = str(Path(CARPETA_SALIDA) / nombre_archivo)

    c = canvas.Canvas(ruta_pdf, pagesize=letter)
    page_width, page_height = letter

    x0 = 24
    table_w = page_width - 48
    logo_top_y = 785
    y_top_content = 705
    y_safe_bottom = 58

    FONT_TITLE = 10.5
    FONT_SECTION = 10.5
    FONT_BODY = 7.1
    FONT_SMALL = 6.7
    FONT_TINY = 6.3

    def iniciar_pagina() -> float:
        draw_logo(c, page_width, logo_top_y)
        return y_top_content

    def cerrar_pagina() -> None:
        c.setFillColor(colors.grey)
        c.setFont("Helvetica", 10)
        c.drawCentredString(page_width / 2, 20, FORMATO_ACTA_CIERRE)
        c.setFillColor(colors.black)
        c.showPage()

    def asegurar_espacio(y_actual: float, alto_requerido: float) -> float:
        if y_actual - alto_requerido < y_safe_bottom:
            cerrar_pagina()
            return iniciar_pagina()
        return y_actual

    def alto_texto_local(texto: str, ancho: float, base: float = 24, font_size: float = FONT_BODY, label: str | None = None, max_h: float = 80) -> float:
        padding = 10
        usable_w = max(ancho - padding, 20)
        texto_total = f"{label or ''} {texto or ''}".strip()
        lineas = wrap_text(texto_total, "Helvetica", font_size, usable_w)
        h = max(base, len(lineas) * (font_size + 2.2) + 12)
        return min(h, max_h)

    y = iniciar_pagina()

    # Título
    h = 22
    draw_cell(
    c,
    x0,
    y - h,
    table_w,
    h,
    f"ACTA No. 03 del proyecto No {datos.get('codigo_proyecto', '')}",
    font="Helvetica-Bold",
    size=FONT_TITLE,
    center=True,
)
    y -= h

    # Nombre del comité o reunión
    nombre_comite = f"Acta de cierre del proyecto {datos.get('codigo_proyecto', '')} - {datos.get('nombre_proyecto', '')}"
    h = alto_texto_local(
        nombre_comite,
        table_w,
        base=38,
        font_size=FONT_BODY,
        label="NOMBRE DEL COMITÉ O DE LA REUNIÓN:",
        max_h=58,
    )
    draw_cell(
        c,
        x0,
        y - h,
        table_w,
        h,
        nombre_comite,
        label="NOMBRE DEL COMITÉ O DE LA REUNIÓN:",
        size=FONT_BODY,
    )
    y -= h

    # Ciudad, fecha, hora
    h = 38
    w1 = table_w * 0.58
    w2 = table_w * 0.21
    w3 = table_w * 0.21
    draw_cell(
        c,
        x0,
        y - h,
        w1,
        h,
        f"Campoalegre (Huila) - {datos.get('fecha_iso', '')}",
        label="CIUDAD Y FECHA:",
        size=FONT_BODY,
    )
    draw_cell(c, x0 + w1, y - h, w2, h, datos.get("hora_inicio", ""), label="HORA INICIO:", size=FONT_BODY)
    draw_cell(c, x0 + w1 + w2, y - h, w3, h, datos.get("hora_fin", ""), label="HORA FIN:", size=FONT_BODY)
    y -= h

    # Lugar y centro
    h = 46
    w_lugar = table_w * 0.55
    w_dir = table_w * 0.45
    draw_cell(
        c,
        x0,
        y - h,
        w_lugar,
        h,
        "Tecnoparque Angostura - Campoalegre, Huila",
        label="LUGAR Y/O ENLACE:",
        size=FONT_BODY,
    )
    draw_cell(
        c,
        x0 + w_lugar,
        y - h,
        w_dir,
        h,
        "Centro de Formación Agroindustrial / SENA Regional Huila",
        label="DIRECCIÓN / REGIONAL / CENTRO:",
        size=FONT_BODY,
    )
    y -= h

    # Objetivo de la reunión
    objetivo_reunion = datos.get(
    "objetivo_cierre",
    f"Dar por finalizada la ejecución del proyecto {datos.get('codigo_proyecto', '')} - {datos.get('nombre_proyecto', '')}, revisando los objetivos alcanzados, entregables, desempeño, y obtener la aprobación formal del cierre del proyecto."
)
    h = alto_texto_local(
        objetivo_reunion,
        table_w,
        base=40,
        font_size=FONT_BODY,
        label="OBJETIVO(S) DE LA REUNIÓN:",
        max_h=56,
    )
    draw_cell(c, x0, y - h, table_w, h, objetivo_reunion, label="OBJETIVO(S) DE LA REUNIÓN:", size=FONT_BODY)
    y -= h

    # Desarrollo
    h = 22
    y = asegurar_espacio(y, h)
    draw_cell(c, x0, y - h, table_w, h, "DESARROLLO DE LA REUNIÓN", font="Helvetica-Bold", size=FONT_SECTION, center=True)
    y -= h

    # Código y nombre
    codigo_nombre = f"{datos.get('codigo_proyecto', '')} - {datos.get('nombre_proyecto', '')}"
    h = alto_texto_local(codigo_nombre, table_w, base=26, font_size=FONT_BODY, label="Código y nombre del Proyecto:", max_h=46)
    y = asegurar_espacio(y, h)
    draw_cell(c, x0, y - h, table_w, h, codigo_nombre, label="Código y nombre del Proyecto:", size=FONT_BODY)
    y -= h

    # TRL y aporte
    h = 32
    w_trl = table_w * 0.35
    w_aporte = table_w * 0.65
    y = asegurar_espacio(y, h)
    draw_cell(c, x0, y - h, w_trl, h, datos.get("trl_obtenido", ""), label="TRL OBTENIDO:", size=FONT_BODY)
    draw_cell(
        c,
        x0 + w_trl,
        y - h,
        w_aporte,
        h,
        formato_moneda_colombiana(datos.get("aporte_tecnoparque", 0)),
        label="APORTE ESTIMADO DE TECNOPARQUE:",
        size=FONT_BODY,
    )
    y -= h

    # Objetivo de cierre
    h = 22
    y = asegurar_espacio(y, h)
    draw_cell(c, x0, y - h, table_w, h, "OBJETIVO DEL CIERRE", font="Helvetica-Bold", size=FONT_SECTION, center=True)
    y -= h

    objetivo_cierre = datos.get("objetivo_cierre", "")
    h = alto_texto_local(objetivo_cierre, table_w, base=42, font_size=FONT_BODY, max_h=72)
    y = asegurar_espacio(y, h)
    draw_cell(c, x0, y - h, table_w, h, objetivo_cierre, size=FONT_BODY)
    y -= h

        # Objetivos iniciales y cumplimiento
    h = 22
    y = asegurar_espacio(y, h)
    draw_cell(
        c,
        x0,
        y - h,
        table_w,
        h,
        "OBJETIVOS INICIALES DEL PROYECTO Y CUMPLIMIENTO",
        font="Helvetica-Bold",
        size=FONT_SECTION,
        center=True,
    )
    y -= h

    objetivos_iniciales = datos.get("objetivos_iniciales", [])
    if not objetivos_iniciales:
        objetivos_iniciales = ["No se registraron objetivos iniciales."]

    col_num = 38
    col_obj = table_w - 130
    col_cumple = 92

    h = 24
    y = asegurar_espacio(y, h)
    draw_cell(c, x0, y - h, col_num, h, "No.", font="Helvetica-Bold", size=FONT_SMALL, center=True)
    draw_cell(c, x0 + col_num, y - h, col_obj, h, "Objetivo inicial", font="Helvetica-Bold", size=FONT_SMALL, center=True)
    draw_cell(c, x0 + col_num + col_obj, y - h, col_cumple, h, "Cumplió", font="Helvetica-Bold", size=FONT_SMALL, center=True)
    y -= h

    for idx, obj in enumerate(objetivos_iniciales, start=1):
        h = alto_texto_local(obj, col_obj, base=28, font_size=FONT_SMALL, max_h=58)
        y = asegurar_espacio(y, h)

        draw_cell(c, x0, y - h, col_num, h, str(idx), font="Helvetica-Bold", size=FONT_SMALL, center=True)
        draw_cell(c, x0 + col_num, y - h, col_obj, h, obj, size=FONT_SMALL)
        draw_cell(c, x0 + col_num + col_obj, y - h, col_cumple, h, "SI", font="Helvetica-Bold", size=FONT_SMALL, center=True)

        y -= h

    # Evidencias
    h = 22
    y = asegurar_espacio(y, h)
    draw_cell(
        c,
        x0,
        y - h,
        table_w,
        h,
        "EVIDENCIAS DEL PROYECTO",
        font="Helvetica-Bold",
        size=FONT_SECTION,
        center=True,
    )
    y -= h

    contenido = datos.get("evidencias_generadas", {})

    evidencias = [
        ("Evidencias de Normatividad:", contenido.get("evidencias_normatividad", "")),
        ("Evidencias de Modelo de Negocio:", contenido.get("evidencias_modelo_negocio", "")),
        ("Evidencias de Pruebas Documentadas:", contenido.get("evidencias_pruebas_documentadas", "")),
        ("Evidencias de Prototipo:", contenido.get("evidencias_prototipo", "")),
    ]

    for titulo, texto in evidencias:
        if not str(texto).strip():
            texto = "No se registró información específica para esta evidencia."

        h = alto_texto_local(
            texto,
            table_w,
            base=44,
            font_size=FONT_TINY,
            label=titulo,
            max_h=95,
        )

        y = asegurar_espacio(y, h)
        draw_cell(c, x0, y - h, table_w, h, texto, label=titulo, size=FONT_TINY)
        y -= h
        
    # Conclusiones
    h = 22
    y = asegurar_espacio(y, h)
    draw_cell(c, x0, y - h, table_w, h, "CONCLUSIONES", font="Helvetica-Bold", size=FONT_SECTION, center=True)
    y -= h

    conclusiones = "Se cumplieron a cabalidad todos los objetivos del proyecto."
    if contenido.get("conclusion_adicional"):
        conclusiones += " " + contenido.get("conclusion_adicional")

    h = alto_texto_local(conclusiones, table_w, base=34, font_size=FONT_BODY, max_h=64)
    y = asegurar_espacio(y, h)
    draw_cell(c, x0, y - h, table_w, h, conclusiones, size=FONT_BODY)
    y -= h

    # Asistentes y aprobación
    asistentes_alto = 22 + 34 + 36 + 36
    y = asegurar_espacio(y, asistentes_alto)

    draw_cell(c, x0, y - 22, table_w, 22, "ASISTENTES Y APROBACIÓN DE DECISIONES", font="Helvetica-Bold", size=FONT_SECTION, center=True)
    y -= 22

    col_w = [table_w * 0.22, table_w * 0.22, table_w * 0.16, table_w * 0.18, table_w * 0.22]
    headers = ["NOMBRE", "DEPENDENCIA / EMPRESA", "APRUEBA (SI/NO)", "OBSERVACIÓN", "FIRMA"]
    x = x0
    for w, header in zip(col_w, headers):
        draw_cell(c, x, y - 34, w, 34, header, font="Helvetica-Bold", size=FONT_SMALL, center=True)
        x += w
    y -= 34

    filas = [
        [datos.get("nombre_talento", ""), "Emprendedor/Empresario", "SI", "", ""],
        [datos.get("nombre_experto", ""), "SENA", "SI", "", ""],
    ]

    for fila in filas:
        x = x0
        for w, value in zip(col_w, fila):
            center = value in ["Emprendedor/Empresario", "SENA", "SI", ""]
            draw_cell(c, x, y - 36, w, 36, value, size=FONT_SMALL, center=center)
            x += w
        y -= 36

    cerrar_pagina()
    c.save()

    datos_json = dict(datos)
    datos_json["ruta_pdf"] = ruta_pdf
    guardar_datos_json(datos_json, ruta="datos_acta_cierre.json")

    return ruta_pdf


# =====================================================
# INFORME TÉCNICO FINAL - FASE DE CIERRE / DOCX
# =====================================================

TIPOS_PROYECTO_INFORME = [
    "Software",
    "Diseño industrial",
    "Electrónica y automatización",
    "Inteligencia artificial",
    "Desarrollo de marca e identidad visual",
    "Prototipo agroindustrial",
    "Ficha técnica",
    "Otro",
]

METODOLOGIAS_INFORME = [
    "Design Thinking",
    "Waterfall",
    "Agile",
    "DFMA",
    "UCD",
    "HCD",
    "Engineering Design Process",
    "Otra",
]


def conteo_palabras(texto: str) -> int:
    return len(str(texto or "").split())


def obtener_evidencias_sugeridas(tipo_proyecto: str) -> list[str]:
    mapa = {
        "Software": [
            "Especificación o documento de requerimientos (SRS)",
            "Diagramas de arquitectura, procesos o modelo de datos",
            "Capturas del funcionamiento de la solución",
            "Registro de pruebas funcionales y resultados",
        ],
        "Diseño industrial": [
            "Modelos CAD o representaciones tridimensionales",
            "Planos técnicos y detalles constructivos",
            "Simulaciones o validaciones de diseño",
            "Fotografías del prototipo desarrollado",
        ],
        "Electrónica y automatización": [
            "Diagramas eléctricos, electrónicos o de control",
            "Evidencias de montaje e integración del sistema",
            "Código fuente, configuración o lógica de control",
            "Resultados de pruebas de funcionamiento",
        ],
        "Inteligencia artificial": [
            "Flujo metodológico y arquitectura de la solución",
            "Resultados de entrenamiento, inferencia o procesamiento",
            "Métricas o criterios de evaluación aplicados",
            "Registro de validación con casos de prueba",
        ],
        "Desarrollo de marca e identidad visual": [
            "Logotipo y variantes de aplicación",
            "Manual de identidad visual o lineamientos de uso",
            "Aplicaciones gráficas desarrolladas",
            "Mockups o visualizaciones de implementación",
        ],
        "Prototipo agroindustrial": [
            "Descripción del proceso o formulación desarrollada",
            "Evidencia del prototipo o producto resultante",
            "Registro fotográfico del desarrollo",
            "Pruebas o validaciones técnicas realizadas",
        ],
        "Ficha técnica": [
            "Documento de especificaciones técnicas",
            "Imágenes o esquemas del producto",
            "Criterios de uso, operación o aplicación",
            "Registro de validación de la información técnica",
        ],
    }
    return mapa.get(tipo_proyecto, [
        "Descripción técnica del resultado entregado",
        "Registro documental o fotográfico del desarrollo",
        "Pruebas o validaciones reportadas por el usuario",
        "Archivos finales o soportes de entrega",
    ])


def contexto_informe_para_ia(datos: dict) -> str:
    aceptacion = (
        "El usuario confirmó que el producto final fue presentado, validado y aceptado por el talento beneficiario."
        if datos.get("aceptacion_confirmada")
        else "El usuario NO confirmó aceptación final. No afirmes que el producto fue aceptado; indica que la aceptación debe formalizarse cuando corresponda."
    )
    return f"""
Nombre del proyecto: {datos.get('nombre_proyecto', '')}
Código del proyecto: {datos.get('codigo_proyecto', '')}
Talento o beneficiario: {datos.get('nombre_talento', '')}
Experto o asesor Tecnoparque: {datos.get('nombre_experto', '')}
Línea tecnológica: {datos.get('linea_tecnologica', '')}
Tipo de proyecto: {datos.get('tipo_proyecto_detallado', '')}
Contexto, necesidad o problema: {datos.get('contexto_proyecto', '')}
Metodología: {datos.get('metodologia_detallada', '')}
Aplicación de la metodología: {datos.get('aplicacion_metodologia', '')}
Tipo de innovación y valor diferencial: {datos.get('innovacion_valor', '')}
Producto final entregado: {datos.get('producto_final', '')}
Costo total estimado: {formato_moneda_colombiana(datos.get('costo_total', 0))}
Pruebas y validación reportadas: {datos.get('pruebas_validacion', '')}
Evidencias disponibles: {datos.get('evidencias_disponibles', '')}
Entregables reportados: {datos.get('entregables', '')}
Condición sobre aceptación: {aceptacion}
""".strip()


def generar_apartado_informe_modo_prueba(titulo: str, datos: dict, minimo_palabras: int) -> str:
    base = {
        "Introducción, Contexto y Antecedentes": (
            f"El presente Informe Técnico Final documenta el proceso desarrollado para el proyecto {datos.get('nombre_proyecto', '')}, "
            f"identificado con el código {datos.get('codigo_proyecto', '')}, en el marco del acompañamiento técnico de Tecnoparque Nodo Angostura. "
            f"La iniciativa se relaciona con la necesidad descrita por el talento beneficiario: {datos.get('contexto_proyecto', '')}. "
            "El informe organiza los antecedentes, el enfoque metodológico aplicado, la estimación técnica de recursos, las validaciones reportadas "
            "y los productos entregados, con el propósito de dejar trazabilidad del desarrollo sin convertir el documento en una certificación contable, "
            "normativa o de aceptación no acreditada. "
        ),
        "Metodología": (
            f"El proyecto se abordó mediante la metodología {datos.get('metodologia_detallada', '')}. "
            f"De acuerdo con la información suministrada, su aplicación se desarrolló así: {datos.get('aplicacion_metodologia', '')}. "
            "La metodología se interpreta como una ruta de trabajo para reconocer la necesidad, definir criterios técnicos, construir alternativas, "
            "producir resultados verificables y revisar oportunidades de ajuste. Cada fase debe estar respaldada por archivos, decisiones, evidencias "
            "o validaciones efectivamente disponibles en el expediente del proyecto. "
        ),
        "Normatividad": (
            f"Para el tipo de proyecto {datos.get('tipo_proyecto_detallado', '')}, la revisión normativa debe orientarse a referentes técnicos y buenas prácticas "
            "realmente aplicables al producto desarrollado. Este documento no acredita certificaciones, conformidades ni cumplimiento de normas específicas. "
            "La identificación final de requisitos dependerá del uso previsto, condiciones de seguridad, datos tratados, materiales utilizados, entorno de operación "
            "y evidencias de prueba disponibles. Se recomienda conservar trazabilidad de versiones, decisiones técnicas, resultados de validación y soportes de entrega. "
        ),
        "Análisis y Estimación de Costos del Proyecto": (
            f"El costo total suministrado para el desarrollo corresponde a {formato_moneda_colombiana(datos.get('costo_total', 0))}. "
            "La tabla incluida en este informe distribuye dicho total de forma estimativa entre componentes coherentes con el tipo de proyecto y el producto final. "
            "Su propósito es facilitar una lectura técnica del esfuerzo y de los recursos asociados, sin representar facturación, avalúo, certificación financiera "
            "o verificación contable. Los valores pueden corresponder a actividades de diseño, desarrollo, integración, documentación, pruebas o recursos técnicos. "
        ),
        "Pruebas Documentadas y Validación del Prototipo": (
            f"Las pruebas y validaciones informadas para el proyecto son las siguientes: {datos.get('pruebas_validacion', '')}. "
            f"Como soportes disponibles se reportan: {datos.get('evidencias_disponibles', '')}. "
            "La validación técnica debe entenderse en función de los registros efectivamente anexados, de las condiciones bajo las cuales se realizaron las pruebas "
            "y de los criterios definidos para comprobar funcionamiento, calidad, presentación o desempeño del producto final. "
        ),
        "Entregables": (
            f"Los entregables declarados por el usuario son: {datos.get('entregables', '')}. "
            "Estos productos representan los resultados que deben quedar identificados, organizados y vinculados con su respectiva evidencia documental, digital, "
            "gráfica o física. La tabla de entregables se construye a partir de lo reportado y permite revisar el estado de cada resultado sin inventar archivos no aportados. "
        ),
        "Análisis y Conclusiones": (
            f"El proyecto {datos.get('nombre_proyecto', '')} consolida un resultado técnico asociado a {datos.get('producto_final', '')}. "
            f"Su valor diferencial se relaciona con: {datos.get('innovacion_valor', '')}. "
            "La revisión final integra contexto, metodología, costos estimativos, pruebas reportadas y entregables identificados. "
        ),
    }.get(titulo, "El apartado se construye con base en la información suministrada para el proyecto. ")
    if titulo == "Análisis y Conclusiones":
        if datos.get("aceptacion_confirmada"):
            base += "De acuerdo con la confirmación diligenciada, se deja constancia de que el producto final fue presentado, validado y aceptado por el talento beneficiario. "
        else:
            base += "No se deja constancia de aceptación final, debido a que dicha confirmación no fue marcada en el formulario; esta deberá formalizarse mediante el soporte correspondiente. "
    ampliaciones = [
        "La documentación técnica cumple un papel central porque permite relacionar la necesidad inicial con los resultados alcanzados, los criterios de revisión y los soportes disponibles. En ese sentido, cada afirmación debe corresponder a evidencias identificables y conservar una redacción prudente sobre los alcances reales del desarrollo.",
        "El acompañamiento de Tecnoparque se comprende como un proceso de orientación técnica y metodológica que fortalece la estructuración del proyecto. Las decisiones de diseño, desarrollo o validación deben quedar asociadas a las condiciones reportadas por el talento y a los recursos efectivamente utilizados durante la ejecución.",
        "La revisión del producto final debe considerar su propósito, las funcionalidades o características desarrolladas y la forma en que responde a la necesidad planteada. Cuando exista información pendiente de verificación, esta debe presentarse como una oportunidad de complemento documental y no como un resultado ya certificado.",
        "Para efectos de trazabilidad, es recomendable conservar archivos fuente, versiones finales, registros fotográficos, resultados de prueba, documentos de diseño, actas y demás soportes que permitan reconstruir el proceso de desarrollo. Esta organización facilita evaluaciones posteriores y procesos de mejora o continuidad.",
        "La estimación de recursos y actividades técnicas permite dimensionar el esfuerzo asociado al proyecto sin sustituir procesos contables o contractuales. Su lectura debe enfocarse en los componentes del desarrollo, la integración del resultado y los recursos vinculados a documentación, revisión, pruebas o presentación final.",
        "La validación debe estar respaldada por criterios explícitos, observaciones registradas y soportes disponibles. En proyectos tecnológicos, esta práctica ayuda a identificar ajustes, documentar resultados funcionales y sustentar decisiones sobre evolución, mantenimiento, escalamiento o nuevas etapas de desarrollo.",
        "El valor diferencial del proyecto debe analizarse a partir del problema atendido, la pertinencia de la solución y las características que aportan una mejora frente al contexto inicial. Este análisis evita afirmaciones generales y orienta la comunicación del resultado hacia elementos técnicos verificables.",
        "La presentación final de entregables constituye un cierre documental del proceso y permite establecer cuáles productos fueron reportados, qué soportes los acompañan y qué elementos requieren conservación o actualización. Esta información es relevante para la memoria técnica del proyecto y su eventual continuidad.",
        "Las buenas prácticas aplicables pueden relacionarse con gestión documental, control de versiones, seguridad, calidad, pruebas, usabilidad, diseño o trazabilidad, según corresponda al tipo de desarrollo. La adopción de referentes concretos debe verificarse posteriormente frente al producto, sector y uso previsto.",
        "Finalmente, el informe debe facilitar la comprensión del proyecto por parte de sus interesados, integrando información técnica clara, verificable y ordenada. Su contenido constituye una síntesis del desarrollo reportado y una base para determinar acciones futuras, sin reemplazar certificaciones o evaluaciones externas que no hayan sido aportadas.",
    ]
    texto = base
    indice = 0
    while conteo_palabras(texto) < minimo_palabras:
        texto += " " + ampliaciones[indice % len(ampliaciones)]
        indice += 1
    return texto.strip()


def generar_apartado_informe_con_chatgpt(
    titulo: str,
    datos: dict,
    minimo_palabras: int,
    modelo: str = "gpt-4.1-mini",
) -> str:
    if OpenAI is None:
        raise ImportError("No está instalada la librería openai. Instálala con: pip install openai")
    api_key = obtener_api_key()
    if not api_key:
        raise ValueError("No se encontró OPENAI_API_KEY.")
    client = OpenAI(api_key=api_key)
    instrucciones = f"""
Eres un redactor técnico institucional para proyectos de SENA Tecnoparque Nodo Angostura.
Redacta únicamente el apartado titulado: {titulo}.
El texto debe ser profesional, técnico, coherente y tener mínimo {minimo_palabras} palabras.
No uses markdown, listas con viñetas ni encabezados adicionales.
No inventes datos, pruebas, certificaciones, normas específicas, aceptación, cifras ni documentos no suministrados.
En normatividad, menciona referentes y buenas prácticas pertinentes de manera prudente, aclarando que la aplicabilidad debe verificarse según el producto y uso previsto.
En costos, aclara que se trata de una estimación técnica y no de información contable certificada.
En conclusiones, solo afirma aceptación final cuando la condición suministrada indique que fue confirmada.
"""
    prompt = contexto_informe_para_ia(datos)
    respuesta = client.responses.create(
        model=modelo,
        instructions=instrucciones,
        input=prompt,
        temperature=0.28,
    )
    texto = str(getattr(respuesta, "output_text", "") or "").strip()
    if conteo_palabras(texto) < minimo_palabras:
        ampliacion = client.responses.create(
            model=modelo,
            instructions=instrucciones,
            input=(
                prompt
                + "\n\nBorrador inicial que debes ampliar y mejorar manteniendo los hechos suministrados:\n"
                + texto
                + f"\n\nDevuelve una versión integral de al menos {minimo_palabras} palabras."
            ),
            temperature=0.25,
        )
        texto_ampliado = str(getattr(ampliacion, "output_text", "") or "").strip()
        if conteo_palabras(texto_ampliado) > conteo_palabras(texto):
            texto = texto_ampliado
    return texto or generar_apartado_informe_modo_prueba(titulo, datos, minimo_palabras)


def categorias_costos_por_tipo(tipo_proyecto: str) -> list[tuple[str, str, float]]:
    catalogo = {
        "Software": [
            ("Análisis y arquitectura", "Levantamiento técnico, estructura de solución y diseño funcional", 0.18),
            ("Desarrollo e implementación", "Construcción de módulos, lógica y componentes de software", 0.42),
            ("Interfaz y experiencia de usuario", "Diseño de pantallas, navegación y ajustes de usabilidad", 0.14),
            ("Pruebas y validación", "Verificación funcional y ajustes derivados", 0.14),
            ("Documentación técnica", "Manuales, soportes y consolidación de entregables", 0.12),
        ],
        "Diseño industrial": [
            ("Conceptualización", "Investigación, criterios de diseño y alternativas", 0.16),
            ("Modelado CAD", "Diseño tridimensional y ajustes geométricos", 0.27),
            ("Planos y especificaciones", "Despieces, medidas y documentación técnica", 0.17),
            ("Prototipado", "Materialización o validación de la propuesta", 0.28),
            ("Pruebas y ajustes", "Verificaciones y mejoras del diseño", 0.12),
        ],
        "Electrónica y automatización": [
            ("Diseño electrónico", "Diagramas, selección de componentes y arquitectura", 0.18),
            ("Componentes e integración", "Elementos, montaje e interconexión", 0.34),
            ("Programación y control", "Firmware, lógica o configuración de automatización", 0.20),
            ("Pruebas funcionales", "Medición, puesta a punto y validación", 0.17),
            ("Documentación", "Soportes técnicos y evidencias", 0.11),
        ],
        "Inteligencia artificial": [
            ("Preparación de información", "Estructuración de insumos y datos disponibles", 0.17),
            ("Diseño de solución IA", "Arquitectura, flujo y criterios de análisis", 0.21),
            ("Desarrollo e integración", "Implementación del componente inteligente", 0.30),
            ("Evaluación y validación", "Métricas, pruebas y revisión de resultados", 0.20),
            ("Documentación técnica", "Registro metodológico y soportes finales", 0.12),
        ],
        "Desarrollo de marca e identidad visual": [
            ("Investigación y concepto", "Análisis inicial y direccionamiento creativo", 0.18),
            ("Diseño de identidad", "Logotipo, sistema visual y variantes", 0.30),
            ("Manual de marca", "Lineamientos técnicos de uso", 0.22),
            ("Aplicaciones y mockups", "Visualización de piezas y soportes", 0.20),
            ("Validación y entrega", "Ajustes y consolidación de archivos", 0.10),
        ],
        "Prototipo agroindustrial": [
            ("Formulación y diseño de proceso", "Definición técnica de la solución", 0.18),
            ("Materias primas o insumos", "Elementos estimados para prototipo", 0.27),
            ("Desarrollo del prototipo", "Preparación, fabricación o integración", 0.27),
            ("Pruebas y validación", "Evaluación técnica del resultado", 0.17),
            ("Documentación y entrega", "Ficha, registro y evidencias", 0.11),
        ],
        "Ficha técnica": [
            ("Levantamiento de información", "Obtención y depuración de datos técnicos", 0.22),
            ("Estructuración documental", "Organización de especificaciones y contenido", 0.30),
            ("Recursos gráficos", "Imágenes, esquemas o diagramación", 0.18),
            ("Revisión y validación", "Verificación de información y ajustes", 0.18),
            ("Entrega final", "Consolidación del documento final", 0.12),
        ],
    }
    return catalogo.get(tipo_proyecto, catalogo["Software"])


def ajustar_tabla_costos(tabla: list[dict], total: int) -> list[dict]:
    total = max(0, int(round(total)))
    filas = []
    for idx, item in enumerate(tabla, start=1):
        categoria = str(item.get("categoria", item.get("componente", "Componente técnico"))).strip() or "Componente técnico"
        descripcion = str(item.get("descripcion", item.get("descripcion_costo", "Estimación técnica asociada al desarrollo"))).strip()
        try:
            valor = max(0, int(round(float(item.get("valor", item.get("valor_estimado", 0))))))
        except Exception:
            valor = 0
        filas.append({"item": idx, "categoria": categoria, "descripcion": descripcion, "valor_estimado": valor})
    if not filas:
        filas = [{"item": 1, "categoria": "Desarrollo técnico", "descripcion": "Estimación global del proyecto", "valor_estimado": total}]
    suma = sum(f["valor_estimado"] for f in filas)
    if suma == 0:
        base = total // len(filas) if filas else total
        for fila in filas:
            fila["valor_estimado"] = base
        suma = sum(f["valor_estimado"] for f in filas)
    diferencia = total - suma
    filas[-1]["valor_estimado"] = max(0, filas[-1]["valor_estimado"] + diferencia)
    nueva_suma = sum(f["valor_estimado"] for f in filas)
    if nueva_suma != total:
        filas[-1]["valor_estimado"] += total - nueva_suma
    return filas


def generar_tabla_costos_modo_prueba(datos: dict) -> list[dict]:
    total = int(datos.get("costo_total", 0))
    filas = []
    for idx, (categoria, descripcion, proporcion) in enumerate(categorias_costos_por_tipo(datos.get("tipo_proyecto", "")), start=1):
        filas.append({
            "item": idx,
            "categoria": categoria,
            "descripcion": descripcion,
            "valor_estimado": int(round(total * proporcion)),
        })
    return ajustar_tabla_costos(filas, total)


def generar_tabla_costos_con_chatgpt(datos: dict, modelo: str = "gpt-4.1-mini") -> list[dict]:
    if OpenAI is None:
        raise ImportError("No está instalada la librería openai. Instálala con: pip install openai")
    api_key = obtener_api_key()
    if not api_key:
        raise ValueError("No se encontró OPENAI_API_KEY.")
    client = OpenAI(api_key=api_key)
    instrucciones = """
Genera una tabla estimativa de costos para un informe técnico final de Tecnoparque.
Responde únicamente JSON válido con la clave "costos" y entre 4 y 7 filas.
Cada fila debe contener: categoria, descripcion, valor_estimado como número entero en pesos colombianos.
La distribución debe ser coherente con el tipo de proyecto y el producto final entregado.
No presentes los valores como costos certificados, facturas o contabilidad verificada.
"""
    entrada = contexto_informe_para_ia(datos) + "\nEl total que debe distribuirse es: " + str(int(datos.get("costo_total", 0)))
    respuesta = client.responses.create(model=modelo, instructions=instrucciones, input=entrada, temperature=0.2)
    contenido = json.loads(limpiar_respuesta_json(respuesta.output_text))
    return ajustar_tabla_costos(contenido.get("costos", []), int(datos.get("costo_total", 0)))


def dividir_entregables_texto(texto: str) -> list[str]:
    texto = str(texto or "").replace(";", "\n")
    elementos = [linea.strip(" -•\t") for linea in texto.splitlines() if linea.strip(" -•\t")]
    if len(elementos) <= 1 and "," in texto:
        elementos = [item.strip() for item in texto.split(",") if item.strip()]
    return elementos or ["Entregable descrito en el informe técnico final"]


def generar_tabla_entregables(datos: dict) -> list[dict]:
    evidencias = dividir_entregables_texto(datos.get("evidencias_disponibles", ""))
    entregables = dividir_entregables_texto(datos.get("entregables", ""))
    filas = []
    for idx, entregable in enumerate(entregables, start=1):
        evidencia = evidencias[idx - 1] if idx - 1 < len(evidencias) else "Soporte por verificar o anexar"
        filas.append({
            "item": idx,
            "entregable": entregable,
            "evidencia": evidencia,
            "estado": "Reportado por el usuario",
        })
    return filas


def sombrear_celda_docx(celda, color: str) -> None:
    propiedades = celda._tc.get_or_add_tcPr()
    sombreado = OxmlElement("w:shd")
    sombreado.set(qn("w:fill"), color)
    propiedades.append(sombreado)


def configurar_celda_docx(celda, texto: str, negrita: bool = False, tamano: int = 9) -> None:
    celda.text = ""
    parrafo = celda.paragraphs[0]
    run = parrafo.add_run(str(texto))
    run.bold = negrita
    run.font.size = DocxPt(tamano)
    celda.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def agregar_titulo_docx(documento, titulo: str) -> None:
    parrafo = documento.add_paragraph()
    parrafo.paragraph_format.space_before = DocxPt(10)
    parrafo.paragraph_format.space_after = DocxPt(6)
    run = parrafo.add_run(titulo)
    run.bold = True
    run.font.size = DocxPt(13)
    run.font.color.rgb = RGBColor(57, 169, 53)


def agregar_texto_docx(documento, texto: str) -> None:
    for bloque in [p.strip() for p in str(texto or "").split("\n") if p.strip()]:
        parrafo = documento.add_paragraph(bloque)
        parrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        parrafo.paragraph_format.space_after = DocxPt(6)
        parrafo.paragraph_format.line_spacing = 1.15
        for run in parrafo.runs:
            run.font.size = DocxPt(10)


def generar_docx_informe_tecnico_final(datos: dict) -> str:
    if Document is None:
        raise ImportError("No está instalada python-docx. Agrégala a requirements.txt e instala con: pip install python-docx")
    Path(CARPETA_SALIDA).mkdir(parents=True, exist_ok=True)
    archivo = f"Informe_Tecnico_Final_{safe_filename(datos.get('codigo_proyecto', 'proyecto'))}.docx"
    ruta_docx = str(Path(CARPETA_SALIDA) / archivo)
    documento = Document()
    seccion = documento.sections[0]
    seccion.top_margin = DocxCm(2.3)
    seccion.bottom_margin = DocxCm(2.0)
    seccion.left_margin = DocxCm(2.2)
    seccion.right_margin = DocxCm(2.2)

    estilo_normal = documento.styles["Normal"]
    estilo_normal.font.name = "Arial"
    estilo_normal.font.size = DocxPt(10)

    encabezado = seccion.header
    tabla_header = encabezado.add_table(rows=2, cols=3, width=DocxCm(17))
    tabla_header.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabla_header.columns[0].width = DocxCm(4.0)
    tabla_header.columns[1].width = DocxCm(9.0)
    tabla_header.columns[2].width = DocxCm(4.0)
    ruta_tecno = obtener_ruta_logo_tecnoparque()
    ruta_sena = obtener_ruta_logo_sena()
    if ruta_tecno and Path(ruta_tecno).exists():
        tabla_header.cell(0, 0).paragraphs[0].add_run().add_picture(ruta_tecno, width=DocxCm(3.7))
    else:
        configurar_celda_docx(tabla_header.cell(0, 0), "TECNOPARQUE", True, 9)
    titulo_p = tabla_header.cell(0, 1).paragraphs[0]
    titulo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo_r = titulo_p.add_run("INFORME TÉCNICO FINAL")
    titulo_r.bold = True
    titulo_r.font.size = DocxPt(16)
    titulo_r.font.color.rgb = RGBColor(57, 169, 53)
    if ruta_sena and Path(ruta_sena).exists():
        p_logo = tabla_header.cell(0, 2).paragraphs[0]
        p_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_logo.add_run().add_picture(ruta_sena, width=DocxCm(2.0))
    else:
        configurar_celda_docx(tabla_header.cell(0, 2), "SENA", True, 12)
    subtitulo = tabla_header.cell(1, 0).merge(tabla_header.cell(1, 2)).paragraphs[0]
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = subtitulo.add_run(f"{datos.get('codigo_proyecto', '')} - {datos.get('nombre_proyecto', '')}")
    run_sub.bold = True
    run_sub.font.size = DocxPt(10)

    p_inicial = documento.add_paragraph()
    p_inicial.add_run(f"Talento o beneficiario: ").bold = True
    p_inicial.add_run(datos.get("nombre_talento", ""))
    p_inicial.add_run("   |   Experto Tecnoparque: ").bold = True
    p_inicial.add_run(datos.get("nombre_experto", ""))
    p_linea = documento.add_paragraph()
    p_linea.add_run("Línea tecnológica: ").bold = True
    p_linea.add_run(datos.get("linea_tecnologica", ""))
    p_linea.add_run("   |   Tipo de proyecto: ").bold = True
    p_linea.add_run(datos.get("tipo_proyecto_detallado", ""))

    secciones = [
        "Introducción, Contexto y Antecedentes",
        "Metodología",
        "Normatividad",
        "Análisis y Estimación de Costos del Proyecto",
        "Pruebas Documentadas y Validación del Prototipo",
        "Entregables",
        "Análisis y Conclusiones",
    ]
    contenidos = datos.get("contenido_generado", {})
    for titulo in secciones:
        agregar_titulo_docx(documento, titulo)
        agregar_texto_docx(documento, contenidos.get(titulo, ""))
        if titulo == "Análisis y Estimación de Costos del Proyecto":
            p_nota = documento.add_paragraph()
            r_nota = p_nota.add_run("Nota: La siguiente distribución corresponde a una estimación técnica y no constituye información contable certificada.")
            r_nota.italic = True
            r_nota.font.size = DocxPt(9)
            tabla = documento.add_table(rows=1, cols=4)
            tabla.style = "Table Grid"
            tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
            encabezados = ["Ítem", "Categoría o componente", "Descripción del costo estimado", "Valor estimado"]
            for col, encabezado_texto in enumerate(encabezados):
                configurar_celda_docx(tabla.cell(0, col), encabezado_texto, True, 9)
                sombrear_celda_docx(tabla.cell(0, col), "D9EAD3")
            for fila in datos.get("tabla_costos", []):
                celdas = tabla.add_row().cells
                configurar_celda_docx(celdas[0], fila.get("item", ""), False, 9)
                configurar_celda_docx(celdas[1], fila.get("categoria", ""), False, 9)
                configurar_celda_docx(celdas[2], fila.get("descripcion", ""), False, 9)
                configurar_celda_docx(celdas[3], formato_moneda_colombiana(fila.get("valor_estimado", 0)), False, 9)
            total_cells = tabla.add_row().cells
            total_cells[0].merge(total_cells[2])
            configurar_celda_docx(total_cells[0], "TOTAL ESTIMADO DEL PROTOTIPO O DESARROLLO", True, 9)
            configurar_celda_docx(total_cells[3], formato_moneda_colombiana(datos.get("costo_total", 0)), True, 9)
            sombrear_celda_docx(total_cells[0], "D9EAD3")
            sombrear_celda_docx(total_cells[3], "D9EAD3")
        if titulo == "Pruebas Documentadas y Validación del Prototipo":
            documento.add_paragraph("Espacios sugeridos para anexar evidencias:")
            tabla_ev = documento.add_table(rows=1, cols=3)
            tabla_ev.style = "Table Grid"
            for col, head in enumerate(["Evidencia sugerida", "Referencia / archivo anexo", "Observación"]):
                configurar_celda_docx(tabla_ev.cell(0, col), head, True, 9)
                sombrear_celda_docx(tabla_ev.cell(0, col), "D9EAD3")
            for evidencia in obtener_evidencias_sugeridas(datos.get("tipo_proyecto", "")):
                row = tabla_ev.add_row().cells
                configurar_celda_docx(row[0], evidencia, False, 9)
                configurar_celda_docx(row[1], "", False, 9)
                configurar_celda_docx(row[2], "", False, 9)
        if titulo == "Entregables":
            tabla_en = documento.add_table(rows=1, cols=4)
            tabla_en.style = "Table Grid"
            for col, head in enumerate(["Ítem", "Entregable", "Evidencia asociada", "Estado"]):
                configurar_celda_docx(tabla_en.cell(0, col), head, True, 9)
                sombrear_celda_docx(tabla_en.cell(0, col), "D9EAD3")
            for fila in datos.get("tabla_entregables", []):
                row = tabla_en.add_row().cells
                configurar_celda_docx(row[0], fila.get("item", ""), False, 9)
                configurar_celda_docx(row[1], fila.get("entregable", ""), False, 9)
                configurar_celda_docx(row[2], fila.get("evidencia", ""), False, 9)
                configurar_celda_docx(row[3], fila.get("estado", ""), False, 9)

    documento.save(ruta_docx)
    datos_json = dict(datos)
    datos_json["ruta_docx"] = ruta_docx
    guardar_datos_json(datos_json, ruta="datos_informe_tecnico_final.json")
    return ruta_docx


# =====================================================
# =====================================================
# MODELO DE NEGOCIOS - FASE DE CIERRE / PDF
# Basado en formato de identificación de modelo de negocios TRL 6 / TRL 7
# =====================================================

ITEMS_LEAN_CANVAS = [
    "Propuesta de valor",
    "Segmento de clientes y estrategia de adopción",
    "Canales de distribución",
    "Relaciones con clientes",
    "Flujo de Ingresos",
    "Recursos Claves",
    "Actividades Claves",
    "Alianzas claves",
    "Estructura de Costos",
]

PREGUNTAS_MODELO_NEGOCIO = {
    "Propuesta de valor": "¿Qué beneficios se ofrece a los clientes? Precio, calidad, novedad, marca, diseño.",
    "Segmento de clientes y estrategia de adopción": "¿Quiénes son los clientes más importantes? ¿Quiénes pagarán por el producto o servicio? ¿Mercado de masas, nichos o mercado diversificado?",
    "Canales de distribución": "¿Cómo pueden comprar, acceder o implementar los productos o servicios los clientes? Canales directos o indirectos.",
    "Relaciones con clientes": "¿Cómo obtener, retener y aumentar los clientes? Relación personal, servicios automatizados, soporte o acompañamiento.",
    "Flujo de Ingresos": "¿De dónde se generan los ingresos del negocio? Venta, uso, suscripción, publicidad, licencia o servicios asociados.",
    "Recursos Claves": "¿Cuáles son los recursos más importantes del negocio? Físicos, intelectuales y humanos.",
    "Actividades Claves": "¿Cuáles son las actividades más importantes para que el modelo funcione? Producción, comercialización, plataforma, logística o gestión humana.",
    "Alianzas claves": "¿Cuáles son los socios más importantes para que la idea funcione? Economía a escala, compra de recursos, actividades o proveedores.",
    "Estructura de Costos": "¿Cuáles son los recursos y actividades claves más costosas? Costos fijos, variables, gastos e infraestructura.",
}


def limpiar_texto_canvas(texto: str) -> str:
    """Normaliza texto para PDF conservando párrafos simples y eliminando puntos suspensivos."""
    texto = str(texto or "").strip()
    texto = texto.replace("\\n", "\n")
    texto = texto.replace("…", ".").replace("...", ".")
    texto = texto.replace("- ", "• ").replace("* ", "• ")
    while "\n\n\n" in texto:
        texto = texto.replace("\n\n\n", "\n\n")
    return texto.strip()


def interpretar_tipo_producto(nombre_proyecto: str, descripcion_producto: str, aspectos_generacion: str, region_contexto: str) -> str:
    texto_total = f"{nombre_proyecto} {descripcion_producto} {aspectos_generacion} {region_contexto}".lower()
    categorias = []
    if any(p in texto_total for p in ["software", "app", "plataforma", "sistema", "web", "aplicación", "aplicacion", "dashboard", "base de datos", "módulo", "modulo"]):
        categorias.append("sistema digital")
    if any(p in texto_total for p in ["sensor", "electrónica", "electronica", "automatización", "automatizacion", "iot", "microcontrolador", "dispositivo", "control"]):
        categorias.append("solución electrónica o automatizada")
    if any(p in texto_total for p in ["marca", "logotipo", "identidad", "branding", "manual de marca", "diseño gráfico"]):
        categorias.append("marca e identidad visual")
    if any(p in texto_total for p in ["agro", "alimento", "café", "cafe", "proceso", "producto agroindustrial", "cultivo", "rural"]):
        categorias.append("solución agroindustrial")
    if any(p in texto_total for p in ["diseño", "diseño industrial", "cad", "prototipo", "producto físico", "producto fisico", "mecanismo"]):
        categorias.append("prototipo de producto")
    if any(p in texto_total for p in ["ia", "inteligencia artificial", "modelo", "clasificación", "clasificacion", "predicción", "prediccion", "machine learning"]):
        categorias.append("solución con inteligencia artificial")

    if categorias:
        return ", ".join(dict.fromkeys(categorias))
    return "producto o servicio de base tecnológica"


def contexto_lean_para_ia(
    nombre_proyecto: str,
    codigo_proyecto: str,
    descripcion_producto: str,
    aspectos_generacion: str,
    region_contexto: str,
    usuarios_objetivo: str = "",
    trl_nivel: str = "TRL 6",
) -> str:
    tipo = interpretar_tipo_producto(nombre_proyecto, descripcion_producto, aspectos_generacion, region_contexto)
    usuarios = usuarios_objetivo.strip() or "No especificado; inferir con prudencia a partir del producto, contexto y usuarios probables."
    return f"""
Nombre del proyecto: {nombre_proyecto}
Código del proyecto: {codigo_proyecto}
Nivel TRL del informe: {trl_nivel}
Tipo inferido de solución: {tipo}
Descripción del prototipo, producto o servicio suministrada por el usuario:
{descripcion_producto}

Región, territorio o contexto de implementación suministrado por el usuario:
{region_contexto}

Usuarios, clientes o beneficiarios objetivo indicados por el usuario:
{usuarios}

Aspectos estratégicos a tener en cuenta suministrados por el usuario:
{aspectos_generacion}
""".strip()


def remover_nombre_proyecto_en_bloque(texto: str, nombre_proyecto: str, codigo_proyecto: str = "") -> str:
    texto = limpiar_texto_canvas(texto)
    for valor in [nombre_proyecto, codigo_proyecto]:
        valor = str(valor or "").strip()
        if valor:
            texto = texto.replace(valor, "la solución")
            texto = texto.replace(valor.upper(), "la solución")
            texto = texto.replace(valor.lower(), "la solución")
    return texto


def limitar_caracteres_canvas(texto: str, nombre_proyecto: str = "", codigo_proyecto: str = "", min_chars: int = 150, max_chars: int = 250) -> str:
    """Ajusta el texto entre 150 y 250 caracteres, sin puntos suspensivos ni cortes con '...'."""
    texto = remover_nombre_proyecto_en_bloque(texto, nombre_proyecto, codigo_proyecto)
    texto = " ".join(texto.replace("\n", " ").split())
    texto = texto.replace("…", ".").replace("...", ".")

    frases_no_deseadas = [
        "se debe diligenciar", "debe contener", "este bloque debe", "se recomienda poner",
        "debería incluir", "como llenar", "cómo llenar", "en este campo", "debe responder",
        "debe analizar", "el usuario indicó", "el usuario señala", "según la descripción",
        "de acuerdo con la descripción", "este apartado", "este bloque"
    ]
    texto_lower = texto.lower()
    for frase in frases_no_deseadas:
        if frase in texto_lower:
            texto = texto.replace(frase, "")

    if len(texto) > max_chars:
        corte = texto[:max_chars]
        ultimo_punto = max(corte.rfind("."), corte.rfind(";"))
        ultimo_espacio = corte.rfind(" ")
        if ultimo_punto >= min_chars:
            texto = corte[:ultimo_punto + 1]
        elif ultimo_espacio >= min_chars:
            texto = corte[:ultimo_espacio].rstrip(" ,;:") + "."
        else:
            texto = corte.rstrip(" ,;:") + "."

    while len(texto) < min_chars:
        texto += " Incluye soporte, validación y mejora continua para facilitar su adopción."
        if len(texto) > max_chars:
            texto = texto[:max_chars].rsplit(" ", 1)[0].rstrip(" ,;:") + "."
            break

    return texto.strip().replace("...", ".").replace("…", ".")


def generar_modelo_negocio_modo_prueba(
    nombre_proyecto: str,
    codigo_proyecto: str,
    descripcion_producto: str,
    aspectos_generacion: str,
    region_contexto: str,
    usuarios_objetivo: str = "",
    trl_nivel: str = "TRL 6",
) -> dict:
    """Fallback local: genera respuestas concretas, cortas y aplicadas, no instrucciones."""
    tipo = interpretar_tipo_producto(nombre_proyecto, descripcion_producto, aspectos_generacion, region_contexto)
    region = region_contexto.strip() or "el contexto de implementación definido"
    usuarios = usuarios_objetivo.strip() or "usuarios y beneficiarios con necesidad directa"

    contenido = {
        "Propuesta de valor": f"Ofrece una {tipo} práctica, ajustada a {region}, que mejora tiempos, orden, trazabilidad y toma de decisiones. Entrega valor por su utilidad funcional, adaptación local y capacidad de validar resultados.",
        "Segmento de clientes y estrategia de adopción": f"Los clientes iniciales son {usuarios}. La adopción se impulsa con pilotos, demostraciones, capacitación breve y evidencia de beneficios operativos antes de escalar a otros usuarios similares.",
        "Canales de distribución": f"La solución se entrega mediante contacto directo, demostraciones, soporte digital y aliados del territorio. Los canales priorizan explicación técnica, prueba funcional y acompañamiento durante la implementación.",
        "Relaciones con clientes": "La relación se basa en atención personalizada, soporte técnico continuo, canales de comunicación para incidentes y seguimiento a necesidades específicas. La solución se presenta como herramienta funcional en mejora constante.",
        "Flujo de Ingresos": "Los ingresos pueden generarse por implementación, licenciamiento, suscripción, soporte, mantenimiento, capacitación o personalización. También crea valor por ahorro, eficiencia y reducción de reprocesos.",
        "Recursos Claves": f"Requiere talento técnico, documentación, prototipo funcional, herramientas de desarrollo, infraestructura y usuarios validadores. El conocimiento del contexto regional fortalece su ajuste y adopción.",
        "Actividades Claves": "Las actividades centrales son desarrollo, pruebas, documentación, capacitación, soporte y ajustes funcionales. También incluye demostraciones y validación con usuarios reales para confirmar utilidad.",
        "Alianzas claves": f"Las alianzas estratégicas incluyen actores de {region}, usuarios piloto, proveedores técnicos, instituciones de apoyo y aliados de implementación que faciliten validación, acceso y adopción.",
        "Estructura de Costos": "Los costos se concentran en desarrollo, diseño, pruebas, infraestructura, soporte, documentación, capacitación y mantenimiento. También pueden incluir adaptación local y acompañamiento inicial.",
    }

    return {
        clave: limitar_caracteres_canvas(valor, nombre_proyecto, codigo_proyecto, 150, 250)
        for clave, valor in contenido.items()
    }


# Alias para compatibilidad con llamadas anteriores del módulo.
def generar_lean_canvas_modo_prueba(
    nombre_proyecto: str,
    codigo_proyecto: str,
    descripcion_producto: str,
    aspectos_generacion: str,
    region_contexto: str = "",
    usuarios_objetivo: str = "",
    trl_nivel: str = "TRL 6",
) -> dict:
    return generar_modelo_negocio_modo_prueba(
        nombre_proyecto,
        codigo_proyecto,
        descripcion_producto,
        aspectos_generacion,
        region_contexto,
        usuarios_objetivo,
        trl_nivel,
    )


def generar_modelo_negocio_con_chatgpt(
    nombre_proyecto: str,
    codigo_proyecto: str,
    descripcion_producto: str,
    aspectos_generacion: str,
    region_contexto: str,
    usuarios_objetivo: str = "",
    trl_nivel: str = "TRL 6",
    modelo: str = "gpt-4.1-mini",
) -> dict:
    if OpenAI is None:
        raise ImportError("No está instalada la librería openai. Instálala con: pip install openai")

    api_key = obtener_api_key()
    if not api_key:
        raise ValueError("No se encontró OPENAI_API_KEY.")

    client = OpenAI(api_key=api_key)

    instrucciones = """
Eres un consultor senior en modelos de negocio, innovación, emprendimiento tecnológico y proyectos de base tecnológica del SENA Tecnoparque Nodo Angostura.

Construye un informe de identificación del modelo de negocios aplicado específicamente al producto descrito por el usuario. El nivel TRL será TRL 6 o TRL 7 según el dato suministrado.

Reglas obligatorias:
1. Usa el nombre del proyecto solo como contexto interno. No lo incluyas dentro de las respuestas de los 9 campos.
2. No transcribas literalmente ninguna frase del formulario. Interpreta, corrige ortografía y gramática, amplía el concepto y redacta una respuesta propia.
3. Cada respuesta debe resolver directamente la pregunta del campo, no explicar qué debería escribirse allí.
4. No uses frases instructivas como: “se debe diligenciar”, “este bloque debe”, “se recomienda”, “debe contener”, “según la descripción” o “el usuario indica”.
5. Cada respuesta debe estar aplicada al producto, a sus características, al contexto regional de implementación, a los usuarios, soporte, adopción, ingresos, costos, recursos o alianzas según corresponda.
6. No inventes ventas, clientes reales, cifras certificadas, alianzas confirmadas, certificaciones ni validaciones no suministradas.
7. Puedes proponer hipótesis de mercado, adopción, ingresos, costos y alianzas, siempre redactadas como planteamientos iniciales aplicados al producto.
8. Cada campo debe tener entre 150 y 250 caracteres, no palabras.
9. No uses puntos suspensivos.
10. Responde únicamente JSON válido, sin markdown, sin explicación y sin texto adicional.
"""

    entrada = f"""
Contexto del proyecto:
{contexto_lean_para_ia(nombre_proyecto, codigo_proyecto, descripcion_producto, aspectos_generacion, region_contexto, usuarios_objetivo, trl_nivel)}

Genera exactamente estos 9 campos, resolviendo cada pregunta del formato institucional de modelo de negocios. Recuerda: no incluyas el nombre del proyecto dentro de los campos y no copies literal lo escrito por el usuario.

Formato JSON obligatorio:
{{
  "Propuesta de valor": "...",
  "Segmento de clientes y estrategia de adopción": "...",
  "Canales de distribución": "...",
  "Relaciones con clientes": "...",
  "Flujo de Ingresos": "...",
  "Recursos Claves": "...",
  "Actividades Claves": "...",
  "Alianzas claves": "...",
  "Estructura de Costos": "..."
}}
"""

    try:
        respuesta = client.responses.create(
            model=modelo,
            instructions=instrucciones,
            input=entrada,
            temperature=0.22,
        )

        datos = json.loads(limpiar_respuesta_json(respuesta.output_text))

        respaldo = generar_modelo_negocio_modo_prueba(
            nombre_proyecto,
            codigo_proyecto,
            descripcion_producto,
            aspectos_generacion,
            region_contexto,
            usuarios_objetivo,
            trl_nivel,
        )

        frases_no_deseadas = [
            "se debe diligenciar", "debe contener", "este bloque debe", "se recomienda",
            "debería incluir", "como llenar", "cómo llenar", "en este campo",
            "debe responder", "debe analizar", "el usuario indicó", "el usuario señala",
            "según la descripción", "de acuerdo con la descripción", "este apartado", "este bloque",
        ]

        for item in ITEMS_LEAN_CANVAS:
            texto = datos.get(item, "") if isinstance(datos, dict) else ""
            if not isinstance(texto, str) or not texto.strip():
                texto = respaldo[item]

            texto_limpio = remover_nombre_proyecto_en_bloque(texto, nombre_proyecto, codigo_proyecto)
            texto_lower = texto_limpio.lower()

            if any(frase in texto_lower for frase in frases_no_deseadas):
                texto_limpio = respaldo[item]

            datos[item] = limitar_caracteres_canvas(texto_limpio, nombre_proyecto, codigo_proyecto, 150, 250)

        return {item: datos[item] for item in ITEMS_LEAN_CANVAS}

    except Exception:
        return generar_modelo_negocio_modo_prueba(
            nombre_proyecto,
            codigo_proyecto,
            descripcion_producto,
            aspectos_generacion,
            region_contexto,
            usuarios_objetivo,
            trl_nivel,
        )


# Alias para compatibilidad con llamadas anteriores del módulo.
def generar_lean_canvas_con_chatgpt(
    nombre_proyecto: str,
    codigo_proyecto: str,
    descripcion_producto: str,
    aspectos_generacion: str,
    modelo: str = "gpt-4.1-mini",
    region_contexto: str = "",
    usuarios_objetivo: str = "",
    trl_nivel: str = "TRL 6",
) -> dict:
    return generar_modelo_negocio_con_chatgpt(
        nombre_proyecto,
        codigo_proyecto,
        descripcion_producto,
        aspectos_generacion,
        region_contexto,
        usuarios_objetivo,
        trl_nivel,
        modelo,
    )


def parrafo_canvas_pdf(texto: str, estilo: ParagraphStyle) -> Paragraph:
    from html import escape
    texto = limpiar_texto_canvas(texto)
    lineas = []
    for linea in texto.splitlines():
        l = linea.strip()
        if not l:
            lineas.append("<br/>")
        elif l.startswith("•"):
            lineas.append("• " + escape(l.lstrip("• ").strip()))
        else:
            lineas.append(escape(l))
    return Paragraph("<br/>".join(lineas), estilo)


def generar_pdf_lean_canvas(datos: dict) -> str:
    if canvas is None:
        raise ImportError("No está instalada reportlab. Instálala con: pip install reportlab")

    Path(CARPETA_SALIDA).mkdir(parents=True, exist_ok=True)

    archivo = f"Modelo_Negocios_{safe_filename(datos.get('codigo_proyecto', 'proyecto'))}.pdf"
    ruta_pdf = str(Path(CARPETA_SALIDA) / archivo)

    page_size = landscape(letter)
    page_width, page_height = page_size

    verde_claro = colors.HexColor("#EAF5EA")
    gris = colors.HexColor("#F4F4F4")
    amarillo = colors.HexColor("#FFF2CC")
    azul = colors.HexColor("#D9EAF7")
    rojo = colors.HexColor("#F4CCCC")

    doc = SimpleDocTemplate(
        ruta_pdf,
        pagesize=page_size,
        rightMargin=0.75 * cm,
        leftMargin=0.75 * cm,
        topMargin=1.0 * cm,
        bottomMargin=0.75 * cm,
    )

    estilo_titulo = ParagraphStyle(
        name="TituloModeloNegocios",
        fontName="Helvetica-Bold",
        fontSize=13,
        leading=15,
        alignment=TA_CENTER,
        textColor=colors.black,
    )
    estilo_info = ParagraphStyle(
        name="InfoModeloNegocios",
        fontName="Helvetica",
        fontSize=7.4,
        leading=9,
        alignment=TA_LEFT,
    )
    estilo_info_bold = ParagraphStyle(
        name="InfoModeloNegociosBold",
        fontName="Helvetica-Bold",
        fontSize=7.4,
        leading=9,
        alignment=TA_LEFT,
    )
    estilo_pregunta = ParagraphStyle(
        name="PreguntaModeloNegocios",
        fontName="Helvetica-Bold",
        fontSize=8,
        leading=9.8,
        alignment=TA_LEFT,
        textColor=colors.black,
        spaceAfter=3,
    )
    estilo_respuesta = ParagraphStyle(
        name="RespuestaModeloNegocios",
        fontName="Helvetica",
        fontSize=7.4,
        leading=9.2,
        alignment=TA_JUSTIFY,
        spaceAfter=4,
    )
    estilo_canvas = ParagraphStyle(
        name="CanvasModeloNegocios",
        fontName="Helvetica",
        fontSize=6.4,
        leading=7.3,
        alignment=TA_LEFT,
    )

    contenido = {
        item: limitar_caracteres_canvas(datos.get("contenido_lean_canvas", {}).get(item, ""), datos.get("nombre_proyecto", ""), datos.get("codigo_proyecto", ""), 150, 250)
        for item in ITEMS_LEAN_CANVAS
    }
    trl_nivel = str(datos.get("trl_nivel", "TRL 6")).strip() or "TRL 6"
    historia = []

    logo_sena = obtener_ruta_logo_sena()
    if logo_sena and Path(logo_sena).exists():
        try:
            logo_izq = Image(logo_sena, width=2.0 * cm, height=1.25 * cm)
        except Exception:
            logo_izq = Paragraph("SENA", estilo_titulo)
    else:
        logo_izq = Paragraph("SENA", estilo_titulo)

    titulo_header = Paragraph(f"INFORME DE IDENTIFICACIÓN<br/>DEL MODELO DE NEGOCIOS<br/>{trl_nivel}", estilo_titulo)
    fecha_header = Paragraph("Fecha: Marzo 02 de 2020<br/>Versión: 02", estilo_info)

    tabla_header = Table(
        [[logo_izq, titulo_header, fecha_header]],
        colWidths=[5.0 * cm, 13.5 * cm, 7.0 * cm],
        rowHeights=[1.8 * cm],
    )
    tabla_header.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.45, colors.black),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ALIGN", (0, 0), (0, 0), "CENTER"),
        ("ALIGN", (1, 0), (1, 0), "CENTER"),
        ("ALIGN", (2, 0), (2, 0), "CENTER"),
    ]))
    historia.append(tabla_header)
    historia.append(Spacer(1, 0.25 * cm))

    info_data = [
        [Paragraph("TALENTO QUE REALIZA EL INFORME:", estilo_info_bold), Paragraph(str(datos.get("nombre_talento", "No diligenciado")), estilo_info)],
        [Paragraph("NIVEL TRL DEL INFORME:", estilo_info_bold), Paragraph(trl_nivel, estilo_info)],
        [Paragraph("Código del Proyecto:", estilo_info_bold), Paragraph(str(datos.get("codigo_proyecto", "")), estilo_info)],
        [Paragraph("Nombre del Proyecto:", estilo_info_bold), Paragraph(str(datos.get("nombre_proyecto", "")), estilo_info)],
        [Paragraph("GESTOR DE PROYECTO:", estilo_info_bold), Paragraph(str(datos.get("nombre_experto", "No diligenciado")), estilo_info)],
        [Paragraph("LÍNEA DE DESARROLLO TECNOLÓGICO:", estilo_info_bold), Paragraph(str(datos.get("linea_tecnologica", "No diligenciado")), estilo_info)],
        [Paragraph("REGIÓN / CONTEXTO DE IMPLEMENTACIÓN:", estilo_info_bold), Paragraph(str(datos.get("region_contexto", "No diligenciado")), estilo_info)],
    ]

    tabla_info = Table(info_data, colWidths=[7.2 * cm, 18.3 * cm])
    tabla_info.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.35, colors.black),
        ("BACKGROUND", (0, 0), (0, -1), gris),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    historia.append(tabla_info)
    historia.append(Spacer(1, 0.25 * cm))
    historia.append(Paragraph("A continuación se presenta el informe de identificación del modelo de negocios desarrollado para el producto o prototipo descrito.", estilo_info_bold))
    historia.append(Spacer(1, 0.15 * cm))

    colores_preguntas = [verde_claro, azul, amarillo, rojo, verde_claro, azul, amarillo, rojo, gris]
    for idx, item in enumerate(ITEMS_LEAN_CANVAS, start=1):
        titulo = f"{idx}. {item}: {PREGUNTAS_MODELO_NEGOCIO[item]}"
        respuesta = contenido.get(item, "")
        caja = Table(
            [[parrafo_canvas_pdf(titulo, estilo_pregunta)], [parrafo_canvas_pdf(respuesta, estilo_respuesta)]],
            colWidths=[25.5 * cm],
        )
        caja.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.35, colors.black),
            ("BACKGROUND", (0, 0), (0, 0), colores_preguntas[(idx - 1) % len(colores_preguntas)]),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 5),
            ("RIGHTPADDING", (0, 0), (-1, -1), 5),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        historia.append(caja)
        historia.append(Spacer(1, 0.14 * cm))

    historia.append(PageBreak())
    historia.append(Paragraph("MODELO DE NEGOCIOS.", estilo_info_bold))
    historia.append(Spacer(1, 0.2 * cm))

    def celda_canvas(titulo: str) -> Paragraph:
        from html import escape
        texto_completo = limpiar_texto_canvas(contenido.get(titulo, ""))
        return Paragraph(f"<b>{escape(titulo)}</b><br/><br/>{escape(texto_completo)}", estilo_canvas)

    tabla_canvas_data = [
        [
            celda_canvas("Propuesta de valor"),
            celda_canvas("Segmento de clientes y estrategia de adopción"),
            celda_canvas("Canales de distribución"),
        ],
        [
            celda_canvas("Relaciones con clientes"),
            celda_canvas("Flujo de Ingresos"),
            celda_canvas("Recursos Claves"),
        ],
        [
            celda_canvas("Actividades Claves"),
            celda_canvas("Alianzas claves"),
            celda_canvas("Estructura de Costos"),
        ],
    ]

    ancho_total = page_width - doc.leftMargin - doc.rightMargin
    tabla_canvas = Table(
        tabla_canvas_data,
        colWidths=[ancho_total / 3 for _ in range(3)],
        rowHeights=[5.75 * cm, 5.75 * cm, 5.75 * cm],
        repeatRows=0,
    )
    tabla_canvas.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.45, colors.black),
        ("BACKGROUND", (0, 0), (0, 0), amarillo),
        ("BACKGROUND", (1, 0), (1, 0), rojo),
        ("BACKGROUND", (2, 0), (2, 0), azul),
        ("BACKGROUND", (0, 1), (0, 1), verde_claro),
        ("BACKGROUND", (1, 1), (1, 1), amarillo),
        ("BACKGROUND", (2, 1), (2, 1), azul),
        ("BACKGROUND", (0, 2), (0, 2), verde_claro),
        ("BACKGROUND", (1, 2), (1, 2), rojo),
        ("BACKGROUND", (2, 2), (2, 2), gris),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    historia.append(tabla_canvas)

    def encabezado_pie(c, doc):
        c.saveState()
        c.setFont("Helvetica", 7)
        c.setFillColor(colors.grey)
        c.drawCentredString(page_width / 2, 0.32 * cm, f"Informe de identificación del modelo de negocios {trl_nivel} - Red Tecnoparque")
        c.restoreState()

    doc.build(historia, onFirstPage=encabezado_pie, onLaterPages=encabezado_pie)

    datos_json = dict(datos)
    datos_json["ruta_pdf"] = ruta_pdf
    guardar_datos_json(datos_json, ruta="datos_modelo_negocios.json")

    return ruta_pdf

# =====================================================
# SIDEBAR DE CONFIGURACIÓN
# =====================================================
with st.sidebar:
    st.header("Configuración")

    modo_prueba = st.checkbox(
        "Activar modo prueba sin consumir API",
        value=False,
        help="Usa textos generados localmente para probar el flujo sin gastar saldo de OpenAI."
    )

    modelo_openai = st.selectbox("Modelo para generar textos", ["gpt-4.1-mini", "gpt-4.1"], index=0)

    if modo_prueba:
        st.info("Modo prueba activo: no se consumirá API")
    else:
        if obtener_api_key():
            st.success("API Key detectada")
        else:
            st.warning("No se detectó OPENAI_API_KEY")
            st.caption("Configúrala en PowerShell o en .streamlit/secrets.toml")

    st.markdown("---")
    st.caption("Logo principal: recursos/logo_sena.png")
    st.caption("Firmas institucionales: recursos/firmas/")


# =====================================================
# ENCABEZADO
# =====================================================
st.markdown('<div class="main-title">Generador de Documentos Tecnoparque</div>', unsafe_allow_html=True)
st.markdown(
    '<div class="subtitle">Sistema para generar documentos institucionales de proyectos de base tecnológica, innovación y desarrollo tecnológico.</div>',
    unsafe_allow_html=True
)
st.markdown(
    """
    <div class="info-box">
    Selecciona la fase del proceso documental. Para esta versión se habilita el flujo de 
    <b>Fase de inicio → Acta de inicio</b>, <b>Uso de infraestructura</b> y 
    <b>Confidencialidad y compromiso</b>.
    </div>
    """,
    unsafe_allow_html=True
)

# =====================================================
# MENÚ PRINCIPAL
# =====================================================
st.subheader("¿Qué deseas hacer?")

col1, col2, col3, col4 = st.columns(4)
with col1:
    if st.button("📌 Fase de inicio"):
        seleccionar_fase("inicio")
with col2:
    if st.button("🗓️ Fase de planeación"):
        seleccionar_fase("planeacion")
with col3:
    if st.button("📝 Acta de ejecución"):
        seleccionar_fase("ejecucion")
with col4:
    if st.button("✅ Documentos de cierre"):
        seleccionar_fase("cierre")

if st.session_state.fase_seleccionada is None:
    st.stop()


# =====================================================
# FASE DE INICIO
# =====================================================
if st.session_state.fase_seleccionada == "inicio":
    st.markdown("---")
    st.subheader("Documentos de la fase de inicio")

    doc_col1, doc_col2, doc_col3 = st.columns(3)
    with doc_col1:
        if st.button("📄 Acta de inicio"):
            seleccionar_documento("acta_inicio")
    with doc_col2:
        if st.button("🏢 Uso de infraestructura"):
            seleccionar_documento("uso_infraestructura")
    with doc_col3:
        if st.button("🔒 Confidencialidad y compromiso"):
            seleccionar_documento("confidencialidad")

    if st.session_state.documento_seleccionado is None:
        st.info("Selecciona un documento para continuar.")
        st.stop()

    # =====================================================
    # MÓDULO USO DE INFRAESTRUCTURA
    # =====================================================
    if st.session_state.documento_seleccionado == "uso_infraestructura":
        st.markdown("---")
        st.subheader("Formulario para Manual de Préstamo y Uso de Infraestructura")

        st.info(
            "Este documento no consume API de OpenAI. Se genera con base en el formato institucional "
            "de uso de infraestructura de la Red Tecnoparque Colombia."
        )

        with st.form("form_uso_infraestructura"):
            col_a, col_b = st.columns(2)

            with col_a:
                codigo_proyecto = st.text_input("Código del proyecto", placeholder="Ejemplo: P2024-143440-16602")
                nombre_proyecto = st.text_area(
                    "Nombre del proyecto",
                    placeholder='Ejemplo: Diseño de un "Precipitómetro" de bajo costo para determinación de valores reales de infiltración de suelos',
                    height=100
                )
                nombre_talento = st.text_input("Nombre del talento", placeholder="Nombre completo del talento")
                telefono_talento = st.text_input("Teléfono del talento", placeholder="Ejemplo: 324 6428300")

            with col_b:
                nombre_experto = st.text_input("Nombre del experto", placeholder="Nombre completo del experto encargado")
                linea_experto = st.text_input("Línea del experto", placeholder="Ejemplo: ingeniería y diseño")
                ciudad = st.text_input("Ciudad", value="Campoalegre")
                fecha_documento = st.date_input("Fecha del documento", value=date.today())

            generar_infraestructura = st.form_submit_button("Generar documento de uso de infraestructura")

        if generar_infraestructura:
            errores = []
            campos_obligatorios = {
                "Código del proyecto": codigo_proyecto,
                "Nombre del proyecto": nombre_proyecto,
                "Nombre del talento": nombre_talento,
                "Teléfono del talento": telefono_talento,
                "Nombre del experto": nombre_experto,
                "Línea del experto": linea_experto,
                "Ciudad": ciudad,
            }

            for campo, valor in campos_obligatorios.items():
                if not str(valor).strip():
                    errores.append(campo)

            if errores:
                st.error("Faltan campos obligatorios: " + ", ".join(errores))
                st.stop()

            datos_infraestructura = {
                "tipo_documento": "Uso de infraestructura",
                "codigo_proyecto": codigo_proyecto,
                "nombre_proyecto": nombre_proyecto,
                "nombre_talento": nombre_talento,
                "telefono_talento": telefono_talento,
                "nombre_experto": nombre_experto,
                "linea_experto": linea_experto,
                "ciudad": ciudad,
                "fecha_documento": fecha_documento,
                "fecha_corta": fecha_documento.strftime("%d/%m/%Y"),
                "fecha_iso": fecha_documento.strftime("%Y-%m-%d"),
            }

            st.session_state.datos_infraestructura_generada = datos_infraestructura
            st.session_state.ruta_pdf_infraestructura_generado = None
            st.success("Información registrada correctamente. Ahora puedes generar el PDF.")

        if st.session_state.datos_infraestructura_generada:
            datos_infraestructura = st.session_state.datos_infraestructura_generada

            st.markdown("## Resumen para validación")
            st.write("**Tipo de documento:**", datos_infraestructura["tipo_documento"])
            st.write("**Código del proyecto:**", datos_infraestructura["codigo_proyecto"])
            st.write("**Nombre del proyecto:**", datos_infraestructura["nombre_proyecto"])
            st.write("**Nombre del talento:**", datos_infraestructura["nombre_talento"])
            st.write("**Teléfono del talento:**", datos_infraestructura["telefono_talento"])
            st.write("**Nombre del experto:**", datos_infraestructura["nombre_experto"])
            st.write("**Línea del experto:**", datos_infraestructura["linea_experto"])
            st.write("**Ciudad y fecha:**", f'{datos_infraestructura["ciudad"]}, {datos_infraestructura["fecha_corta"]}')

            col_json, col_pdf = st.columns(2)

            with col_json:
                datos_json_descarga = dict(datos_infraestructura)
                if isinstance(datos_json_descarga.get("fecha_documento"), date):
                    datos_json_descarga["fecha_documento"] = datos_json_descarga["fecha_documento"].strftime("%d/%m/%Y")

                st.download_button(
                    label="Descargar datos en JSON",
                    data=json.dumps(datos_json_descarga, ensure_ascii=False, indent=4),
                    file_name="datos_uso_infraestructura.json",
                    mime="application/json"
                )

            with col_pdf:
                if st.button("📄 Generar PDF de uso de infraestructura"):
                    try:
                        ruta_pdf = generar_pdf_uso_infraestructura(datos_infraestructura)
                        st.session_state.ruta_pdf_infraestructura_generado = ruta_pdf
                        st.success(f"PDF generado correctamente: {ruta_pdf}")
                    except Exception as e:
                        st.error(f"No se pudo generar el PDF: {e}")

            if (
                st.session_state.ruta_pdf_infraestructura_generado
                and Path(st.session_state.ruta_pdf_infraestructura_generado).exists()
            ):
                ruta_pdf = st.session_state.ruta_pdf_infraestructura_generado

                with open(ruta_pdf, "rb") as f:
                    st.download_button(
                        label="⬇️ Descargar PDF de uso de infraestructura",
                        data=f,
                        file_name=Path(ruta_pdf).name,
                        mime="application/pdf"
                    )

        st.stop()

    # =====================================================
    # MÓDULO CONFIDENCIALIDAD Y COMPROMISO
    # =====================================================
    if st.session_state.documento_seleccionado == "confidencialidad":
        st.markdown("---")
        st.subheader("Formulario de Confidencialidad y Compromiso")

        st.info(
            "Este documento no consume API de OpenAI. Genera el formato GIC-F-041 V03 "
            "con datos dinámicos del proyecto, talento y firmas institucionales."
        )

        with st.expander("Verificación de firmas institucionales"):
            firmas_requeridas = ["fcaro.png", "fcesar.png", "fdiego.png", "ffelix.png", "fmaria.png", "fsergio.png"]
            for firma in firmas_requeridas:
                ruta = Path(CARPETA_FIRMAS) / firma
                if ruta.exists():
                    st.success(f"Encontrada: {ruta}")
                else:
                    st.warning(f"No encontrada: {ruta}")

        with st.form("form_confidencialidad"):
            col_a, col_b = st.columns(2)

            with col_a:
                codigo_proyecto = st.text_input(
                    "Código del proyecto",
                    placeholder="Ejemplo: I2025-1421611-24230"
                )

                nombre_proyecto = st.text_area(
                    "Nombre del proyecto",
                    placeholder="Ejemplo: Sistema de Bombeo Eléctrico para Mieles de Café en la empresa Reservas del Chapuro de Neiva",
                    height=100
                )

                nombre_talento = st.text_input(
                    "Nombres completos del talento",
                    placeholder="Ejemplo: José Lizardo Ninco Ibarra"
                )

                cedula_talento = st.text_input(
                    "Cédula del talento",
                    placeholder="Ejemplo: 7.728.013"
                )

            with col_b:
                ciudad_expedicion = st.text_input(
                    "Ciudad de expedición de la cédula",
                    placeholder="Ejemplo: Neiva-Huila"
                )

                ciudad = st.text_input(
                    "Ciudad del documento",
                    value="Campoalegre"
                )

                fecha_documento = st.date_input(
                    "Fecha del documento",
                    value=date.today()
                )

                firma_talento_upload = st.file_uploader(
                    "Firma del talento en PNG/JPG",
                    type=["png", "jpg", "jpeg"],
                    help="Opcional. Si no se carga, quedará el espacio de firma en blanco."
                )

            generar_confidencialidad = st.form_submit_button("Generar documento de confidencialidad y compromiso")

        if generar_confidencialidad:
            errores = []

            campos_obligatorios = {
                "Código del proyecto": codigo_proyecto,
                "Nombre del proyecto": nombre_proyecto,
                "Nombres completos del talento": nombre_talento,
                "Cédula del talento": cedula_talento,
                "Ciudad de expedición de la cédula": ciudad_expedicion,
                "Ciudad del documento": ciudad,
            }

            for campo, valor in campos_obligatorios.items():
                if not str(valor).strip():
                    errores.append(campo)

            if errores:
                st.error("Faltan campos obligatorios: " + ", ".join(errores))
                st.stop()

            try:
                ruta_firma_talento_tmp = guardar_archivo_subido(firma_talento_upload, "firma_talento")
            except Exception as e:
                st.error(f"No se pudo procesar la firma del talento: {e}")
                st.stop()

            datos_confidencialidad = {
                "tipo_documento": "Confidencialidad y compromiso",
                "codigo_proyecto": codigo_proyecto,
                "nombre_proyecto": nombre_proyecto,
                "nombre_talento": nombre_talento,
                "cedula_talento": cedula_talento,
                "ciudad_expedicion": ciudad_expedicion,
                "ciudad": ciudad,
                "fecha_documento": fecha_documento,
                "fecha_corta": fecha_documento.strftime("%d/%m/%Y"),
                "fecha_iso": fecha_documento.strftime("%Y-%m-%d"),
                "ruta_firma_talento": ruta_firma_talento_tmp,
            }

            st.session_state.datos_confidencialidad_generada = datos_confidencialidad
            st.session_state.ruta_pdf_confidencialidad_generado = None

            st.success("Información registrada correctamente. Ahora puedes generar el PDF.")

        if st.session_state.datos_confidencialidad_generada:
            datos_confidencialidad = st.session_state.datos_confidencialidad_generada

            st.markdown("## Resumen para validación")
            st.write("**Tipo de documento:**", datos_confidencialidad["tipo_documento"])
            st.write("**Código del proyecto:**", datos_confidencialidad["codigo_proyecto"])
            st.write("**Nombre del proyecto:**", datos_confidencialidad["nombre_proyecto"])
            st.write("**Talento:**", datos_confidencialidad["nombre_talento"])
            st.write("**Cédula:**", datos_confidencialidad["cedula_talento"])
            st.write("**Ciudad de expedición:**", datos_confidencialidad["ciudad_expedicion"])
            st.write("**Ciudad y fecha:**", f'{datos_confidencialidad["ciudad"]}, {datos_confidencialidad["fecha_corta"]}')

            if datos_confidencialidad.get("ruta_firma_talento"):
                st.success("Firma del talento cargada correctamente.")
            else:
                st.warning("No se cargó firma del talento. El documento se generará con el espacio en blanco.")

            col_json, col_pdf = st.columns(2)

            with col_json:
                datos_json_descarga = dict(datos_confidencialidad)
                if isinstance(datos_json_descarga.get("fecha_documento"), date):
                    datos_json_descarga["fecha_documento"] = datos_json_descarga["fecha_documento"].strftime("%d/%m/%Y")

                st.download_button(
                    label="Descargar datos en JSON",
                    data=json.dumps(datos_json_descarga, ensure_ascii=False, indent=4),
                    file_name="datos_confidencialidad_compromiso.json",
                    mime="application/json"
                )

            with col_pdf:
                if st.button("📄 Generar PDF de confidencialidad y compromiso"):
                    try:
                        ruta_pdf = generar_pdf_confidencialidad(datos_confidencialidad)
                        st.session_state.ruta_pdf_confidencialidad_generado = ruta_pdf
                        st.success(f"PDF generado correctamente: {ruta_pdf}")
                    except Exception as e:
                        st.error(f"No se pudo generar el PDF: {e}")

            if (
                st.session_state.ruta_pdf_confidencialidad_generado
                and Path(st.session_state.ruta_pdf_confidencialidad_generado).exists()
            ):
                ruta_pdf = st.session_state.ruta_pdf_confidencialidad_generado

                with open(ruta_pdf, "rb") as f:
                    st.download_button(
                        label="⬇️ Descargar PDF de confidencialidad y compromiso",
                        data=f,
                        file_name=Path(ruta_pdf).name,
                        mime="application/pdf"
                    )

        st.stop()

    # =====================================================
    # MÓDULO ACTA DE INICIO
    # =====================================================
    if st.session_state.documento_seleccionado != "acta_inicio":
        st.warning("Este módulo estará disponible en una siguiente versión.")
        st.stop()

    st.markdown("---")
    st.subheader("Formulario para Acta de Inicio")

    with st.form("form_acta_inicio"):
        col_a, col_b = st.columns(2)
        with col_a:
            codigo_acta = st.text_input("Código o número del acta / proyecto", placeholder="Ejemplo: P2026-143440-00001")
            nombre_proyecto = st.text_area("Nombre del proyecto", placeholder="Ejemplo: Diseño de un sistema electrónico para medición...", height=90)
            fecha_acta = st.date_input("Fecha del acta", value=date.today())
            hora_inicio = st.time_input("Hora de inicio", value=time(8, 0))
            ciudad = st.text_input("Ciudad", value="Campoalegre")

        with col_b:
            nombre_talento = st.text_input("Nombre del talento", placeholder="Nombre completo del talento")
            nombre_experto = st.text_input("Nombre del experto", placeholder="Nombre completo del experto")
            linea_sublinea = st.text_input("Línea y sublínea de Tecnoparque", placeholder="Ejemplo: Productos y procesos / Diseño de producto")

        descripcion_proyecto = st.text_area(
            "Descripción general del proyecto",
            placeholder="Describe la necesidad, el problema, la solución propuesta, los componentes tecnológicos y el resultado esperado.",
            height=170
        )

        texto_boton = "Generar resumen del acta en modo prueba" if modo_prueba else "Generar resumen del acta con ChatGPT"
        generar = st.form_submit_button(texto_boton)

    if generar:
        errores = []
        campos_obligatorios = {
            "Código o número del acta / proyecto": codigo_acta,
            "Nombre del proyecto": nombre_proyecto,
            "Nombre del talento": nombre_talento,
            "Nombre del experto": nombre_experto,
            "Línea y sublínea": linea_sublinea,
            "Descripción general del proyecto": descripcion_proyecto,
        }

        for campo, valor in campos_obligatorios.items():
            if not str(valor).strip():
                errores.append(campo)

        if errores:
            st.error("Faltan campos obligatorios: " + ", ".join(errores))
            st.stop()

        hora_fin = calcular_hora_fin(fecha_acta, hora_inicio)
        objetivo_reunion = construir_objetivo_reunion(nombre_proyecto)

        datos_acta = {
            "tipo_documento": "Acta de inicio",
            "codigo_acta": codigo_acta,
            "nombre_proyecto": nombre_proyecto,
            "ciudad": ciudad,
            "fecha_acta": fecha_acta.strftime("%d/%m/%Y"),
            "fecha_iso": fecha_acta.strftime("%Y-%m-%d"),
            "hora_inicio": hora_inicio.strftime("%H:%M"),
            "hora_fin": hora_fin.strftime("%H:%M"),
            "nombre_talento": nombre_talento,
            "nombre_experto": nombre_experto,
            "linea_sublinea": linea_sublinea,
            "descripcion_proyecto": descripcion_proyecto,
            "objetivo_reunion": objetivo_reunion,
        }

        mensaje_spinner = "Generando textos en modo prueba..." if modo_prueba else "Generando objetivo general, objetivos específicos y alcance con ChatGPT..."
        with st.spinner(mensaje_spinner):
            try:
                if modo_prueba:
                    textos_ia = generar_textos_modo_prueba(nombre_proyecto, descripcion_proyecto)
                else:
                    textos_ia = generar_textos_con_chatgpt(nombre_proyecto, descripcion_proyecto, modelo_openai)

                datos_acta["objetivo_general"] = textos_ia.get("objetivo_general", "")
                datos_acta["objetivos_especificos"] = textos_ia.get("objetivos_especificos", [])
                datos_acta["alcance"] = textos_ia.get("alcance", "")
                datos_acta["modo_generacion"] = "Prueba local" if modo_prueba else "ChatGPT API"

            except Exception as e:
                st.error(f"No se pudo generar el contenido: {e}")
                st.stop()

        guardar_datos_json(datos_acta)
        st.session_state.datos_acta_generada = datos_acta
        st.session_state.ruta_pdf_generado = None
        st.success("Resumen generado correctamente. La hora de finalización fue calculada automáticamente.")

    if st.session_state.datos_acta_generada:
        datos_acta = st.session_state.datos_acta_generada

        st.markdown("## Resumen para validación")
        st.write("**Modo de generación:**", datos_acta["modo_generacion"])
        st.write("**Código / Acta:**", datos_acta["codigo_acta"])
        st.write("**Nombre del proyecto:**", datos_acta["nombre_proyecto"])
        st.write("**Ciudad y fecha:**", f'{datos_acta["ciudad"]}, {datos_acta["fecha_acta"]}')
        st.write("**Hora de inicio:**", datos_acta["hora_inicio"])
        st.write("**Hora de finalización:**", datos_acta["hora_fin"])
        st.write("**Talento:**", datos_acta["nombre_talento"])
        st.write("**Experto:**", datos_acta["nombre_experto"])
        st.write("**Línea y sublínea:**", datos_acta["linea_sublinea"])
        st.write("**Objetivo de la reunión:**", datos_acta["objetivo_reunion"])
        st.write("**Descripción general:**", datos_acta["descripcion_proyecto"])

        st.markdown("## Textos generados")
        st.markdown("### Objetivo general")
        st.write(datos_acta["objetivo_general"])

        st.markdown("### Objetivos específicos")
        if datos_acta["objetivos_especificos"]:
            for i, objetivo in enumerate(datos_acta["objetivos_especificos"], start=1):
                st.write(f"{i}. {objetivo}")
        else:
            st.warning("No se generaron objetivos específicos en formato de lista.")

        st.markdown("### Alcance del proyecto")
        st.write(datos_acta["alcance"])

        col_json, col_pdf = st.columns(2)

        with col_json:
            st.download_button(
                label="Descargar datos generados en JSON",
                data=json.dumps(datos_acta, ensure_ascii=False, indent=4),
                file_name="datos_acta_inicio.json",
                mime="application/json"
            )

        with col_pdf:
            if st.button("📄 Generar PDF del acta"):
                try:
                    ruta_pdf = generar_pdf_acta_inicio(datos_acta)
                    st.session_state.ruta_pdf_generado = ruta_pdf
                    st.success(f"PDF generado correctamente: {ruta_pdf}")
                except Exception as e:
                    st.error(f"No se pudo generar el PDF: {e}")

        if st.session_state.ruta_pdf_generado and Path(st.session_state.ruta_pdf_generado).exists():
            ruta_pdf = st.session_state.ruta_pdf_generado
            with open(ruta_pdf, "rb") as f:
                st.download_button(
                    label="⬇️ Descargar PDF del acta",
                    data=f,
                    file_name=Path(ruta_pdf).name,
                    mime="application/pdf"
                )

elif st.session_state.fase_seleccionada == "planeacion":
    st.markdown("---")
    st.subheader("Documentos de la fase de planeación")

    plan_col1, plan_col2 = st.columns(2)

    with plan_col1:
        if st.button("📊 Cronograma de actividades"):
            seleccionar_documento("cronograma")

    with plan_col2:
        if st.button("📚 Estado del arte"):
            seleccionar_documento("estado_arte")

    if st.session_state.documento_seleccionado is None:
        st.info("Selecciona un documento de planeación para continuar.")
        st.stop()

    # =====================================================
    # MÓDULO ESTADO DEL ARTE
    # =====================================================
    if st.session_state.documento_seleccionado == "estado_arte":
        st.markdown("---")
        st.subheader("Formulario para Estado del Arte")

        st.info(
            "Este módulo genera un documento académico e investigativo con búsqueda de proyectos similares, "
            "tecnologías relevantes, artículos de validación y referencias en APA 7."
        )

        with st.form("form_estado_arte"):
            col_a, col_b = st.columns(2)

            with col_a:
                codigo_proyecto = st.text_input(
                    "Código del proyecto",
                    placeholder="Ejemplo: P2024-143440-16600"
                )

                nombre_proyecto = st.text_area(
                    "Nombre del proyecto",
                    placeholder="Título oficial de la iniciativa",
                    height=80
                )

            with col_b:
                fecha_documento = st.date_input(
                    "Fecha del documento",
                    value=date.today()
                )

                tecnologias_previstas_texto = st.text_area(
                    "Tecnologías previstas",
                    placeholder="Ejemplo: Inteligencia artificial, animación 3D, códigos QR, realidad aumentada",
                    height=80
                )

            descripcion_proyecto = st.text_area(
                "Háblame sobre el proyecto",
                placeholder="Describe de qué trata, quién es el talento detrás, cuál es el origen de la iniciativa y qué la hace innovadora.",
                height=220
            )

            generar_estado_arte = st.form_submit_button("Generar Estado del Arte")

        if generar_estado_arte:
            errores = []

            campos_obligatorios = {
                "Código del proyecto": codigo_proyecto,
                "Nombre del proyecto": nombre_proyecto,
                "Descripción detallada": descripcion_proyecto,
                "Tecnologías previstas": tecnologias_previstas_texto,
            }

            for campo, valor in campos_obligatorios.items():
                if not str(valor).strip():
                    errores.append(campo)

            if errores:
                st.error("Faltan campos obligatorios: " + ", ".join(errores))
                st.stop()

            tecnologias_previstas = limpiar_lista_tecnologias(tecnologias_previstas_texto)

            with st.spinner("Generando Estado del Arte con búsqueda académica y referencias APA 7..."):
                try:
                    if modo_prueba:
                        contenido_estado_arte = generar_estado_arte_modo_prueba(
                            nombre_proyecto,
                            codigo_proyecto,
                            descripcion_proyecto,
                            tecnologias_previstas,
                        )
                    else:
                        contenido_estado_arte = generar_estado_arte_con_chatgpt(
                            nombre_proyecto,
                            codigo_proyecto,
                            descripcion_proyecto,
                            tecnologias_previstas,
                            modelo_openai,
                        )
                except Exception as e:
                    st.error(f"No se pudo generar el Estado del Arte: {e}")
                    st.stop()

            datos_estado_arte = {
                "tipo_documento": "Estado del Arte",
                "codigo_proyecto": codigo_proyecto,
                "nombre_proyecto": nombre_proyecto,
                "fecha_documento": fecha_documento,
                "descripcion_proyecto": descripcion_proyecto,
                "tecnologias_previstas": tecnologias_previstas,
                "contenido_estado_arte": contenido_estado_arte,
                "modo_generacion": "Prueba local" if modo_prueba else "ChatGPT API con búsqueda web",
            }

            st.session_state.datos_estado_arte_generado = datos_estado_arte
            st.session_state.ruta_pdf_estado_arte_generado = None

            st.success("Estado del Arte generado correctamente. Ahora puedes revisar y generar el PDF.")

        if st.session_state.datos_estado_arte_generado:
            datos_estado_arte = st.session_state.datos_estado_arte_generado
            contenido = datos_estado_arte["contenido_estado_arte"]

            st.markdown("## Resumen para validación")

            st.write("**Modo de generación:**", datos_estado_arte["modo_generacion"])
            st.write("**Código del proyecto:**", datos_estado_arte["codigo_proyecto"])
            st.write("**Nombre del proyecto:**", datos_estado_arte["nombre_proyecto"])
            st.write("**Fecha:**", datos_estado_arte["fecha_documento"].strftime("%d/%m/%Y"))
            st.write("**Tecnologías previstas:**", ", ".join(datos_estado_arte["tecnologias_previstas"]))

            st.markdown("### Introducción")
            st.write(contenido.get("introduccion", ""))

            st.markdown("### Objetivos")
            for objetivo in contenido.get("objetivos", []):
                st.write(f"- {objetivo}")

            st.markdown("### Proyectos similares")
            for item in contenido.get("proyectos_similares", []):
                st.write(f"**{item.get('numero', '')}. {item.get('nombre', '')}**")
                st.write(item.get("enlace", ""))

            st.markdown("### Tecnologías emergentes")
            for item in contenido.get("tecnologias_emergentes", []):
                st.write(f"**{item.get('tecnologia', '')}:** {item.get('analisis', '')}")

            col_json, col_pdf = st.columns(2)

            with col_json:
                datos_json_descarga = dict(datos_estado_arte)
                if isinstance(datos_json_descarga.get("fecha_documento"), date):
                    datos_json_descarga["fecha_documento"] = datos_json_descarga["fecha_documento"].strftime("%d/%m/%Y")

                st.download_button(
                    label="Descargar datos en JSON",
                    data=json.dumps(datos_json_descarga, ensure_ascii=False, indent=4),
                    file_name="datos_estado_arte.json",
                    mime="application/json"
                )

            with col_pdf:
                if st.button("📄 Generar PDF del Estado del Arte"):
                    try:
                        ruta_pdf = generar_pdf_estado_arte(datos_estado_arte)
                        st.session_state.ruta_pdf_estado_arte_generado = ruta_pdf
                        st.success(f"PDF generado correctamente: {ruta_pdf}")
                    except Exception as e:
                        st.error(f"No se pudo generar el PDF: {e}")

            if (
                st.session_state.ruta_pdf_estado_arte_generado
                and Path(st.session_state.ruta_pdf_estado_arte_generado).exists()
            ):
                ruta_pdf = st.session_state.ruta_pdf_estado_arte_generado

                with open(ruta_pdf, "rb") as f:
                    st.download_button(
                        label="⬇️ Descargar PDF del Estado del Arte",
                        data=f,
                        file_name=Path(ruta_pdf).name,
                        mime="application/pdf"
                    )

        st.stop()

    # =====================================================
    # MÓDULO CRONOGRAMA DE ACTIVIDADES
    # =====================================================
    if st.session_state.documento_seleccionado == "cronograma":
        st.markdown("---")
        st.subheader("Formulario para Cronograma de Actividades")

        st.info(
            "Este módulo genera un cronograma horizontal tipo diagrama de Gantt. "
            "Las actividades se generan con la API de OpenAI de acuerdo con la descripción del proyecto."
        )

        with st.form("form_cronograma"):
            col_a, col_b = st.columns(2)

            with col_a:
                codigo_proyecto = st.text_input(
                    "Código del proyecto",
                    placeholder="Ejemplo: P2024-143440-16602"
                )

                nombre_proyecto = st.text_area(
                    "Nombre del proyecto",
                    placeholder='Ejemplo: Diseño de un "Precipitómetro" de bajo costo para determinación de valores reales de infiltración de suelos',
                    height=90
                )

                nombre_talento = st.text_input(
                    "Nombre del talento",
                    placeholder="Nombre completo del talento"
                )

                nombre_experto = st.text_input(
                    "Nombre del experto",
                    placeholder="Nombre completo del experto"
                )

            with col_b:
                linea = st.text_input(
                    "Línea",
                    placeholder="Ejemplo: Diseño de productos"
                )

                cantidad_actividades = st.number_input(
                    "Cantidad de actividades",
                    min_value=3,
                    max_value=20,
                    value=7,
                    step=1
                )

                fecha_inicio = st.date_input(
                    "Fecha de inicio",
                    value=date.today()
                )

                fecha_fin = st.date_input(
                    "Fecha de finalización",
                    value=date.today() + timedelta(days=60)
                )

                dias_semana = st.multiselect(
                    "Día(s) de la semana para programar actividades",
                    options=["Lunes", "Martes", "Miércoles", "Jueves", "Viernes", "Sábado", "Domingo"],
                    default=["Sábado", "Domingo"]
                )

            descripcion_proyecto = st.text_area(
                "Describe el proyecto",
                placeholder="Describe la necesidad, la solución propuesta, los componentes técnicos, el prototipo o resultado esperado.",
                height=160
            )

            generar_cronograma = st.form_submit_button(
                "Generar cronograma de actividades"
            )

        if generar_cronograma:
            errores = []

            campos_obligatorios = {
                "Código del proyecto": codigo_proyecto,
                "Nombre del proyecto": nombre_proyecto,
                "Nombre del talento": nombre_talento,
                "Nombre del experto": nombre_experto,
                "Línea": linea,
                "Descripción del proyecto": descripcion_proyecto,
            }

            for campo, valor in campos_obligatorios.items():
                if not str(valor).strip():
                    errores.append(campo)

            if not dias_semana:
                errores.append("Día(s) de la semana")

            if fecha_fin < fecha_inicio:
                errores.append("La fecha de finalización no puede ser anterior a la fecha de inicio")

            if errores:
                st.error("Revisa los siguientes campos: " + ", ".join(errores))
                st.stop()

            fechas_programadas = obtener_fechas_programadas(
                fecha_inicio,
                fecha_fin,
                dias_semana
            )

            if not fechas_programadas:
                st.error("No se encontraron fechas programadas con los días seleccionados dentro del rango indicado.")
                st.stop()

            with st.spinner("Generando actividades técnicas con IA..."):
                try:
                    if modo_prueba:
                        actividades = generar_actividades_cronograma_modo_prueba(
                            descripcion_proyecto,
                            int(cantidad_actividades)
                        )
                    else:
                        actividades = generar_actividades_cronograma_con_chatgpt(
                            descripcion_proyecto,
                            int(cantidad_actividades),
                            modelo_openai
                        )
                except Exception as e:
                    st.error(f"No se pudieron generar las actividades: {e}")
                    st.stop()

            bloques_fechas = dividir_fechas_por_actividad(
                fechas_programadas,
                int(cantidad_actividades)
            )

            datos_cronograma = {
                "tipo_documento": "Cronograma de actividades",
                "codigo_proyecto": codigo_proyecto,
                "nombre_proyecto": nombre_proyecto,
                "nombre_talento": nombre_talento,
                "nombre_experto": nombre_experto,
                "linea": linea,
                "cantidad_actividades": int(cantidad_actividades),
                "fecha_inicio": fecha_inicio,
                "fecha_fin": fecha_fin,
                "dias_semana": dias_semana,
                "descripcion_proyecto": descripcion_proyecto,
                "actividades": actividades,
                "fechas_programadas": fechas_programadas,
                "bloques_fechas": bloques_fechas,
                "modo_generacion": "Prueba local" if modo_prueba else "ChatGPT API",
            }

            st.session_state.datos_cronograma_generado = datos_cronograma
            st.session_state.ruta_pdf_cronograma_generado = None

            st.success("Cronograma generado correctamente. Ahora puedes revisar y generar el PDF.")

        if st.session_state.datos_cronograma_generado:
            datos_cronograma = st.session_state.datos_cronograma_generado

            st.markdown("## Resumen para validación")

            st.write("**Modo de generación:**", datos_cronograma["modo_generacion"])
            st.write("**Código del proyecto:**", datos_cronograma["codigo_proyecto"])
            st.write("**Nombre del proyecto:**", datos_cronograma["nombre_proyecto"])
            st.write("**Talento:**", datos_cronograma["nombre_talento"])
            st.write("**Experto:**", datos_cronograma["nombre_experto"])
            st.write("**Línea:**", datos_cronograma["linea"])
            st.write(
                "**Periodo:**",
                f'{datos_cronograma["fecha_inicio"].strftime("%d/%m/%Y")} al {datos_cronograma["fecha_fin"].strftime("%d/%m/%Y")}'
            )
            st.write("**Días programados:**", ", ".join(datos_cronograma["dias_semana"]))

            st.markdown("### Actividades generadas")
            for idx, actividad in enumerate(datos_cronograma["actividades"], start=1):
                fechas_actividad = datos_cronograma["bloques_fechas"][idx - 1]
                if fechas_actividad:
                    periodo = f"{fechas_actividad[0].strftime('%d/%m/%Y')} al {fechas_actividad[-1].strftime('%d/%m/%Y')}"
                else:
                    periodo = "Sin fecha asignada"

                st.write(f"**{idx}.** {actividad} — {periodo}")

            col_json, col_pdf = st.columns(2)

            with col_json:
                datos_json_descarga = dict(datos_cronograma)
                datos_json_descarga["fecha_inicio"] = datos_json_descarga["fecha_inicio"].strftime("%d/%m/%Y")
                datos_json_descarga["fecha_fin"] = datos_json_descarga["fecha_fin"].strftime("%d/%m/%Y")
                datos_json_descarga["fechas_programadas"] = [
                    f.strftime("%d/%m/%Y") for f in datos_json_descarga["fechas_programadas"]
                ]
                datos_json_descarga["bloques_fechas"] = [
                    [f.strftime("%d/%m/%Y") for f in bloque]
                    for bloque in datos_json_descarga["bloques_fechas"]
                ]

                st.download_button(
                    label="Descargar datos en JSON",
                    data=json.dumps(datos_json_descarga, ensure_ascii=False, indent=4),
                    file_name="datos_cronograma_actividades.json",
                    mime="application/json"
                )

            with col_pdf:
                if st.button("📄 Generar PDF del cronograma"):
                    try:
                        ruta_pdf = generar_pdf_cronograma(datos_cronograma)
                        st.session_state.ruta_pdf_cronograma_generado = ruta_pdf
                        st.success(f"PDF generado correctamente: {ruta_pdf}")
                    except Exception as e:
                        st.error(f"No se pudo generar el PDF: {e}")

            if (
                st.session_state.ruta_pdf_cronograma_generado
                and Path(st.session_state.ruta_pdf_cronograma_generado).exists()
            ):
                ruta_pdf = st.session_state.ruta_pdf_cronograma_generado

                with open(ruta_pdf, "rb") as f:
                    st.download_button(
                        label="⬇️ Descargar PDF del cronograma",
                        data=f,
                        file_name=Path(ruta_pdf).name,
                        mime="application/pdf"
                    )

        st.stop()

    st.warning("Selecciona una opción válida de la fase de planeación.")
    st.stop()

    
elif st.session_state.fase_seleccionada == "ejecucion":
    st.markdown("---")
    st.subheader("Acta de ejecución")

    st.info(
        "Este módulo genera el documento de seguimiento de asesorías y usos de infraestructura. "
        "Incluye cálculo automático de honorarios, equipos, materiales y total general."
    )

    metodo_asesorias = st.radio(
        "Método de generación de asesorías y usos",
        options=["Generación con IA", "Generación manual"],
        horizontal=True,
    )

    st.markdown("### Equipos y materiales")
    cantidad_equipos = st.number_input(
        "¿Cuántos registros de uso de equipos y materiales desea ingresar?",
        min_value=0,
        max_value=30,
        value=1,
        step=1,
    )

    st.markdown("### Asesorías y usos")
    if metodo_asesorias == "Generación manual":
        cantidad_asesorias_manual = st.number_input(
            "Cantidad de registros de asesorías y usos",
            min_value=1,
            max_value=30,
            value=4,
            step=1,
        )
    else:
        cantidad_asesorias_manual = 0

    with st.form("form_acta_ejecucion"):
        st.markdown("## Información general")

        col_a, col_b = st.columns(2)

        with col_a:
            codigo_proyecto = st.text_input(
                "Código del proyecto",
                placeholder="Ejemplo: P2025-1431026-17218"
            )

            nombre_proyecto = st.text_area(
                "Nombre del proyecto",
                placeholder="Nombre oficial del proyecto",
                height=90
            )

            sublinea_tecnologica = st.text_input(
                "Sublínea tecnológica",
                placeholder="Ejemplo: IND - Productos y procesos"
            )

            nombre_experto = st.text_input(
                "Nombre del experto",
                placeholder="Nombre completo del experto"
            )

        with col_b:
            nombre_talento = st.text_input(
                "Nombre del talento",
                placeholder="Nombre completo del talento interlocutor"
            )

            documento_talento = st.text_input(
                "Documento de identidad del talento",
                placeholder="Número de documento"
            )

            telefono_talento = st.text_input(
                "Teléfono del talento",
                placeholder="Número de contacto"
            )

        st.markdown("## Equipos, materiales e insumos utilizados")

        equipos_materiales = []

        for i in range(int(cantidad_equipos)):
            st.markdown(f"**Registro equipo/material {i + 1}**")
            col1, col2, col3 = st.columns([2, 1, 1])

            with col1:
                nombre_equipo = st.text_input(
                    f"Nombre del equipo / Material usado {i + 1}",
                    key=f"equipo_nombre_{i}"
                )

            with col2:
                cantidad_horas = st.text_input(
                    f"Horas de uso / Cantidad {i + 1}",
                    key=f"equipo_cantidad_{i}"
                )

            with col3:
                valor_total = st.number_input(
                    f"Valor total {i + 1}",
                    min_value=0,
                    value=0,
                    step=1000,
                    key=f"equipo_valor_{i}"
                )

            equipos_materiales.append(
                {
                    "nombre": nombre_equipo,
                    "cantidad_horas": cantidad_horas,
                    "valor_total": valor_total,
                }
            )

        st.markdown("## Asesorías y usos")

        asesorias = []
        descripcion_proyecto = ""
        fecha_inicio = None
        fecha_fin = None
        dias_ejecucion = []
        cantidad_asesorias_ia = 0
        horas_por_asesoria_ia = 0

        if metodo_asesorias == "Generación manual":
            for i in range(int(cantidad_asesorias_manual)):
                st.markdown(f"**Asesoría / uso {i + 1}**")
                col1, col2, col3 = st.columns([1.2, 1, 3])

                with col1:
                    fecha_asesoria = st.date_input(
                        f"Fecha asesoría {i + 1}",
                        value=date.today(),
                        key=f"asesoria_fecha_manual_{i}"
                    )

                with col2:
                    horas_asesoria = st.number_input(
                        f"Horas asesoría {i + 1}",
                        min_value=0.5,
                        max_value=12.0,
                        value=2.0,
                        step=0.5,
                        key=f"asesoria_horas_manual_{i}"
                    )

                with col3:
                    descripcion_asesoria = st.text_area(
                        f"Descripción {i + 1}",
                        height=70,
                        key=f"asesoria_desc_manual_{i}"
                    )

                asesorias.append(
                    {
                        "fecha": fecha_asesoria.strftime("%d/%m/%Y"),
                        "horas": horas_asesoria,
                        "descripcion": descripcion_asesoria,
                    }
                )

        else:
            col1, col2 = st.columns(2)

            with col1:
                cantidad_asesorias_ia = st.number_input(
                    "Cantidad de asesorías",
                    min_value=1,
                    max_value=30,
                    value=4,
                    step=1
                )

                horas_por_asesoria_ia = st.number_input(
                    "Horas por asesoría",
                    min_value=0.5,
                    max_value=12.0,
                    value=2.0,
                    step=0.5
                )

                fecha_inicio = st.date_input(
                    "Fecha de inicio",
                    value=date.today()
                )

            with col2:
                fecha_fin = st.date_input(
                    "Fecha de fin",
                    value=date.today() + timedelta(days=30)
                )

                dias_ejecucion = st.multiselect(
                    "Días de ejecución",
                    options=["Lunes", "Martes", "Miércoles", "Jueves", "Viernes", "Sábado", "Domingo"],
                    default=["Martes", "Jueves"]
                )

            descripcion_proyecto = st.text_area(
                "Descripción general del proyecto",
                placeholder="Describe el proyecto para que la IA genere actividades de asesoría coherentes.",
                height=140
            )

        generar_acta_ejecucion = st.form_submit_button("Generar Acta de Ejecución")

    if generar_acta_ejecucion:
        errores = []

        campos_obligatorios = {
            "Código del proyecto": codigo_proyecto,
            "Nombre del proyecto": nombre_proyecto,
            "Sublínea tecnológica": sublinea_tecnologica,
            "Nombre del experto": nombre_experto,
            "Nombre del talento": nombre_talento,
            "Documento del talento": documento_talento,
            "Teléfono del talento": telefono_talento,
        }

        for campo, valor in campos_obligatorios.items():
            if not str(valor).strip():
                errores.append(campo)

        equipos_validos = []
        for item in equipos_materiales:
            if str(item["nombre"]).strip() or str(item["cantidad_horas"]).strip() or item["valor_total"] > 0:
                equipos_validos.append(item)

        if metodo_asesorias == "Generación manual":
            asesorias_validas = []
            for item in asesorias:
                if str(item["descripcion"]).strip():
                    asesorias_validas.append(item)

            if not asesorias_validas:
                errores.append("Debe ingresar al menos una asesoría manual con descripción")

        else:
            if not descripcion_proyecto.strip():
                errores.append("Descripción general del proyecto")

            if not dias_ejecucion:
                errores.append("Días de ejecución")

            if fecha_fin < fecha_inicio:
                errores.append("La fecha de fin no puede ser anterior a la fecha de inicio")

            asesorias_validas = []

        if errores:
            st.error("Revisa los siguientes campos: " + ", ".join(errores))
            st.stop()

        if metodo_asesorias == "Generación con IA":
            fechas_asesorias = distribuir_fechas_para_asesorias(
                fecha_inicio,
                fecha_fin,
                dias_ejecucion,
                int(cantidad_asesorias_ia)
            )

            if not fechas_asesorias:
                st.error("No se encontraron fechas válidas para distribuir las asesorías.")
                st.stop()

            with st.spinner("Generando asesorías con IA..."):
                if modo_prueba:
                    asesorias_validas = generar_asesorias_ejecucion_modo_prueba(
                        descripcion_proyecto,
                        int(cantidad_asesorias_ia),
                        fechas_asesorias,
                        float(horas_por_asesoria_ia)
                    )
                else:
                    asesorias_validas = generar_asesorias_ejecucion_con_chatgpt(
                        descripcion_proyecto,
                        int(cantidad_asesorias_ia),
                        fechas_asesorias,
                        float(horas_por_asesoria_ia),
                        modelo_openai
                    )

        total_equipos_materiales = sum(float(item.get("valor_total", 0)) for item in equipos_validos)

        if metodo_asesorias == "Generación con IA":
            total_honorarios = int(cantidad_asesorias_ia) * float(horas_por_asesoria_ia) * VALOR_HORA_EXPERTO
        else:
            total_horas_manual = sum(float(item.get("horas", 0)) for item in asesorias_validas)
            total_honorarios = total_horas_manual * VALOR_HORA_EXPERTO

        total_general = total_equipos_materiales + total_honorarios

        datos_acta_ejecucion = {
            "tipo_documento": "Acta de Ejecución",
            "titulo_acta": f"Acta 2 - {codigo_proyecto}",
            "codigo_proyecto": codigo_proyecto,
            "nombre_proyecto": nombre_proyecto,
            "sublinea_tecnologica": sublinea_tecnologica,
            "nombre_experto": nombre_experto,
            "nombre_talento": nombre_talento,
            "documento_talento": documento_talento,
            "telefono_talento": telefono_talento,
            "equipos_materiales": equipos_validos,
            "asesorias": asesorias_validas,
            "metodo_asesorias": metodo_asesorias,
            "total_equipos_materiales": total_equipos_materiales,
            "total_honorarios": total_honorarios,
            "total_general": total_general,
            "valor_hora_experto": VALOR_HORA_EXPERTO,
            "modo_generacion": "Prueba local" if modo_prueba else "ChatGPT API" if metodo_asesorias == "Generación con IA" else "Manual",
        }

        st.session_state.datos_acta_ejecucion_generada = datos_acta_ejecucion
        st.session_state.ruta_pdf_acta_ejecucion_generado = None

        st.success("Acta de ejecución generada correctamente. Ahora puedes revisar y generar el PDF.")

    if st.session_state.datos_acta_ejecucion_generada:
        datos = st.session_state.datos_acta_ejecucion_generada

        st.markdown("## Resumen para validación")
        st.write("**Título:**", datos["titulo_acta"])
        st.write("**Proyecto:**", datos["nombre_proyecto"])
        st.write("**Código:**", datos["codigo_proyecto"])
        st.write("**Experto:**", datos["nombre_experto"])
        st.write("**Talento:**", datos["nombre_talento"])
        st.write("**Método de asesorías:**", datos["metodo_asesorias"])
        st.write("**Total equipos y materiales:**", formato_moneda_colombiana(datos["total_equipos_materiales"]))
        st.write("**Total honorarios:**", formato_moneda_colombiana(datos["total_honorarios"]))
        st.write("**Total general:**", formato_moneda_colombiana(datos["total_general"]))

        st.markdown("### Asesorías y usos")
        for idx, item in enumerate(datos["asesorias"], start=1):
            st.write(f"**{idx}.** {item['fecha']} — {item['horas']} horas — {item['descripcion']}")

        st.markdown("### Equipos y materiales")
        if datos["equipos_materiales"]:
            for idx, item in enumerate(datos["equipos_materiales"], start=1):
                st.write(
                    f"**{idx}.** {item['nombre']} — {item['cantidad_horas']} — "
                    f"{formato_moneda_colombiana(item['valor_total'])}"
                )
        else:
            st.info("No se registraron equipos o materiales.")

        col_json, col_pdf = st.columns(2)

        with col_json:
            st.download_button(
                label="Descargar datos en JSON",
                data=json.dumps(datos, ensure_ascii=False, indent=4),
                file_name="datos_acta_ejecucion.json",
                mime="application/json"
            )

        with col_pdf:
            if st.button("📄 Generar PDF del Acta de Ejecución"):
                try:
                    ruta_pdf = generar_pdf_acta_ejecucion(datos)
                    st.session_state.ruta_pdf_acta_ejecucion_generado = ruta_pdf
                    st.success(f"PDF generado correctamente: {ruta_pdf}")
                except Exception as e:
                    st.error(f"No se pudo generar el PDF: {e}")

        if (
            st.session_state.ruta_pdf_acta_ejecucion_generado
            and Path(st.session_state.ruta_pdf_acta_ejecucion_generado).exists()
        ):
            ruta_pdf = st.session_state.ruta_pdf_acta_ejecucion_generado

            with open(ruta_pdf, "rb") as f:
                st.download_button(
                    label="⬇️ Descargar PDF del Acta de Ejecución",
                    data=f,
                    file_name=Path(ruta_pdf).name,
                    mime="application/pdf"
                )

elif st.session_state.fase_seleccionada == "cierre":
    st.markdown("---")
    st.subheader("Documentos de la fase de cierre")

    cierre_col1, cierre_col2, cierre_col3 = st.columns(3)

    with cierre_col1:
        if st.button("📄 Generar acta de cierre y ficha de caracterización"):
            seleccionar_documento("acta_cierre_ficha")

    with cierre_col2:
        if st.button("🧩 Modelo de negocio Lean Canvas"):
            seleccionar_documento("lean_canvas")

    with cierre_col3:
        if st.button("📘 Informe técnico final"):
            seleccionar_documento("informe_tecnico_final")

    if st.session_state.documento_seleccionado is None:
        st.info("Selecciona un documento de cierre para continuar.")
        st.stop()

    if st.session_state.documento_seleccionado == "lean_canvas":
        st.markdown("---")
        st.subheader("Modelo de negocio Lean Canvas")

        st.info(
            "Este módulo genera un Informe de Identificación del Modelo de Negocios en PDF. "
            "Selecciona el nivel TRL correspondiente y diligencia los insumos mínimos para que el sistema construya respuestas aplicadas al producto, al contexto y a la región de implementación."
        )

        with st.form("form_lean_canvas"):
            col_a, col_b = st.columns(2)

            with col_a:
                codigo_proyecto_lc = st.text_input(
                    "Código del proyecto",
                    placeholder="Ejemplo: P2026-143440-00001"
                )

                trl_nivel_lc = st.selectbox(
                    "Nivel TRL del informe",
                    options=["TRL 6", "TRL 7"],
                    index=0
                )

                nombre_proyecto_lc = st.text_area(
                    "Nombre del proyecto",
                    placeholder="Nombre oficial del proyecto",
                    height=90
                )

                region_contexto_lc = st.text_area(
                    "Región o contexto de implementación del proyecto",
                    placeholder="Ejemplo: Regional Huila, Centro de Formación Agroindustrial, sector cafetero del Huila, centros de formación SENA, etc.",
                    height=110
                )

            with col_b:
                descripcion_producto_lc = st.text_area(
                    "Descripción del prototipo, producto o servicio",
                    placeholder="Describe qué es el producto, servicio, prototipo o desarrollo, qué hace, cómo funciona y a quién podría servir.",
                    height=145
                )

                aspectos_generacion_lc = st.text_area(
                    "Aspectos a tener en cuenta para la generación del modelo de negocio",
                    placeholder="Ejemplo: diferenciadores, restricciones, canales deseados, aliados, precios estimados, forma de adopción, soporte esperado, etc.",
                    height=120
                )

                usuarios_objetivo_lc = st.text_area(
                    "Clientes, usuarios o beneficiarios objetivo (opcional)",
                    placeholder="Ejemplo: centros de formación, productores cafeteros, emprendedores rurales, empresas del sector, aprendices, clientes finales, etc.",
                    height=85
                )

            with st.expander("Datos opcionales para el encabezado del informe"):
                col_opt1, col_opt2, col_opt3 = st.columns(3)
                with col_opt1:
                    nombre_talento_lc = st.text_input("Talento que realiza el informe", placeholder="Nombre del talento", key="lc_nombre_talento")
                with col_opt2:
                    nombre_experto_lc = st.text_input("Gestor o experto del proyecto", placeholder="Nombre del experto", key="lc_nombre_experto")
                with col_opt3:
                    linea_tecnologica_lc = st.text_input("Línea de desarrollo tecnológico", placeholder="Ejemplo: Ingeniería y Diseño", key="lc_linea")

            generar_lean_canvas = st.form_submit_button("Generar Informe de Modelo de Negocios")

        if generar_lean_canvas:
            errores = []

            campos_obligatorios = {
                "Código del proyecto": codigo_proyecto_lc,
                "Nombre del proyecto": nombre_proyecto_lc,
                "Región o contexto de implementación": region_contexto_lc,
                "Descripción del prototipo, producto o servicio": descripcion_producto_lc,
                "Aspectos a tener en cuenta": aspectos_generacion_lc,
            }

            for campo, valor in campos_obligatorios.items():
                if not str(valor).strip():
                    errores.append(campo)

            if errores:
                st.error("Faltan campos obligatorios: " + ", ".join(errores))
                st.stop()

            progreso = st.progress(0)
            estado = st.empty()

            estado.info("Generando los 9 cuadros del modelo de negocios...")
            progreso.progress(20)

            try:
                if modo_prueba:
                    contenido_lean_canvas = generar_modelo_negocio_modo_prueba(
                        nombre_proyecto_lc,
                        codigo_proyecto_lc,
                        descripcion_producto_lc,
                        aspectos_generacion_lc,
                        region_contexto_lc,
                        usuarios_objetivo_lc,
                        trl_nivel_lc,
                    )
                else:
                    contenido_lean_canvas = generar_modelo_negocio_con_chatgpt(
                        nombre_proyecto_lc,
                        codigo_proyecto_lc,
                        descripcion_producto_lc,
                        aspectos_generacion_lc,
                        region_contexto_lc,
                        usuarios_objetivo_lc,
                        trl_nivel_lc,
                        modelo_openai,
                    )
            except Exception as e:
                st.warning(f"No se pudo usar IA. Se generará una versión base. Detalle: {e}")
                contenido_lean_canvas = generar_lean_canvas_modo_prueba(
                    nombre_proyecto_lc,
                    codigo_proyecto_lc,
                    descripcion_producto_lc,
                    aspectos_generacion_lc,
                    region_contexto_lc,
                    usuarios_objetivo_lc,
                    trl_nivel_lc,
                )

            progreso.progress(65)
            estado.info("Construyendo informe PDF del modelo de negocios...")

            datos_lean_canvas = {
                "tipo_documento": "Modelo de negocio Lean Canvas",
                "codigo_proyecto": codigo_proyecto_lc,
                "nombre_proyecto": nombre_proyecto_lc,
                "trl_nivel": trl_nivel_lc,
                "descripcion_producto": descripcion_producto_lc,
                "aspectos_generacion": aspectos_generacion_lc,
                "region_contexto": region_contexto_lc,
                "usuarios_objetivo": usuarios_objetivo_lc,
                "nombre_talento": nombre_talento_lc,
                "nombre_experto": nombre_experto_lc,
                "linea_tecnologica": linea_tecnologica_lc,
                "contenido_lean_canvas": contenido_lean_canvas,
                "modo_generacion": "Prueba local" if modo_prueba else "ChatGPT API",
            }

            try:
                ruta_pdf = generar_pdf_lean_canvas(datos_lean_canvas)
                st.session_state.datos_lean_canvas_generado = datos_lean_canvas
                st.session_state.ruta_pdf_lean_canvas_generado = ruta_pdf
                progreso.progress(100)
                estado.success("Informe de modelo de negocios generado correctamente.")
            except Exception as e:
                progreso.empty()
                estado.empty()
                st.error(f"No se pudo generar el PDF: {e}")
                st.stop()

        if st.session_state.datos_lean_canvas_generado:
            datos = st.session_state.datos_lean_canvas_generado

            st.markdown("## Resumen para validación")
            st.write("**Código del proyecto:**", datos["codigo_proyecto"])
            st.write("**Nombre del proyecto:**", datos["nombre_proyecto"])
            st.write("**Nivel TRL:**", datos.get("trl_nivel", "TRL seleccionado"))
            st.write("**Modo de generación:**", datos["modo_generacion"])

            st.markdown("### Cuadros generados")
            for item in ITEMS_LEAN_CANVAS:
                with st.expander(item):
                    st.write(datos["contenido_lean_canvas"].get(item, ""))

            col_json, col_pdf = st.columns(2)

            with col_json:
                st.download_button(
                    label="Descargar datos en JSON",
                    data=json.dumps(datos, ensure_ascii=False, indent=4),
                    file_name="datos_lean_canvas.json",
                    mime="application/json"
                )

            with col_pdf:
                ruta_pdf = st.session_state.ruta_pdf_lean_canvas_generado

                if ruta_pdf and Path(ruta_pdf).exists():
                    with open(ruta_pdf, "rb") as f:
                        st.download_button(
                            label="⬇️ Descargar PDF Modelo de Negocios",
                            data=f,
                            file_name=Path(ruta_pdf).name,
                            mime="application/pdf"
                        )

        st.stop()

    if st.session_state.documento_seleccionado == "informe_tecnico_final":
        st.markdown("---")
        st.subheader("Informe Técnico Final")
        st.info(
            "Este módulo genera un documento Word institucional con apartados técnicos, tabla estimativa de costos, "
            "evidencias sugeridas y entregables del proyecto."
        )

        tipo_proyecto = st.selectbox(
            "Tipo de proyecto",
            options=TIPOS_PROYECTO_INFORME,
            key="itf_tipo_proyecto",
        )
        tipo_otro = ""
        if tipo_proyecto == "Otro":
            tipo_otro = st.text_input(
                "Describe el tipo de proyecto",
                placeholder="Ejemplo: Solución de realidad aumentada aplicada al turismo",
                key="itf_tipo_otro",
            )

        metodologia = st.selectbox(
            "Metodología",
            options=METODOLOGIAS_INFORME,
            key="itf_metodologia",
        )
        metodologia_otra = ""
        if metodologia == "Otra":
            metodologia_otra = st.text_input(
                "Describe la metodología aplicada",
                placeholder="Nombre o enfoque metodológico utilizado",
                key="itf_metodologia_otra",
            )

        with st.form("form_informe_tecnico_final"):
            col_a, col_b = st.columns(2)

            with col_a:
                nombre_proyecto_itf = st.text_area(
                    "Nombre completo del proyecto",
                    placeholder="Nombre oficial del proyecto",
                    height=82,
                )
                codigo_proyecto_itf = st.text_input(
                    "Código del proyecto",
                    placeholder="Ejemplo: P2026-143440-00001",
                )
                nombre_talento_itf = st.text_input(
                    "Nombre del talento o beneficiario",
                    placeholder="Nombre completo",
                )
                nombre_experto_itf = st.text_input(
                    "Nombre del experto o asesor Tecnoparque",
                    placeholder="Nombre completo",
                )
                linea_tecnologica_itf = st.text_input(
                    "Línea tecnológica del proyecto",
                    placeholder="Ejemplo: Ingeniería y diseño",
                )

            with col_b:
                costo_total_itf = st.number_input(
                    "Costo total estimado del prototipo o desarrollo (COP)",
                    min_value=0,
                    value=0,
                    step=10000,
                    help="Solo ingresa el valor total. El sistema generará la distribución estimativa.",
                )
                contexto_itf = st.text_area(
                    "Contexto, necesidad o problema del proyecto",
                    placeholder="Describe el origen, necesidad atendida y contexto del desarrollo.",
                    height=145,
                )
                innovacion_itf = st.text_area(
                    "Tipo de innovación y valor diferencial",
                    placeholder="Explica el aporte innovador y qué diferencia al resultado obtenido.",
                    height=120,
                )

            aplicacion_metodologia_itf = st.text_area(
                "Descripción de cómo se aplicó la metodología",
                placeholder="Describe las etapas, decisiones y actividades efectivamente realizadas.",
                height=150,
            )
            producto_final_itf = st.text_area(
                "Descripción del producto final entregado",
                placeholder="Describe el prototipo, producto, desarrollo o documento final entregado.",
                height=145,
            )
            pruebas_itf = st.text_area(
                "Pruebas y validación realizadas",
                placeholder="Describe únicamente pruebas o validaciones realmente realizadas.",
                height=145,
            )
            evidencias_itf = st.text_area(
                "Evidencias disponibles para el informe",
                placeholder="Lista soportes disponibles: capturas, planos, enlaces, fotografías, actas, archivos, resultados, etc.",
                height=145,
            )
            entregables_itf = st.text_area(
                "Entregables del proyecto",
                placeholder="Escribe un entregable por línea: producto, documento, archivo, prototipo o resultado entregado.",
                height=145,
            )
            aceptacion_itf = st.checkbox(
                "Confirmo que el prototipo, producto o desarrollo final fue presentado, validado y aceptado por el talento beneficiario."
            )
            generar_informe_itf = st.form_submit_button("Generar Informe Técnico Final en Word")

        if generar_informe_itf:
            tipo_detallado = tipo_otro.strip() if tipo_proyecto == "Otro" else tipo_proyecto
            metodologia_detallada = metodologia_otra.strip() if metodologia == "Otra" else metodologia
            errores = []
            campos_obligatorios = {
                "Nombre completo del proyecto": nombre_proyecto_itf,
                "Código del proyecto": codigo_proyecto_itf,
                "Nombre del talento o beneficiario": nombre_talento_itf,
                "Nombre del experto o asesor Tecnoparque": nombre_experto_itf,
                "Línea tecnológica": linea_tecnologica_itf,
                "Tipo de proyecto": tipo_detallado,
                "Contexto, necesidad o problema": contexto_itf,
                "Metodología": metodologia_detallada,
                "Aplicación de la metodología": aplicacion_metodologia_itf,
                "Tipo de innovación y valor diferencial": innovacion_itf,
                "Producto final entregado": producto_final_itf,
                "Pruebas y validación": pruebas_itf,
                "Evidencias disponibles": evidencias_itf,
                "Entregables del proyecto": entregables_itf,
            }
            for campo, valor in campos_obligatorios.items():
                if not str(valor).strip():
                    errores.append(campo)
            if int(costo_total_itf) <= 0:
                errores.append("Costo total estimado del prototipo o desarrollo")
            if errores:
                st.error("Faltan datos obligatorios o inválidos: " + ", ".join(errores))
                st.stop()

            datos_itf = {
                "tipo_documento": "Informe Técnico Final",
                "nombre_proyecto": nombre_proyecto_itf.strip(),
                "codigo_proyecto": codigo_proyecto_itf.strip(),
                "nombre_talento": nombre_talento_itf.strip(),
                "nombre_experto": nombre_experto_itf.strip(),
                "linea_tecnologica": linea_tecnologica_itf.strip(),
                "tipo_proyecto": tipo_proyecto,
                "tipo_proyecto_detallado": tipo_detallado,
                "contexto_proyecto": contexto_itf.strip(),
                "metodologia": metodologia,
                "metodologia_detallada": metodologia_detallada,
                "aplicacion_metodologia": aplicacion_metodologia_itf.strip(),
                "innovacion_valor": innovacion_itf.strip(),
                "producto_final": producto_final_itf.strip(),
                "costo_total": int(costo_total_itf),
                "pruebas_validacion": pruebas_itf.strip(),
                "evidencias_disponibles": evidencias_itf.strip(),
                "entregables": entregables_itf.strip(),
                "aceptacion_confirmada": bool(aceptacion_itf),
                "modo_generacion": "Prueba local" if modo_prueba else "ChatGPT API",
            }
            apartados = [
                ("Introducción, Contexto y Antecedentes", 400),
                ("Metodología", 600),
                ("Normatividad", 400),
                ("Análisis y Estimación de Costos del Proyecto", 400),
                ("Pruebas Documentadas y Validación del Prototipo", 400),
                ("Entregables", 400),
                ("Análisis y Conclusiones", 400),
            ]
            progreso = st.progress(0, text="Preparando generación del informe...")
            contenido_generado = {}
            try:
                for indice, (titulo_apartado, minimo) in enumerate(apartados, start=1):
                    progreso.progress(
                        int((indice - 1) / (len(apartados) + 2) * 100),
                        text=f"Generando apartado: {titulo_apartado}...",
                    )
                    if modo_prueba:
                        contenido_generado[titulo_apartado] = generar_apartado_informe_modo_prueba(
                            titulo_apartado, datos_itf, minimo
                        )
                    else:
                        contenido_generado[titulo_apartado] = generar_apartado_informe_con_chatgpt(
                            titulo_apartado, datos_itf, minimo, modelo_openai
                        )
                progreso.progress(82, text="Construyendo tabla estimativa de costos...")
                if modo_prueba:
                    tabla_costos = generar_tabla_costos_modo_prueba(datos_itf)
                else:
                    try:
                        tabla_costos = generar_tabla_costos_con_chatgpt(datos_itf, modelo_openai)
                    except Exception:
                        tabla_costos = generar_tabla_costos_modo_prueba(datos_itf)
                datos_itf["contenido_generado"] = contenido_generado
                datos_itf["tabla_costos"] = tabla_costos
                datos_itf["tabla_entregables"] = generar_tabla_entregables(datos_itf)
                progreso.progress(92, text="Generando documento Word...")
                ruta_docx = generar_docx_informe_tecnico_final(datos_itf)
                st.session_state.datos_informe_tecnico_final_generado = datos_itf
                st.session_state.ruta_docx_informe_tecnico_final_generado = ruta_docx
                progreso.progress(100, text="Informe Técnico Final generado correctamente.")
                st.success("Informe Técnico Final generado correctamente en formato Word.")
            except Exception as e:
                st.error(f"No se pudo generar el Informe Técnico Final: {e}")
                st.stop()

        if st.session_state.datos_informe_tecnico_final_generado:
            datos = st.session_state.datos_informe_tecnico_final_generado
            st.markdown("## Resumen para validación")
            st.write("**Proyecto:**", datos.get("nombre_proyecto", ""))
            st.write("**Código:**", datos.get("codigo_proyecto", ""))
            st.write("**Tipo:**", datos.get("tipo_proyecto_detallado", ""))
            st.write("**Metodología:**", datos.get("metodologia_detallada", ""))
            st.write("**Costo total estimado:**", formato_moneda_colombiana(datos.get("costo_total", 0)))
            st.write("**Aceptación confirmada:**", "Sí" if datos.get("aceptacion_confirmada") else "No")
            st.caption("La tabla de costos se presenta como estimación técnica, no como información contable certificada.")
            if st.session_state.ruta_docx_informe_tecnico_final_generado and Path(st.session_state.ruta_docx_informe_tecnico_final_generado).exists():
                ruta_docx = st.session_state.ruta_docx_informe_tecnico_final_generado
                with open(ruta_docx, "rb") as archivo_docx:
                    st.download_button(
                        label="⬇️ Descargar Informe Técnico Final en Word",
                        data=archivo_docx,
                        file_name=Path(ruta_docx).name,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
        st.stop()

    if st.session_state.documento_seleccionado == "acta_cierre_ficha":
        st.markdown("---")
        st.subheader("Formulario para Acta de Cierre")

        st.info(
            "Este módulo genera el Acta de Cierre del proyecto. "
            "La ficha de caracterización quedará integrada en este mismo flujo cuando se definan sus campos."
        )

        with st.form("form_acta_cierre"):
            col_a, col_b = st.columns(2)

            with col_a:
                codigo_proyecto = st.text_input(
                    "Código del proyecto",
                    placeholder="Ejemplo: P2025-143440-00001"
                )

                nombre_proyecto = st.text_area(
                    "Nombre del proyecto",
                    placeholder="Nombre oficial del proyecto",
                    height=90
                )

                fecha_acta = st.date_input(
                    "Fecha del acta",
                    value=date.today()
                )

                hora_inicio = st.time_input(
                    "Hora de inicio",
                    value=time(8, 0)
                )

            with col_b:
                nombre_talento = st.text_input(
                    "Nombre del talento",
                    placeholder="Nombre completo del talento"
                )

                nombre_experto = st.text_input(
                    "Nombre del experto",
                    placeholder="Nombre completo del experto"
                )

                trl_obtenido = st.selectbox(
                    "TRL obtenido",
                    options=["TRL 6", "TRL 7", "TRL 8"]
                )

                aporte_tecnoparque = st.number_input(
                    "Aporte estimado de Tecnoparque",
                    min_value=0,
                    value=0,
                    step=10000
                )
            objetivos_iniciales = st.text_area(
                "Objetivos iniciales del proyecto",
                placeholder="Copia aquí los objetivos iniciales del proyecto. Cada objetivo puede ir en una línea diferente.",
                height=140
)
           
            evidencias_producto = st.text_area(
                "Evidencias del Producto",
                placeholder=(
                    "Describe de forma general los productos, prototipos, documentos, pruebas, diseños, "
                    "implementaciones o entregables generados. Si aplica, menciona si el proyecto es susceptible "
                    "de inscribir un nuevo proyecto."
                ),
                height=180
            )

            generar_acta_cierre = st.form_submit_button("Generar Acta de Cierre")

        if generar_acta_cierre:
            errores = []

            campos_obligatorios = {
                "Código del proyecto": codigo_proyecto,
                "Nombre del proyecto": nombre_proyecto,
                "Nombre del talento": nombre_talento,
                "Nombre del experto": nombre_experto,
                "Objetivos iniciales del proyecto": objetivos_iniciales,
                "Evidencias del Producto": evidencias_producto,
            }

            for campo, valor in campos_obligatorios.items():
                if not str(valor).strip():
                    errores.append(campo)

            if errores:
                st.error("Faltan campos obligatorios: " + ", ".join(errores))
                st.stop()

            hora_fin = calcular_hora_fin(fecha_acta, hora_inicio)
            objetivo_cierre = generar_objetivo_cierre(codigo_proyecto, nombre_proyecto)

            objetivos_lista = [
                obj.strip()
                for obj in objetivos_iniciales.replace("•", "\n").replace(";", "\n").split("\n")
                if obj.strip()
            ]

            with st.spinner("Generando evidencias del acta de cierre con IA..."):
                try:
                    if modo_prueba:
                        evidencias_generadas = generar_evidencias_cierre_modo_prueba(
                            nombre_proyecto,
                            evidencias_producto
                        )
                    else:
                        evidencias_generadas = generar_evidencias_cierre_con_chatgpt(
                            nombre_proyecto,
                            evidencias_producto,
                            modelo_openai
                        )
                except Exception as e:
                    st.warning(f"No se pudo usar IA. Se generará una versión base. Detalle: {e}")
                    evidencias_generadas = generar_evidencias_cierre_modo_prueba(
                        nombre_proyecto,
                        evidencias_producto
                    )

            datos_acta_cierre = {
                "tipo_documento": "Acta de Cierre",
                "titulo_acta": f"ACTA No. 03 del proyecto No {codigo_proyecto}",
                "codigo_proyecto": codigo_proyecto,
                "nombre_proyecto": nombre_proyecto,
                "fecha_acta": fecha_acta.strftime("%d/%m/%Y"),
                "fecha_iso": fecha_acta.strftime("%Y-%m-%d"),
                "hora_inicio": hora_inicio.strftime("%H:%M"),
                "hora_fin": hora_fin.strftime("%H:%M"),
                "nombre_talento": nombre_talento,
                "nombre_experto": nombre_experto,
                "trl_obtenido": trl_obtenido,
                "aporte_tecnoparque": aporte_tecnoparque,
                "evidencias_producto": evidencias_producto,
                "objetivo_cierre": objetivo_cierre,
                "objetivos_iniciales": objetivos_lista,
                "evidencias_generadas": evidencias_generadas,
                "modo_generacion": "Prueba local" if modo_prueba else "ChatGPT API",
            }

            st.session_state.datos_acta_cierre_generada = datos_acta_cierre
            st.session_state.ruta_pdf_acta_cierre_generado = None

            st.success("Acta de cierre generada correctamente. Ahora puedes revisar y generar el PDF.")

        if st.session_state.datos_acta_cierre_generada:
            datos = st.session_state.datos_acta_cierre_generada
            evidencias = datos["evidencias_generadas"]

            st.markdown("## Resumen para validación")
            st.write("**Título:**", datos["titulo_acta"])
            st.write("**Proyecto:**", datos["nombre_proyecto"])
            st.write("**Código:**", datos["codigo_proyecto"])
            st.write("**Fecha:**", datos["fecha_acta"])
            st.write("**Hora inicio:**", datos["hora_inicio"])
            st.write("**Hora fin:**", datos["hora_fin"])
            st.write("**Talento:**", datos["nombre_talento"])
            st.write("**Experto:**", datos["nombre_experto"])
            st.write("**TRL obtenido:**", datos["trl_obtenido"])
            st.write("**Aporte Tecnoparque:**", formato_moneda_colombiana(datos["aporte_tecnoparque"]))

            st.markdown("### Objetivo de cierre")
            st.write(datos.get("objetivo_cierre", ""))

            st.markdown("### Objetivos iniciales del proyecto")
            for i, objetivo in enumerate(datos.get("objetivos_iniciales", []), start=1):
                st.write(f"{i}. SI Cumplió — {objetivo}")

            st.markdown("### Evidencias generadas")
            st.write("**Normatividad:**", evidencias.get("evidencias_normatividad", ""))
            st.write("**Modelo de negocio:**", evidencias.get("evidencias_modelo_negocio", ""))
            st.write("**Pruebas documentadas:**", evidencias.get("evidencias_pruebas_documentadas", ""))
            st.write("**Prototipo:**", evidencias.get("evidencias_prototipo", ""))

            if evidencias.get("conclusion_adicional"):
                st.write("**Conclusión adicional:**", evidencias.get("conclusion_adicional"))

            col_json, col_pdf = st.columns(2)

            with col_json:
                st.download_button(
                    label="Descargar datos en JSON",
                    data=json.dumps(datos, ensure_ascii=False, indent=4),
                    file_name="datos_acta_cierre.json",
                    mime="application/json"
                )

            with col_pdf:
                if st.button("📄 Generar PDF del Acta de Cierre"):
                    try:
                        ruta_pdf = generar_pdf_acta_cierre(datos)
                        st.session_state.ruta_pdf_acta_cierre_generado = ruta_pdf
                        st.success(f"PDF generado correctamente: {ruta_pdf}")
                    except Exception as e:
                        st.error(f"No se pudo generar el PDF: {e}")

            if (
                st.session_state.ruta_pdf_acta_cierre_generado
                and Path(st.session_state.ruta_pdf_acta_cierre_generado).exists()
            ):
                ruta_pdf = st.session_state.ruta_pdf_acta_cierre_generado

                with open(ruta_pdf, "rb") as f:
                    st.download_button(
                        label="⬇️ Descargar PDF del Acta de Cierre",
                        data=f,
                        file_name=Path(ruta_pdf).name,
                        mime="application/pdf"
                    )

        st.stop()

    st.warning("Selecciona una opción válida de la fase de cierre.")
    st.stop()